import os
import logging
import sqlite3
import re
import time
import gc
from datetime import datetime, timezone
from typing import List, Dict, Any, Optional, Iterator, Iterable
import traceback
import difflib
import json
import tracemalloc
import zipfile
import xml.etree.ElementTree as ET

from flask import request, redirect, url_for, render_template, flash, jsonify, session, render_template_string
from werkzeug.security import generate_password_hash

# Parsers are optional imports - provide helpful messages if missing
try:
    import docx
except Exception:
    docx = None

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    from academic_setup_validator import validate_staged_slots
except Exception as val_imp_err:
    validate_staged_slots = None

logger = logging.getLogger("app.timetable")

# Tunables for import batching and preview limits
BATCH_INSERT_SIZE = int(os.environ.get("TIMETABLE_BATCH_SIZE", 5))
PREVIEW_ROW_CAP = int(os.environ.get("TIMETABLE_PREVIEW_CAP", 0))
SKIPPED_ROW_SAMPLE_CAP = int(os.environ.get("TIMETABLE_SKIPPED_SAMPLE_CAP", 5))
TIMETABLE_MAX_TABLES = int(os.environ.get("TIMETABLE_MAX_TABLES", 1))
TIMETABLE_SINGLE_SECTION_ONLY = os.environ.get("TIMETABLE_SINGLE_SECTION_ONLY", "false").strip().lower() in ("1", "true", "yes", "on")
ENABLE_IMPORT_TRACEMALLOC = os.environ.get("TIMETABLE_ENABLE_TRACEMALLOC", "false").strip().lower() in ("1", "true", "yes", "on")
PDF_DIAG_SAMPLE_CAP = int(os.environ.get("TIMETABLE_PDF_DIAG_SAMPLE_CAP", 12))
PDF_TABLE_SETTINGS = (
    {
        "vertical_strategy": "lines",
        "horizontal_strategy": "lines",
        "intersection_tolerance": 5,
        "snap_tolerance": 3,
        "join_tolerance": 3,
        "edge_min_length": 3,
        "min_words_vertical": 1,
        "min_words_horizontal": 1,
    },
    {
        "vertical_strategy": "text",
        "horizontal_strategy": "text",
        "intersection_tolerance": 5,
        "snap_tolerance": 3,
        "join_tolerance": 3,
        "edge_min_length": 3,
        "min_words_vertical": 1,
        "min_words_horizontal": 1,
    },
)


class TimetablePDFValidationError(ValueError):
    pass


def _append_skipped_sample(skipped_rows: List[Dict], skipped_rows_omitted: int, payload: Dict) -> int:
    if len(skipped_rows) < SKIPPED_ROW_SAMPLE_CAP:
        skipped_rows.append(payload)
        return skipped_rows_omitted
    return skipped_rows_omitted + 1

# Database helper functions - use existing app.get_db() pattern where called from app.py

def _is_postgres_db(db) -> bool:
    try:
        name = type(db).__name__
        if name == "_PostgresDB":
            return True
        if hasattr(db, "_conn"):
            mod = type(getattr(db, "_conn")).__module__
            if "psycopg2" in mod:
                return True
        mod = type(db).__module__
        if "psycopg2" in mod:
            return True
        return False
    except Exception:
        return False


def _db_execute(db, query, params=()):
    """Execute a query using the DB connection or cursor.

    - Prefer PostgreSQL-style `%s` placeholders in code.
    - If running against SQLite, convert `%s` -> `?` so queries remain compatible.
    - If the module uses `?` placeholders, convert them to `%s` for Postgres.
    """
    try:
        is_pg = _is_postgres_db(db)
        if is_pg:
            if "?" in query:
                query = query.replace("?", "%s")
        else:
            if "%s" in query:
                query = query.replace("%s", "?")

        if hasattr(db, "execute") and callable(getattr(db, "execute")):
            return db.execute(query, params)
        elif hasattr(db, "cursor") and callable(getattr(db, "cursor")):
            cur = db.cursor()
            cur.execute(query, params)
            return cur
        else:
            raise AttributeError(f"{type(db).__name__} object has no attribute execute or cursor")
    except Exception:
        raise


def _create_tables_sql(db):
    if _is_postgres_db(db):
        return [
            """
            CREATE TABLE IF NOT EXISTS timetable_slots (
                id SERIAL PRIMARY KEY,
                branch TEXT NOT NULL,
                section TEXT NOT NULL,
                semester INTEGER,
                day TEXT NOT NULL,
                start_time TEXT NOT NULL,
                end_time TEXT NOT NULL,
                subject_name TEXT NOT NULL,
                faculty_name TEXT,
                is_lab INTEGER DEFAULT 0,
                room TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            );
            """,
        ]
    else:
        return [
            """
            CREATE TABLE IF NOT EXISTS timetable_slots (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                branch TEXT NOT NULL,
                section TEXT NOT NULL,
                semester INTEGER,
                day TEXT NOT NULL,
                start_time TEXT NOT NULL,
                end_time TEXT NOT NULL,
                subject_name TEXT NOT NULL,
                faculty_name TEXT,
                is_lab INTEGER DEFAULT 0,
                room TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            );
            """,
        ]


def _create_normalized_sql(db):
    if _is_postgres_db(db):
        return [
            """
            CREATE TABLE IF NOT EXISTS timetable_entries (
                id SERIAL PRIMARY KEY,
                branch_id INTEGER,
                section TEXT,
                semester INTEGER,
                day TEXT,
                start_time TEXT,
                end_time TEXT,
                subject_id INTEGER,
                teacher_id INTEGER,
                subject_name TEXT,
                faculty_name TEXT,
                is_lab INTEGER DEFAULT 0,
                room TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            );
            """,
        ]
    else:
        return [
            """
            CREATE TABLE IF NOT EXISTS timetable_entries (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                branch_id INTEGER,
                section TEXT,
                semester INTEGER,
                day TEXT,
                start_time TEXT,
                end_time TEXT,
                subject_id INTEGER,
                teacher_id INTEGER,
                subject_name TEXT,
                faculty_name TEXT,
                is_lab INTEGER DEFAULT 0,
                room TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            );
            """,
        ]


def ensure_timetable_tables(db):
    for sql in _create_tables_sql(db):
        try:
            db.execute(sql)
        except Exception:
            # ignore individual table creation errors; caller will handle
            logger.exception("create table failed")
    for sql in _create_normalized_sql(db):
        try:
            db.execute(sql)
        except Exception:
            logger.exception("create normalized table failed")
    try:
        db.commit()
    except Exception:
        pass

    # Add DB-level duplicate protection. If legacy duplicates exist, continue with warnings.
    unique_index_sql = [
        "CREATE UNIQUE INDEX IF NOT EXISTS uq_timetable_slots_dedupe ON timetable_slots (branch, section, semester, day, start_time, end_time, subject_name, faculty_name, room)",
        "CREATE UNIQUE INDEX IF NOT EXISTS uq_timetable_entries_dedupe ON timetable_entries (branch_id, section, semester, day, start_time, end_time, subject_id, teacher_id, room)",
    ]
    for sql in unique_index_sql:
        try:
            _db_execute(db, sql)
        except Exception as e:
            logger.warning("Unique index creation skipped: %s", repr(e))
    try:
        db.commit()
    except Exception:
        pass

    # Backfill text fallback columns if an older schema is already present.
    try:
        cols = set()
        if _is_postgres_db(db):
            rows = _db_execute(db, "SELECT column_name FROM information_schema.columns WHERE table_schema = 'public' AND table_name = 'timetable_entries'").fetchall()
            cols = {_clean_text(row[0] if not hasattr(row, 'keys') else row['column_name']).lower() for row in rows}
        else:
            rows = _db_execute(db, "PRAGMA table_info(timetable_entries)").fetchall()
            cols = {_clean_text(row[1] if not hasattr(row, 'keys') else row['name']).lower() for row in rows}
        if 'subject_name' not in cols:
            _db_execute(db, "ALTER TABLE timetable_entries ADD COLUMN subject_name TEXT")
        if 'faculty_name' not in cols:
            _db_execute(db, "ALTER TABLE timetable_entries ADD COLUMN faculty_name TEXT")
        try:
            db.commit()
        except Exception:
            pass
    except Exception:
        logger.exception("Failed to ensure fallback timetable columns")


def _clean_text(value) -> str:
    return (str(value).strip() if value is not None else "")


def _normalize_display_text(value: str) -> str:
    text = _clean_text(value)
    if not text:
        return ""
    text = re.sub(r"[‐‑‒–—―]", "-", text)
    text = re.sub(r"\s+", " ", text)
    text = text.replace("_", " ")
    text = text.strip(" ,;:/|-")
    if not text:
        return ""
    parts = []
    for token in text.split(" "):
        if not token:
            continue
        if re.fullmatch(r"[A-Z0-9]{2,6}", token):
            parts.append(token.upper())
            continue
        if re.fullmatch(r"[A-Z]\.", token):
            parts.append(token.upper())
            continue
        if "-" in token:
            parts.append("-".join(piece[:1].upper() + piece[1:].lower() if piece and not re.fullmatch(r"[A-Z0-9]{2,6}", piece) else piece.upper() for piece in token.split("-")))
            continue
        parts.append(token[:1].upper() + token[1:].lower())
    return re.sub(r"\s+", " ", " ".join(parts)).strip()


def _normalize_timetable_branch_name(value: str, row_text: str = "") -> str:
    text = _clean_text(value)
    if not text:
        text = _clean_text(row_text)
    if not text or _contains_blocked_timetable_text(text):
        return ""
    text = re.sub(r"[‐‑‒–—―]", "-", text)
    text = re.sub(r"[\s_/]+", " ", text).strip()
    match = re.search(r"\b(CSM|CSE|ECE|EEE|IT|MECH|CIVIL|AIML|AIDS|DS)\b", text, flags=re.IGNORECASE)
    if match:
        return match.group(1).upper()
    normalized = _normalize_academic_department_code(text.replace(" ", ""))
    if normalized:
        return normalized
    return ""


def split_branch_section(value: str) -> tuple[str, str]:
    """Split branch-section values like CSE-A, cse a, or CSEA."""
    text = _clean_text(value)
    if not text:
        return "", ""
    text = re.sub(r"[\s._]+", "-", text)
    text = re.sub(r"-+", "-", text).strip("-")
    parts = [part for part in re.split(r"[-/]+", text) if part]
    if len(parts) >= 2:
        return parts[0].upper(), parts[-1].upper()
    match = re.fullmatch(r"([A-Za-z]{2,5})([A-Za-z0-9]{1,4})", text)
    if match:
        return match.group(1).upper(), match.group(2).upper()
    return text.upper(), ""


def _normalize_timetable_section_name(value: str, branch_value: str = "", row_text: str = "") -> str:
    text = _clean_text(value)
    if not text:
        text = _clean_text(row_text)
    if not text or _contains_blocked_timetable_text(text):
        return ""
    text = re.sub(r"[‐‑‒–—―]", "-", text)
    text = re.sub(r"[\s_/]+", "-", text)
    text = re.sub(r"-+", "-", text).strip("-")
    if not text:
        return ""

    branch_part, section_part = split_branch_section(text)
    if branch_part and section_part and branch_part in _ACADEMIC_DEPARTMENT_CODES:
        return f"{branch_part}-{section_part}"

    match = re.fullmatch(r"(CSM|CSE|ECE|EEE|IT|MECH|CIVIL|AIML|AIDS|DS)[- ]([A-Z0-9]{1,4})", text, flags=re.IGNORECASE)
    if match:
        return f"{match.group(1).upper()}-{match.group(2).upper()}"

    branch = _normalize_timetable_branch_name(branch_value)
    if branch and re.fullmatch(r"[A-Z0-9]{1,4}", text, flags=re.IGNORECASE):
        return f"{branch}-{text.upper()}"
    if branch and re.fullmatch(r"(?:I|II|III|IV|V|VI|VII|VIII|IX|X)(?:-\d{1,2})?", text, flags=re.IGNORECASE):
        return f"{branch}-{text.upper()}"

    generic = re.fullmatch(r"([A-Z0-9]{1,8})", text, flags=re.IGNORECASE)
    if generic:
        return generic.group(1).upper()
    return text.upper()


def _normalize_timetable_branch_section(branch_value: str, section_value: str = "", row_text: str = "") -> tuple[str, str]:
    """Normalize branch/section pairs from either separate or combined labels."""
    branch_text = _clean_text(branch_value)
    section_text = _clean_text(section_value)

    branch_part, section_part = split_branch_section(branch_text or section_text or row_text)
    if branch_part in _ACADEMIC_DEPARTMENT_CODES and section_part:
        branch_text = branch_part
        section_text = f"{branch_part}-{section_part}"

    branch = _normalize_timetable_branch_name(branch_text, row_text=row_text)
    if not branch and branch_part in _ACADEMIC_DEPARTMENT_CODES:
        branch = branch_part

    section = _normalize_timetable_section_name(section_text, branch_value=branch, row_text=row_text)
    if not section and branch and section_part:
        section = f"{branch}-{section_part}"
    if not section and branch and branch_text and branch_text != branch:
        section = _normalize_timetable_section_name(branch_text, branch_value=branch, row_text=row_text)
    return branch, section


def _normalize_timetable_faculty_name(value: str) -> str:
    text = _clean_text(value)
    if not text or _contains_blocked_timetable_text(text):
        return ""
    text = re.sub(r"[‐‑‒–—―]", "-", text)
    text = re.sub(r"[\s_/]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return _normalize_display_text(text)


def _normalize_timetable_lab_theory(value: str, row_text: str = "") -> str:
    text = _clean_text(value)
    haystack = f"{text} {row_text}".strip()
    if _row_has_token(haystack, "lab", "practical"):
        return "LAB"
    if _row_has_token(haystack, "theory"):
        return "THEORY"
    return _normalize_display_text(text)


def _dup_key(*parts) -> str:
    return "|".join(_clean_text(p) for p in parts)


def _timetable_semantic_key(row: Dict[str, str]) -> str:
    return _dup_key(
        _clean_text(row.get("day")),
        _clean_text(row.get("start_time")),
        _clean_text(row.get("end_time")),
        _clean_text(row.get("branch")),
        _clean_text(row.get("section")),
        _clean_text(row.get("subject_name")),
        _clean_text(row.get("faculty_name")),
    )


def _insert_ignore_sql(db, table: str, columns: List[str]) -> str:
    column_list = ", ".join(columns)
    placeholder_list = ", ".join(["%s"] * len(columns))
    if _is_postgres_db(db):
        return f"INSERT INTO {table} ({column_list}) VALUES ({placeholder_list}) ON CONFLICT DO NOTHING"
    return f"INSERT OR IGNORE INTO {table} ({column_list}) VALUES ({placeholder_list})"


def _table_diagnostics(db):
    """Return duplicate/null diagnostics for timetable tables."""
    result = {}
    try:
        def _safe_fetch(sql):
            try:
                return _db_execute(db, sql).fetchone()
            except Exception:
                return None

        slots_dup = _safe_fetch(
            """
            SELECT COUNT(1) AS c
            FROM (
                SELECT branch, section, day, start_time, end_time, subject_name, faculty_name, room
                FROM timetable_slots
                GROUP BY branch, section, day, start_time, end_time, subject_name, faculty_name, room
                HAVING COUNT(1) > 1
            ) dup
            """
        )
        slots_dup_rows = _safe_fetch(
            """
            SELECT COALESCE(SUM(c), 0) AS c FROM (
                SELECT COUNT(1) AS c
                FROM timetable_slots
                GROUP BY branch, section, day, start_time, end_time, subject_name, faculty_name, room
                HAVING COUNT(1) > 1
            ) dup
            """
        )
        entries_dup = _safe_fetch(
            """
            SELECT COUNT(1) AS c
            FROM (
                SELECT branch_id, section, semester, day, start_time, end_time, subject_id, teacher_id, room
                FROM timetable_entries
                GROUP BY branch_id, section, semester, day, start_time, end_time, subject_id, teacher_id, room
                HAVING COUNT(1) > 1
            ) dup
            """
        )
        entries_dup_rows = _safe_fetch(
            """
            SELECT COALESCE(SUM(c), 0) AS c FROM (
                SELECT COUNT(1) AS c
                FROM timetable_entries
                GROUP BY branch_id, section, semester, day, start_time, end_time, subject_id, teacher_id, room
                HAVING COUNT(1) > 1
            ) dup
            """
        )

        slots_nulls = _safe_fetch(
            """
            SELECT
              COALESCE(SUM(CASE WHEN branch IS NULL OR TRIM(branch) = '' THEN 1 ELSE 0 END), 0) AS branch_null,
              COALESCE(SUM(CASE WHEN section IS NULL OR TRIM(section) = '' THEN 1 ELSE 0 END), 0) AS section_null,
              COALESCE(SUM(CASE WHEN day IS NULL OR TRIM(day) = '' THEN 1 ELSE 0 END), 0) AS day_null,
              COALESCE(SUM(CASE WHEN start_time IS NULL OR TRIM(start_time) = '' THEN 1 ELSE 0 END), 0) AS start_null,
              COALESCE(SUM(CASE WHEN end_time IS NULL OR TRIM(end_time) = '' THEN 1 ELSE 0 END), 0) AS end_null,
              COALESCE(SUM(CASE WHEN subject_name IS NULL OR TRIM(subject_name) = '' THEN 1 ELSE 0 END), 0) AS subject_null,
              COALESCE(SUM(CASE WHEN faculty_name IS NULL OR TRIM(faculty_name) = '' THEN 1 ELSE 0 END), 0) AS faculty_null,
              COALESCE(SUM(CASE WHEN room IS NULL OR TRIM(room) = '' THEN 1 ELSE 0 END), 0) AS room_null
            FROM timetable_slots
            """
        )
        entries_nulls = _safe_fetch(
            """
            SELECT
              COALESCE(SUM(CASE WHEN branch_id IS NULL THEN 1 ELSE 0 END), 0) AS branch_id_null,
              COALESCE(SUM(CASE WHEN section IS NULL OR TRIM(section) = '' THEN 1 ELSE 0 END), 0) AS section_null,
              COALESCE(SUM(CASE WHEN semester IS NULL THEN 1 ELSE 0 END), 0) AS semester_null,
              COALESCE(SUM(CASE WHEN day IS NULL OR TRIM(day) = '' THEN 1 ELSE 0 END), 0) AS day_null,
              COALESCE(SUM(CASE WHEN start_time IS NULL OR TRIM(start_time) = '' THEN 1 ELSE 0 END), 0) AS start_null,
              COALESCE(SUM(CASE WHEN end_time IS NULL OR TRIM(end_time) = '' THEN 1 ELSE 0 END), 0) AS end_null,
              COALESCE(SUM(CASE WHEN subject_id IS NULL THEN 1 ELSE 0 END), 0) AS subject_null,
              COALESCE(SUM(CASE WHEN teacher_id IS NULL THEN 1 ELSE 0 END), 0) AS teacher_null,
              COALESCE(SUM(CASE WHEN room IS NULL OR TRIM(room) = '' THEN 1 ELSE 0 END), 0) AS room_null
            FROM timetable_entries
            """
        )

        result = {
            "slots_duplicate_groups": int(slots_dup[0] if slots_dup else 0),
            "slots_duplicate_rows": int(slots_dup_rows[0] if slots_dup_rows else 0),
            "entries_duplicate_groups": int(entries_dup[0] if entries_dup else 0),
            "entries_duplicate_rows": int(entries_dup_rows[0] if entries_dup_rows else 0),
            "slots_nulls": dict(slots_nulls) if slots_nulls else {},
            "entries_nulls": dict(entries_nulls) if entries_nulls else {},
        }
    except Exception:
        logger.exception("timetable diagnostics failed")
    return result


# --- Simple parsing helpers -------------------------------------------------

def iter_docx_table_slots(path: str) -> Iterator[Dict]:
    """Yield DOCX timetable slots row-by-row to keep memory usage low."""
    if docx is None:
        raise RuntimeError("python-docx is not installed. Install with: pip install python-docx")
    doc = docx.Document(path)
    parsed_rows = 0
    skipped_rows = 0
    failed_rows = 0
    skip_tokens = _DOCX_IGNORE_ROW_TOKENS
    seen_slots = set()
    for table_index, table in enumerate(doc.tables):
        if not table.rows:
            continue
        headers = [_normalize_key(c.text) for c in table.rows[0].cells]
        for row_index, row in enumerate(table.rows[1:], start=1):
            values = [_clean_text(c.text) for c in row.cells]
            row_text = " | ".join(values)
            row_map = {headers[i]: values[i] for i in range(min(len(headers), len(values))) if headers[i]}
            try:
                normalized = _normalize_timetable_row(_normalize_slot_row(row_map, row_text=row_text), row_text=row_text)
                subject_raw = normalized["subject_name"]
                faculty_raw = normalized["faculty_name"]
                if _row_has_token(subject_raw, *skip_tokens) or _row_has_token(faculty_raw, *skip_tokens) or _row_has_token(row_text, *skip_tokens):
                    logger.info("skipped_empty_row table=%s row=%s reason=blocked_tokens text=%s", table_index, row_index, row_text)
                    skipped_rows += 1
                    continue
                if not _valid_slot_row(normalized):
                    logger.info("skipped_empty_row table=%s row=%s reason=invalid_slot slot=%s text=%s", table_index, row_index, normalized, row_text)
                    skipped_rows += 1
                    continue
                subjects = _split_subjects(subject_raw)
                if not subjects:
                    logger.info("skipped_empty_row table=%s row=%s reason=no_subjects text=%s", table_index, row_index, row_text)
                    skipped_rows += 1
                    continue
                for subject in subjects:
                    slot = dict(normalized)
                    slot["subject_name"] = _clean_text(subject)
                    slot["is_lab"] = int(bool(_row_has_token(subject, "lab", "practical") or normalized["is_lab"]))
                    dup = _timetable_semantic_key(slot)
                    if dup in seen_slots:
                        logger.info("skipped_duplicate_row table=%s row=%s dup=%s slot=%s", table_index, row_index, dup, slot)
                        skipped_rows += 1
                        continue
                    seen_slots.add(dup)
                    if not _valid_slot_row(slot):
                        logger.info("skipped_empty_row table=%s row=%s reason=invalid_split_slot slot=%s", table_index, row_index, slot)
                        skipped_rows += 1
                        continue
                    parsed_rows += 1
                    yield slot
            except Exception:
                failed_rows += 1
                logger.exception("iter_docx_table_slots failed table=%s row=%s", table_index, row_index)
    logger.info("iter_docx_table_slots summary: tables=%d parsed_rows=%d skipped_rows=%d failed_rows=%d", len(doc.tables), parsed_rows, skipped_rows, failed_rows)


def iter_docx_grid_slots(path: str) -> Iterator[Dict]:
    """Yield grid-style DOCX slots row-by-row without retaining full table structures."""
    if docx is None:
        raise RuntimeError("python-docx is not installed. Install with: pip install python-docx")
    doc = docx.Document(path)
    base = os.path.splitext(os.path.basename(path))[0]
    title_text = " ".join([p.text for p in doc.paragraphs[:3] if p.text])
    inferred_branch = ""
    inferred_section = ""
    m = re.search(r"([A-Za-z0-9\.\s&/]+)\s+(I\s*-?\d+|I[-\s]?\d+|II|III|IV|I-\d+)", title_text, flags=re.I)
    if m:
        inferred_branch = m.group(1).strip()
        inferred_section = m.group(2).strip()
    else:
        parts = re.split(r"[_\-\s]+", base)
        if parts:
            inferred_branch = parts[0]
            if len(parts) > 1:
                inferred_section = parts[1]

    emitted = 0
    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = table.rows[0].cells
        headers = [_clean_text(c.text) for c in header_cells]
        day_col = 0
        time_cols = []
        for i, h in enumerate(headers):
            low = h.lower()
            if "day" in low:
                day_col = i
            elif re.search(r"\d", h) or "-" in h:
                time_cols.append(i)
        if not time_cols:
            time_cols = [i for i in range(len(headers)) if i != day_col]

        for row in table.rows[1:]:
            day = _clean_text(row.cells[day_col].text)
            if not day:
                continue
            for col in time_cols:
                cell_text = _clean_text(row.cells[col].text) if col < len(row.cells) else ""
                if not cell_text:
                    continue
                head = _clean_text(header_cells[col].text) if col < len(header_cells) else ""
                start_time, end_time = _split_time_range(head)
                for subj in _split_subjects(cell_text):
                    s_text = subj
                    faculty = ""
                    if "\n" in cell_text:
                        lines = [l.strip() for l in cell_text.splitlines() if l.strip()]
                        if len(lines) >= 2:
                            s_text = _clean_text(lines[0])
                            faculty = _clean_text(lines[1])
                    else:
                        parts = re.split(r"[-/\\r\\n]+", subj)
                        if len(parts) >= 2:
                            s_text = _clean_text(parts[0])
                            faculty = _clean_text(parts[1])

                    slot = {
                        "branch": inferred_branch or base,
                        "section": inferred_section or "ALL",
                        "semester": None,
                        "day": day,
                        "start_time": start_time or "",
                        "end_time": end_time or "",
                        "subject_name": s_text,
                        "faculty_name": faculty,
                        "is_lab": int(bool(_row_has_token(s_text, "lab", "practical") or _row_has_token(cell_text, "lab", "practical"))),
                        "room": "",
                    }
                    slot = _normalize_timetable_row(slot, row_text=cell_text)
                    if _valid_slot_row(slot):
                        emitted += 1
                        yield slot
    logger.info("iter_docx_grid_slots summary: tables=%d parsed_slots=%d", len(doc.tables), emitted)


def parse_docx_table(path: str) -> List[Dict]:
    """Compatibility wrapper for existing callers expecting a list."""
    return list(iter_docx_section_slots(path))


def parse_docx_grid(path: str) -> List[Dict]:
    """Compatibility wrapper for existing callers expecting a list."""
    return list(iter_docx_section_slots(path))


DOCX_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W_BODY = f"{{{DOCX_W_NS}}}body"
W_P = f"{{{DOCX_W_NS}}}p"
W_TBL = f"{{{DOCX_W_NS}}}tbl"
W_TR = f"{{{DOCX_W_NS}}}tr"
W_TC = f"{{{DOCX_W_NS}}}tc"
W_T = f"{{{DOCX_W_NS}}}t"
W_TC_PR = f"{{{DOCX_W_NS}}}tcPr"
W_GRID_SPAN = f"{{{DOCX_W_NS}}}gridSpan"
W_VMERGE = f"{{{DOCX_W_NS}}}vMerge"
W_VAL = f"{{{DOCX_W_NS}}}val"
W_BR = f"{{{DOCX_W_NS}}}br"
W_TAB = f"{{{DOCX_W_NS}}}tab"


def _docx_text(elem) -> str:
    parts = []
    for node in elem.iter():
        if node.tag == W_T and node.text:
            parts.append(node.text)
        elif node.tag in (W_BR, W_TAB):
            parts.append(" ")
    return _clean_text("".join(parts).replace("\xa0", " "))


def _docx_preview_rows(table_rows: List[Dict], limit: int = 3) -> List[List[str]]:
    preview = []
    for row in (table_rows or [])[:limit]:
        try:
            preview.append([_clean_text(cell.get("text")) for cell in row.get("cells", [])])
        except Exception:
            preview.append([str(row)])
    return preview


def _docx_cell_span(tc) -> int:
    span = 1
    tc_pr = tc.find(W_TC_PR)
    if tc_pr is not None:
        grid_span = tc_pr.find(W_GRID_SPAN)
        if grid_span is not None:
            try:
                span = max(1, int(grid_span.attrib.get(W_VAL, "1") or 1))
            except Exception:
                span = 1
    return span


def _docx_cell_vmerge(tc) -> str:
    tc_pr = tc.find(W_TC_PR)
    if tc_pr is None:
        return ""
    vmerge = tc_pr.find(W_VMERGE)
    if vmerge is None:
        return ""
    return (vmerge.attrib.get(W_VAL) or "continue").strip().lower()


def _docx_expand_rows(table_elem) -> List[Dict]:
    rows = []
    previous_expanded = []
    for tr in table_elem.findall(W_TR):
        row_cells = []
        expanded = []
        logical_col = 0
        for tc in tr.findall(W_TC):
            text = _docx_text(tc)
            span = _docx_cell_span(tc)
            vmerge = _docx_cell_vmerge(tc)
            if vmerge == "continue" and logical_col < len(previous_expanded):
                text = previous_expanded[logical_col]
            row_cells.append(
                {
                    "text": text,
                    "span": span,
                    "start_col": logical_col,
                    "end_col": logical_col + span - 1,
                    "vmerge": vmerge,
                }
            )
            for _ in range(span):
                expanded.append(text)
                logical_col += 1
        if row_cells:
            if previous_expanded:
                target_len = max(len(expanded), len(previous_expanded))
                if len(expanded) < target_len:
                    expanded.extend([""] * (target_len - len(expanded)))
                for idx in range(min(len(expanded), len(previous_expanded))):
                    if not expanded[idx] and previous_expanded[idx]:
                        expanded[idx] = previous_expanded[idx]
            rows.append({"cells": row_cells, "expanded": expanded})
            previous_expanded = expanded
    return rows


_DOCX_DIRECT_HEADER_ALIASES = {
    "day": ("day",),
    "time": ("time", "period", "slot"),
    "branch": ("branch", "dept", "department", "program", "course", "branch name"),
    "section": ("section", "class", "division", "batch", "group"),
    "semester": ("semester", "sem", "term", "year"),
    "subject": ("subject", "subject name", "course", "paper", "topic", "title"),
    "faculty": ("faculty", "teacher", "instructor", "lecturer", "staff", "faculty name"),
    "room": ("room", "classroom", "hall", "venue", "lab room"),
    "lab_theory": ("lab/theory", "lab theory", "type", "category"),
}

_DOCX_IGNORE_ROW_TOKENS = (
    "principal-vice",
    "principal vice",
    "principal",
    "vice",
    "hod",
    "dean",
    "short break",
    "short braek",
    "lunch",
    "lunch break",
    "timetable",
    "siddhartha",
    "institute",
    "technology",
    "sciences",
)


def _docx_row_text(row: Dict) -> str:
    try:
        return " | ".join(_clean_text(cell) for cell in (row.get("expanded") or []) if _clean_text(cell))
    except Exception:
        return ""


def _docx_collapse_merged_text(value: str) -> str:
    text = _clean_text(value)
    if not text:
        return ""
    # Normalize whitespace and separators
    compact = re.sub(r"[\s|]+", " ", text).strip()
    parts = compact.split()
    # If the sequence of tokens repeats twice, collapse to first half
    if len(parts) >= 2 and len(parts) % 2 == 0:
        half = len(parts) // 2
        if parts[:half] == parts[half:]:
            return _clean_text(" ".join(parts[:half]))
    # If entire string is a direct repetition (e.g. "ABCABC"), try collapsing
    m = re.match(r"^(?P<g>.+?)\s*(?P=g)+$", compact)
    if m:
        return _clean_text(m.group("g"))
    return compact


def _normalize_timetable_subject_name(value: str) -> str:
    text = _clean_text(value)
    if not text or _contains_blocked_timetable_text(text):
        return ""
    normalized_key = _normalize_key(text)
    alias_map = {
        "odevc": "Ordinary Differential Equations and Vector Calculus",
        "bee": "Basic Electrical Engineering",
    }
    if normalized_key in alias_map:
        return alias_map[normalized_key]
    text = re.sub(r"[‐‑‒–—―]", "-", text)
    text = re.sub(r"\s+", " ", text).strip()
    return _normalize_display_text(text)


def _normalize_timetable_row(row: Dict[str, str], row_text: str = "", previous: Optional[Dict[str, str]] = None) -> Dict[str, str]:
    normalized = {key: _clean_text(value) for key, value in row.items()}
    normalized["branch"], normalized["section"] = _normalize_timetable_branch_section(normalized.get("branch"), normalized.get("section"), row_text=row_text)
    normalized["semester"] = _safe_int(normalized.get("semester") or _semester_from_text(f"{normalized.get('semester', '')} {row_text}"))
    normalized["day"] = _normalize_display_text(normalized.get("day"))
    start_time, end_time = _split_time_value(normalized.get("start_time"), normalized.get("end_time"))
    if not start_time and not end_time:
        start_time, end_time = _split_time_range(normalized.get("time") or row_text)
    normalized["start_time"] = start_time
    normalized["end_time"] = end_time
    normalized["subject_name"] = _normalize_timetable_subject_name(normalized.get("subject_name") or row.get("subject_name"))
    normalized["faculty_name"] = _normalize_timetable_faculty_name(normalized.get("faculty_name") or normalized.get("faculty") or row.get("faculty_name"))
    normalized["room"] = _normalize_display_text(normalized.get("room"))
    normalized["lab_theory"] = _normalize_timetable_lab_theory(normalized.get("lab_theory") or row.get("lab_theory") or row.get("type") or row.get("category"), row_text=row_text)
    normalized["is_lab"] = int(bool(normalized.get("is_lab")) or _row_has_token(normalized["lab_theory"], "lab", "practical") or _row_has_token(row_text, "lab", "practical") or _row_has_token(normalized["subject_name"], "lab", "practical"))
    if previous:
        normalized = _merge_timetable_row_values(normalized, previous)
    return normalized


def _merge_timetable_row_values(current: Dict[str, str], previous: Optional[Dict[str, str]]) -> Dict[str, str]:
    if not previous:
        return current
    current_day = _clean_text(current.get("day"))
    current_start = _clean_text(current.get("start_time"))
    current_end = _clean_text(current.get("end_time"))
    if not (current_day and current_start and current_end):
        return current
    previous_day = _clean_text(previous.get("day"))
    previous_start = _clean_text(previous.get("start_time"))
    previous_end = _clean_text(previous.get("end_time"))
    if previous_day and current_day != previous_day:
        return current
    if previous_start and current_start != previous_start:
        return current
    if previous_end and current_end != previous_end:
        return current
    previous_branch = _clean_text(previous.get("branch"))
    previous_section = _clean_text(previous.get("section"))
    current_branch = _clean_text(current.get("branch"))
    current_section = _clean_text(current.get("section"))
    if previous_branch and current_branch and current_branch != previous_branch:
        return current
    if previous_section and current_section and current_section != previous_section:
        return current
    if not current_branch and previous_branch:
        current["branch"] = previous_branch
    if not current_section and previous_section:
        current["section"] = previous_section
    if current.get("semester") in (None, "", 0) and previous.get("semester") not in (None, ""):
        current["semester"] = previous.get("semester")
    if not _clean_text(current.get("subject_name")) and _clean_text(previous.get("subject_name")):
        current["subject_name"] = previous.get("subject_name")
    if not _clean_text(current.get("faculty_name")) and _clean_text(previous.get("faculty_name")):
        current["faculty_name"] = previous.get("faculty_name")
    if not _clean_text(current.get("room")) and _clean_text(previous.get("room")):
        current["room"] = previous.get("room")
    return current


def _is_timetable_continuation_row(current: Dict[str, str], previous: Optional[Dict[str, str]]) -> bool:
    if not previous:
        return False
    if not _clean_text(current.get("day")) or not _clean_text(current.get("start_time")) or not _clean_text(current.get("end_time")):
        return False
    if _clean_text(previous.get("day")) and _clean_text(current.get("day")) != _clean_text(previous.get("day")):
        return False
    if _clean_text(previous.get("start_time")) and _clean_text(current.get("start_time")) != _clean_text(previous.get("start_time")):
        return False
    if _clean_text(previous.get("end_time")) and _clean_text(current.get("end_time")) != _clean_text(previous.get("end_time")):
        return False
    if _clean_text(previous.get("section")) and _clean_text(current.get("section")) and _clean_text(current.get("section")) != _clean_text(previous.get("section")):
        return False
    if _clean_text(previous.get("branch")) and _clean_text(current.get("branch")) and _clean_text(current.get("branch")) != _clean_text(previous.get("branch")):
        return False
    return bool(
        (not _clean_text(current.get("subject_name")) and _clean_text(previous.get("subject_name")))
        or (not _clean_text(current.get("faculty_name")) and _clean_text(previous.get("faculty_name")))
        or (not _clean_text(current.get("room")) and _clean_text(previous.get("room")))
    )


def _docx_header_map(table_rows: List[Dict]) -> tuple[int, Dict[str, int]]:
    best_idx = -1
    best_map: Dict[str, int] = {}
    best_score = 0
    for row_index, row in enumerate((table_rows or [])[:5]):
        expanded = row.get("expanded") or []
        if not expanded:
            continue
        header_map: Dict[str, int] = {}
        for cell_index, cell_text in enumerate(expanded):
            key = _normalize_key(cell_text)
            if not key:
                continue
            for field, aliases in _DOCX_DIRECT_HEADER_ALIASES.items():
                if field in header_map:
                    continue
                for alias in aliases:
                    alias_key = _normalize_key(alias)
                    if alias_key and (alias_key == key or alias_key in key or key in alias_key):
                        header_map[field] = cell_index
                        break
        score = sum(1 for field in ("day", "time", "branch", "section", "semester", "subject", "faculty", "room", "lab_theory") if field in header_map)
        if score > best_score:
            best_score = score
            best_idx = row_index
            best_map = header_map
    if best_score >= 5:
        return best_idx, best_map
    return -1, {}


def _docx_table_has_direct_rows(table_rows: List[Dict]) -> bool:
    header_idx, header_map = _docx_header_map(table_rows)
    return header_idx >= 0 and bool(header_map)


def _docx_parse_direct_rows(table_rows: List[Dict], section_state: Dict, debug_jsonl_path: Optional[str] = None) -> Iterator[Dict]:
    header_idx, header_map = _docx_header_map(table_rows)
    if header_idx < 0:
        return

    last_values = {
        "branch": _clean_text(section_state.get("branch")),
        "section": _clean_text(section_state.get("section")),
        "semester": section_state.get("semester"),
        "room": _clean_text(section_state.get("room")),
    }
    previous_slot: Optional[Dict[str, str]] = None
    seen_slots = section_state.get("seen_slots") or set()

    for row_index, row in enumerate(table_rows[header_idx + 1 :], start=header_idx + 1):
        expanded = row.get("expanded") or []
        if not expanded:
            logger.info("skipped_empty_row table_row=%s reason=no_cells", row_index)
            continue

        row_text = _docx_row_text(row)
        if not row_text or not any(_clean_text(cell) for cell in expanded):
            logger.info("skipped_empty_row table_row=%s reason=empty_text", row_index)
            continue
        if _row_has_token(row_text, *_DOCX_IGNORE_ROW_TOKENS):
            logger.info("skipped_empty_row table_row=%s reason=admin_token text=%s", row_index, row_text)
            continue

        def _cell_value(field: str) -> str:
            idx = header_map.get(field)
            value = _docx_collapse_merged_text(expanded[idx]) if idx is not None and idx < len(expanded) else ""
            if value:
                return value
            if field in last_values and _clean_text(last_values.get(field)):
                return _docx_collapse_merged_text(last_values.get(field))
            return ""

        day = _cell_value("day")
        time_text = _cell_value("time")
        branch = _cell_value("branch")
        section = _cell_value("section")
        semester_text = _cell_value("semester")
        subject_name = _cell_value("subject")
        faculty_name = _cell_value("faculty")
        room = _cell_value("room")
        lab_text = _cell_value("lab_theory")

        start_time, end_time = _split_time_value(time_text)
        if not start_time and not end_time:
            start_time, end_time = _split_time_range(time_text)

        if branch:
            last_values["branch"] = branch
        if section:
            last_values["section"] = section
        if semester_text:
            last_values["semester"] = semester_text
        if room:
            last_values["room"] = room

        slot = {
            "branch": branch,
            "section": section,
            "semester": _safe_int(semester_text),
            "day": day,
            "start_time": start_time,
            "end_time": end_time,
            "subject_name": subject_name,
            "faculty_name": faculty_name,
            "is_lab": int(bool(_row_has_token(lab_text, "lab") or _row_has_token(subject_name, "lab", "practical") or _row_has_token(row_text, "lab"))),
            "room": room,
            "lab_theory": lab_text,
        }

        slot = _merge_timetable_row_values(slot, previous_slot)
        slot = _normalize_timetable_row(slot, row_text=row_text, previous=previous_slot)

        # Skip rows where core timetable fields are missing
        core_fields = [slot.get("day"), slot.get("start_time"), slot.get("end_time"), slot.get("section"), slot.get("subject_name")]
        if not any(_clean_text(v) for v in core_fields):
            logger.info("skipped_empty_row table_row=%s reason=missing_core_fields slot=%s", row_index, {k: slot.get(k) for k in ("day", "start_time", "subject_name", "faculty_name")})
            if debug_jsonl_path:
                _append_jsonl(debug_jsonl_path, {"type": "skipped_empty_core", "table_row": row_index, "slot": slot, "timestamp": time.time()})
            continue

        if _row_has_token(row_text, *_DOCX_IGNORE_ROW_TOKENS):
            logger.info("skipped_empty_row table_row=%s reason=admin_token text=%s", row_index, row_text)
            continue

        if not _valid_slot_row(slot):
            logger.info("skipped_empty_row table_row=%s reason=invalid_slot slot=%s text=%s", row_index, slot, row_text)
            continue

        dup = _timetable_semantic_key(slot)
        if dup in seen_slots:
            logger.info("skipped_duplicate_row table_row=%s dup=%s slot=%s", row_index, dup, slot)
            continue
        seen_slots.add(dup)

        if _is_timetable_continuation_row(slot, previous_slot):
            logger.info("merged_continuation_row table_row=%s slot=%s", row_index, slot)

        logger.info(
            "parsed_row table_row=%s day=%s time=%s-%s branch=%s section=%s semester=%s subject=%s faculty=%s room=%s lab=%s",
            row_index,
            slot["day"],
            slot["start_time"],
            slot["end_time"],
            slot["branch"],
            slot["section"],
            slot["semester"],
            slot["subject_name"],
            slot["faculty_name"],
            slot["room"],
            slot["is_lab"],
        )
        if debug_jsonl_path:
            _append_jsonl(
                debug_jsonl_path,
                {
                    "type": "parsed_row",
                    "table_row": row_index,
                    "row_text": row_text,
                    "slot": slot,
                    "timestamp": time.time(),
                },
            )
        previous_slot = slot
        yield slot


def _docx_iter_blocks(path: str):
    with zipfile.ZipFile(path) as archive:
        with archive.open("word/document.xml") as xml_fp:
            context = ET.iterparse(xml_fp, events=("start", "end"))
            stack = []
            for event, elem in context:
                if event == "start":
                    stack.append(elem)
                    continue
                parent = stack[-2] if len(stack) >= 2 else None
                if parent is not None and parent.tag == W_BODY:
                    if elem.tag == W_P:
                        yield ("paragraph", _docx_text(elem))
                        elem.clear()
                    elif elem.tag == W_TBL:
                        yield ("table", _docx_expand_rows(elem))
                        elem.clear()
                if stack:
                    stack.pop()


def scan_docx_structure(path: str, max_tables: Optional[int] = None) -> Dict[str, object]:
    section_names: List[str] = []
    table_count = 0
    faculty_tables = 0
    timetable_tables = 0
    direct_timetable_tables = 0
    for block_type, payload in _docx_iter_blocks(path):
        if block_type == "paragraph":
            text = _clean_text(payload)
            section_name = _section_from_text(text)
            if section_name and section_name not in section_names:
                section_names.append(section_name)
            continue
        if block_type != "table":
            continue
        table_count += 1
        if _is_faculty_table(payload):
            faculty_tables += 1
        if _is_timetable_table(payload):
            timetable_tables += 1
        if _docx_table_has_direct_rows(payload):
            direct_timetable_tables += 1
        if max_tables is not None and table_count >= max_tables:
            break
    return {
        "section_names": section_names,
        "table_count": table_count,
        "faculty_tables": faculty_tables,
        "timetable_tables": timetable_tables,
        "direct_timetable_tables": direct_timetable_tables,
    }


def _section_from_text(text: str) -> str:
    text = _clean_text(text)
    if not text:
        return ""
    if _row_has_token(text, *_DOCX_IGNORE_ROW_TOKENS):
        return ""
    text = re.sub(r"[‐‑‒–—―]", "-", text)
    # Prefer explicit known section patterns like CSE-A, CSE B, CIVIL, MECH
    known = re.search(r"\b(CSE|CSM|ECE|EEE|IT|MECH|CIVIL|AIML|DS)\s*[-_ ]?\s*([A-Z0-9]{1,4})\b", text, flags=re.IGNORECASE)
    if known:
        dept = known.group(1).upper()
        sec = (known.group(2) or "").upper().strip()
        return (dept + ("-" + sec if sec else "")).strip("-")
    generic = re.search(r"\b(?:section|class|division|div|group)\s*[:\-]?\s*([A-Z0-9]{1,6}(?:\s*[-_/]\s*[A-Z0-9]{1,4})?)\b", text, flags=re.IGNORECASE)
    if generic:
        candidate = _normalize_display_text(generic.group(1))
        if candidate:
            return candidate.upper().replace(" ", "-")
    patterns = [
        r"\b([A-Z]{2,}[A-Z0-9]*(?:\s*[-/]\s*[A-Z0-9]{1,4})+)\b",
        r"\b([A-Z]{2,}[A-Z0-9]*\s*[-]\s*[A-Z0-9]{1,4})\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return re.sub(r"\s*[-/]\s*", "-", match.group(1).strip().upper())
    return ""


def scan_pdf_section_name(path: str) -> str:
    if pdfplumber is None:
        return ""
    try:
        with pdfplumber.open(path) as pdf:
            if not pdf.pages:
                return ""
            text = pdf.pages[0].extract_text() or ""
            for line in text.splitlines():
                section_name = _section_from_text(line)
                if section_name:
                    return section_name
    except Exception:
        return ""
    return ""


_ROMAN_VALUES = {"I": 1, "V": 5, "X": 10, "L": 50, "C": 100}


def _roman_to_int(value: str) -> Optional[int]:
    value = _clean_text(value).upper()
    if not value or any(ch not in _ROMAN_VALUES for ch in value):
        return None
    total = 0
    previous = 0
    for char in reversed(value):
        current = _ROMAN_VALUES[char]
        if current < previous:
            total -= current
        else:
            total += current
            previous = current
    return total or None


def _semester_from_text(text: str) -> Optional[int]:
    text = _clean_text(text)
    match = re.search(r"\b([IVX]+|\d+)\s*[-]?\s*Semester\b", text, flags=re.IGNORECASE)
    if not match:
        if re.search(r"\b(B\.?\s*TECH|BTECH|TIMETABLE|SEM(?:ESTER)?)\b", text, flags=re.IGNORECASE):
            shorthand = re.search(r"\b(?:[IVX]+|\d+)\s*[-_/]\s*(\d+)\b", text, flags=re.IGNORECASE)
            if shorthand:
                try:
                    return int(shorthand.group(1))
                except Exception:
                    return None
        return None
    value = match.group(1)
    if value.isdigit():
        try:
            return int(value)
        except Exception:
            return None
    return _roman_to_int(value)


def _subject_acronym(subject_name: str) -> str:
    tokens = re.findall(r"[A-Za-z0-9]+", _clean_text(subject_name).upper())
    stopwords = {"AND", "OF", "THE", "FOR", "IN", "ON", "TO", "WITH", "A", "AN", "LAB", "PRACTICAL"}
    return "".join(token[0] for token in tokens if token and token not in stopwords)


def _faculty_lookup(entries: List[Dict]) -> Dict[str, Dict]:
    lookup: Dict[str, Dict] = {}
    for entry in entries:
        subject_name = _clean_text(entry.get("subject_name"))
        sub_code = _clean_text(entry.get("sub_code"))
        faculty_name = _clean_text(entry.get("faculty_name"))
        aliases = [
            _normalize_key(subject_name),
            _normalize_key(subject_name.replace("lab", "")),
            _normalize_key(sub_code),
            _normalize_key(_subject_acronym(subject_name)),
        ]
        for alias in aliases:
            if alias and alias not in lookup:
                lookup[alias] = {"subject_name": subject_name, "faculty_name": faculty_name, "sub_code": sub_code}
    return lookup


def _merge_faculty_entries(lookup: Dict[str, Dict], entries: List[Dict]) -> None:
    if not entries:
        return
    for entry in entries:
        subject_name = _clean_text(entry.get("subject_name"))
        sub_code = _clean_text(entry.get("sub_code"))
        faculty_name = _clean_text(entry.get("faculty_name"))
        aliases = [
            _normalize_key(subject_name),
            _normalize_key(subject_name.replace("lab", "")),
            _normalize_key(sub_code),
            _normalize_key(_subject_acronym(subject_name)),
        ]
        for alias in aliases:
            if alias and alias not in lookup:
                lookup[alias] = {
                    "subject_name": subject_name,
                    "faculty_name": faculty_name,
                    "sub_code": sub_code,
                }


def _resolve_subject(subject_text: str, lookup: Dict[str, Dict]) -> Dict:
    cleaned = _clean_text(subject_text)
    for candidate in (_normalize_key(cleaned), _normalize_key(cleaned.replace("lab", "")), _normalize_key(_subject_acronym(cleaned))):
        if candidate and candidate in lookup:
            return lookup[candidate]
    return {}


def _append_jsonl(path: str, payload: Dict):
    try:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "a", encoding="utf-8") as handle:
            handle.write(json.dumps(payload, default=str) + "\n")
    except Exception:
        logger.exception("Failed to write jsonl event")


def _section_branch_name(section_name: str, fallback: str = "") -> str:
    section_name = _clean_text(section_name)
    if section_name and "-" in section_name:
        return _clean_text(section_name.split("-", 1)[0])
    return _clean_text(fallback)


def _is_faculty_table(table_rows: List[Dict]) -> bool:
    if not table_rows:
        return False
    header = _normalize_key(" ".join(table_rows[0].get("expanded") or []))
    return "sub code" in header and ("faculty" in header or "subject name" in header)


def _is_timetable_table(table_rows: List[Dict]) -> bool:
    if not table_rows:
        return False
    header = table_rows[0].get("expanded") or []
    header_text = " ".join(header)
    header_key = _normalize_key(header_text)
    if "day" in header_key and re.search(r"\d{1,2}:\d{2}", header_text):
        return True
    if any("break" in _normalize_key(value) for value in header):
        return True
    return False


def _expand_time_bounds(header_cells: List[str], start_col: int, end_col: int) -> tuple[str, str]:
    start_time = ""
    end_time = ""
    for header in header_cells[start_col:end_col + 1]:
        header_start, header_end = _split_time_range(header)
        if header_start and not start_time:
            start_time = header_start
        if header_end:
            end_time = header_end
    # If end_time still empty, try to infer from the next header to the right
    if not end_time:
        # try subsequent headers to find an end_time
        for header in header_cells[end_col + 1:]:
            _, header_end = _split_time_range(header)
            if header_end:
                end_time = header_end
                break
    return start_time, end_time


def _finalize_section_slots(section_state: Dict) -> List[Dict]:
    # Ensure faculty_map includes any buffered faculty_entries
    faculty_entries = section_state.get("faculty_entries") or []
    faculty_map = section_state.get("faculty_map") or {}
    _merge_faculty_entries(faculty_map, faculty_entries)
    timetable_tables = section_state.get("timetable_tables", [])
    resolved_slots: List[Dict] = []
    seen_slots = section_state.get("seen_slots") or set()
    for table_info in timetable_tables:
        table_rows = table_info.get("rows") or []
        for slot in _iter_section_table_slots(table_rows, section_state, faculty_map):
            dup = _dup_key(
                slot.get("branch"),
                slot.get("section"),
                slot.get("semester"),
                slot.get("day"),
                slot.get("start_time"),
                slot.get("end_time"),
                slot.get("subject_name"),
                slot.get("faculty_name"),
                slot.get("room"),
            )
            if dup in seen_slots:
                continue
            seen_slots.add(dup)
            resolved_slots.append(slot)
    section_state["seen_slots"] = seen_slots
    return resolved_slots


def _iter_section_table_slots(table_rows: List[Dict], section_state: Dict, faculty_map: Dict[str, Dict]) -> Iterator[Dict]:
    if not table_rows:
        return
    section_name = _clean_text(section_state.get("section")) or _clean_text(section_state.get("section_hint"))
    branch_name = _clean_text(section_state.get("branch")) or _section_branch_name(section_name, section_state.get("doc_base", ""))
    semester = section_state.get("semester")
    header_cells = table_rows[0].get("expanded") or []
    day_col = 0
    previous_slot: Optional[Dict[str, str]] = None
    seen_slots = section_state.get("seen_slots") or set()
    for idx, header in enumerate(header_cells):
        if "day" in _normalize_key(header):
            day_col = idx
            break
    for row in table_rows[1:]:
        expanded = row.get("expanded") or []
        if not expanded:
            continue
        day = _clean_text(expanded[day_col]) if day_col < len(expanded) else ""
        if not day or _row_has_token(day, "short break", "lunch break", "break"):
            continue
        for cell in row.get("cells", []):
            if cell.get("start_col") == day_col:
                continue
            cell_text = _clean_text(cell.get("text"))
            if not cell_text or _row_has_token(cell_text, "short break", "lunch break", "break"):
                continue
            start_time, end_time = _expand_time_bounds(header_cells, cell.get("start_col", 0), cell.get("end_col", 0))
            if not start_time and not end_time:
                continue
            for subject_piece in _split_subjects(cell_text):
                resolved = _resolve_subject(subject_piece, faculty_map)
                subject_name = _clean_text(resolved.get("subject_name") or subject_piece)
                faculty_name = _clean_text(resolved.get("faculty_name") or "")
                slot = {
                    "branch": branch_name,
                    "section": section_name,
                    "semester": semester,
                    "day": day,
                    "start_time": start_time,
                    "end_time": end_time,
                    "subject_name": subject_name,
                    "faculty_name": faculty_name,
                    "is_lab": int(bool(_row_has_token(subject_piece, "lab", "practical") or _row_has_token(cell_text, "lab", "practical"))),
                    "room": _clean_text(section_state.get("room", "")),
                    "lab_theory": "LAB" if _row_has_token(cell_text, "lab", "practical") else "",
                }
                slot = _merge_timetable_row_values(slot, previous_slot)
                slot = _normalize_timetable_row(slot, row_text=cell_text, previous=previous_slot)
                # Skip if essential semantic fields are missing
                if not (slot.get("day") and slot.get("start_time") and slot.get("end_time") and slot.get("section") and slot.get("subject_name") and slot.get("faculty_name")):
                    continue
                if not _valid_slot_row(slot):
                    logger.info("skipped_invalid_section_slot section=%s day=%s time=%s-%s subject=%s faculty=%s", slot.get("section"), slot.get("day"), slot.get("start_time"), slot.get("end_time"), slot.get("subject_name"), slot.get("faculty_name"))
                    continue
                dup = _timetable_semantic_key(slot)
                if dup in seen_slots:
                    logger.info("skipped_duplicate_section_slot section=%s dup=%s", slot.get("section"), dup)
                    continue
                seen_slots.add(dup)
                previous_slot = slot
                section_state["seen_slots"] = seen_slots
                yield slot


def iter_docx_section_slots(
    path: str,
    debug_jsonl_path: Optional[str] = None,
    single_section_only: bool = False,
    max_tables: Optional[int] = None,
) -> Iterator[Dict]:
    """Stream DOCX timetable rows section-by-section without materializing the full document."""
    if docx is None:
        raise RuntimeError("python-docx is not installed. Install with: pip install python-docx")

    base_name = os.path.splitext(os.path.basename(path))[0]
    document_semester = _semester_from_text(base_name)
    section_state = {
        "section": "",
        "section_hint": "",
        "branch": "",
        "semester": document_semester,
        "doc_base": base_name,
        "room": "",
        "faculty_map": {},
        "table_count": 0,
        "slot_count": 0,
        "seen_slots": set(),
        "timetable_tables": [],
        "faculty_entries": [],
    }
    primary_section = ""
    total_tables = 0
    total_slots = 0
    parse_failures = 0

    def flush_section(reason: str):
        nonlocal section_state, total_slots
        # Finalize and return resolved slots for this section so caller can yield them.
        if section_state.get("table_count", 0) == 0 and not section_state.get("faculty_map"):
            # reset state
            section_state = {
                "section": "",
                "section_hint": "",
                "branch": "",
                "semester": None,
                "doc_base": base_name,
                "room": "",
                "faculty_map": {},
                "table_count": 0,
                "slot_count": 0,
                "seen_slots": set(),
                "timetable_tables": [],
                "faculty_entries": [],
            }
            return []
        section_name = _clean_text(section_state.get("section")) or _clean_text(section_state.get("section_hint")) or f"section_{total_tables}"
        resolved = _finalize_section_slots(section_state)
        section_state["slot_count"] = len(resolved)
        if debug_jsonl_path:
            _append_jsonl(
                debug_jsonl_path,
                {
                    "type": "section_flush",
                    "reason": reason,
                    "section": section_name,
                    "branch": _clean_text(section_state.get("branch")) or _section_branch_name(section_name, section_state.get("doc_base", "")),
                    "semester": section_state.get("semester"),
                    "table_count": section_state.get("table_count", 0),
                    "faculty_entry_count": len(section_state.get("faculty_map") or {}),
                    "slot_count": section_state.get("slot_count", 0),
                    "timestamp": time.time(),
                },
            )
        logger.info(
            "timetable section flush section=%s reason=%s tables=%d slots=%d",
            section_name,
            reason,
            section_state.get("table_count", 0),
            section_state.get("slot_count", 0),
        )
        # Reset
        section_state = {
            "section": "",
            "section_hint": "",
            "branch": "",
            "semester": document_semester,
            "doc_base": base_name,
            "room": "",
            "faculty_map": {},
            "table_count": 0,
            "slot_count": 0,
            "seen_slots": set(),
            "timetable_tables": [],
            "faculty_entries": [],
        }
        gc.collect()
        return resolved

    try:
        for block_type, payload in _docx_iter_blocks(path):
            if block_type == "paragraph":
                text = _clean_text(payload)
                if not text:
                    continue
                section_name = _section_from_text(text)
                if section_name:
                    if single_section_only:
                        if not primary_section:
                            primary_section = section_name
                        elif section_name != primary_section:
                            raise ValueError(
                                f"Multiple sections detected ({primary_section}, {section_name}). Upload one section per DOCX."
                            )
                    if section_state.get("table_count", 0) > 0 or section_state.get("faculty_map"):
                        for slot in flush_section("new_section"):
                            total_slots += 1
                            yield slot
                    section_state["section"] = section_name
                    section_state["section_hint"] = section_name
                    section_state["branch"] = _section_branch_name(section_name, section_state.get("doc_base", ""))
                    if section_state["semester"] is None:
                        section_state["semester"] = _semester_from_text(text)
                    logger.info("timetable section start section=%s branch=%s semester=%s", section_name, section_state["branch"], section_state["semester"])
                    if debug_jsonl_path:
                        _append_jsonl(
                            debug_jsonl_path,
                            {
                                "type": "section_start",
                                "section": section_name,
                                "branch": section_state["branch"],
                                "semester": section_state["semester"],
                                "timestamp": time.time(),
                            },
                        )
                    continue
                semester = _semester_from_text(text)
                if semester is not None and section_state["semester"] is None:
                    section_state["semester"] = semester
                if not section_state["branch"]:
                    section_state["branch"] = _section_branch_name(text, section_state.get("doc_base", ""))
                continue

            if block_type != "table":
                continue

            table_rows = payload or []
            total_tables += 1
            if max_tables is not None and total_tables > max_tables:
                raise ValueError(
                    f"Too many tables detected ({total_tables}). Split the timetable into section-wise DOCX files."
                )
            try:
                if debug_jsonl_path:
                    _append_jsonl(
                        debug_jsonl_path,
                        {
                            "type": "raw_table_preview",
                            "section": section_state.get("section") or section_state.get("section_hint") or f"section_{total_tables}",
                            "rows": _docx_preview_rows(table_rows),
                            "timestamp": time.time(),
                        },
                    )
                if total_tables == 1:
                    logger.info("DOCX raw table preview section=%s rows=%s", section_state.get("section") or section_state.get("section_hint") or f"section_{total_tables}", _docx_preview_rows(table_rows))
                direct_rows = list(_docx_parse_direct_rows(table_rows, section_state, debug_jsonl_path=debug_jsonl_path))
                if direct_rows:
                    section_state["table_count"] += 1
                    # Resolve now against current faculty_map
                    for slot in direct_rows:
                        dup = _dup_key(slot.get("branch"), slot.get("section"), slot.get("semester"), slot.get("day"), slot.get("start_time"), slot.get("end_time"), slot.get("subject_name"), slot.get("faculty_name"), slot.get("room"))
                        if dup in section_state.get("seen_slots", set()):
                            if debug_jsonl_path:
                                _append_jsonl(debug_jsonl_path, {"type": "skipped_duplicate", "dup": dup, "slot": slot, "timestamp": time.time()})
                            continue
                        section_state.get("seen_slots").add(dup)
                        section_state["slot_count"] += 1
                        total_slots += 1
                        yield slot
                    if debug_jsonl_path:
                        _append_jsonl(debug_jsonl_path, {"type": "direct_timetable_table", "section": section_state.get("section") or section_state.get("section_hint") or f"section_{total_tables}", "rows": len(table_rows), "parsed_rows": len(direct_rows), "timestamp": time.time()})
                    table_rows = None
                    gc.collect()
                    continue
                if _is_faculty_table(table_rows):
                    entries = []
                    # Determine column indices from header if possible
                    header_expanded = table_rows[0].get("expanded") or []
                    hkeys = [_normalize_key(h) for h in header_expanded]
                    code_idx = None
                    name_idx = None
                    faculty_idx = None
                    for i, k in enumerate(hkeys):
                        if code_idx is None and ("sub code" in k or "subject code" in k):
                            code_idx = i
                        elif name_idx is None and "subject name" in k:
                            name_idx = i
                        elif faculty_idx is None and "faculty" in k:
                            faculty_idx = i

                    # Fallback: assume groups of three if header not explicit
                    last_code = last_name = last_faculty = ""
                    for row in table_rows[1:]:
                        expanded = row.get("expanded") or []
                        if code_idx is not None or name_idx is not None or faculty_idx is not None:
                            sub_code = _clean_text(expanded[code_idx]) if code_idx is not None and code_idx < len(expanded) else ""
                            subject_name = _clean_text(expanded[name_idx]) if name_idx is not None and name_idx < len(expanded) else ""
                            faculty_name = _clean_text(expanded[faculty_idx]) if faculty_idx is not None and faculty_idx < len(expanded) else ""
                        else:
                            # group into triplets
                            sub_code = subject_name = faculty_name = ""
                            for offset in range(0, len(expanded), 3):
                                group = expanded[offset:offset + 3]
                                if len(group) < 2:
                                    continue
                                # if 3 elements, treat as code, name, faculty
                                if len(group) >= 3:
                                    sub_code = _clean_text(group[0])
                                    subject_name = _clean_text(group[1])
                                    faculty_name = _clean_text(group[2])
                                else:
                                    # two columns: subject, faculty
                                    subject_name = _clean_text(group[0])
                                    faculty_name = _clean_text(group[1])
                                # prefer non-empty groups
                                if sub_code or subject_name or faculty_name:
                                    break

                        # carry-forward merged rows
                        if not sub_code:
                            sub_code = last_code
                        else:
                            last_code = sub_code
                        if not subject_name:
                            subject_name = last_name
                        else:
                            last_name = subject_name
                        if not faculty_name:
                            faculty_name = last_faculty
                        else:
                            last_faculty = faculty_name

                        if not (sub_code or subject_name or faculty_name):
                            continue
                        if _normalize_key(sub_code) in {"sub code", "subject code"} or _normalize_key(subject_name) == "subject name":
                            continue
                        entries.append({"sub_code": sub_code, "subject_name": _normalize_timetable_subject_name(subject_name), "faculty_name": _normalize_timetable_faculty_name(faculty_name)})
                    _merge_faculty_entries(section_state["faculty_map"], entries)
                    section_state.setdefault("faculty_entries", []).extend(entries)
                    section_state["table_count"] += 1
                    if debug_jsonl_path:
                        _append_jsonl(debug_jsonl_path, {"type": "faculty_table", "section": section_state.get("section") or section_state.get("section_hint") or f"section_{total_tables}", "entries": len(entries), "timestamp": time.time()})
                    entries = None
                    table_rows = None
                    gc.collect()
                    continue

                if _is_timetable_table(table_rows) or section_state.get("table_count", 0) == 0:
                    section_state["table_count"] += 1
                    section_state.setdefault("timetable_tables", []).append({"rows": table_rows})
                    mem_estimate_kb = (len(table_rows) * 180 + len(section_state.get("faculty_map") or {}) * 80) / 1024.0
                    logger.info(
                        "timetable table buffered section=%s rows=%d tables=%d mem_estimate_kb=%.1f",
                        section_state.get("section") or section_state.get("section_hint") or f"section_{total_tables}",
                        len(table_rows),
                        section_state.get("table_count", 0),
                        mem_estimate_kb,
                    )
                    if debug_jsonl_path:
                        _append_jsonl(debug_jsonl_path, {"type": "timetable_table", "section": section_state.get("section") or section_state.get("section_hint") or f"section_{total_tables}", "rows": len(table_rows), "timestamp": time.time()})
                table_rows = None
                gc.collect()
            except Exception:
                parse_failures += 1
                logger.exception("Failed to classify DOCX table index=%d", total_tables)
                continue

    except Exception:
        logger.exception("DOCX section iterator failed for %s", path)
        raise

    for slot in flush_section("eof"):
        total_slots += 1
        yield slot
    logger.info("iter_docx_section_slots summary: tables=%d slots=%d failures=%d file=%s", total_tables, total_slots, parse_failures, os.path.basename(path))


_PDF_DAY_ALIASES = {
    "mon": "Monday",
    "tue": "Tuesday",
    "wed": "Wednesday",
    "thu": "Thursday",
    "fri": "Friday",
    "sat": "Saturday",
    "sun": "Sunday",
}

_PDF_DAY_RE = re.compile(r"\b(mon(?:day)?|tue(?:sday)?|wed(?:nesday)?|thu(?:rsday)?|fri(?:day)?|sat(?:urday)?|sun(?:day)?)\b", re.IGNORECASE)
_PDF_TIME_RANGE_RE = re.compile(r"(\d{1,2}[:\.]\d{2}\s*(?:AM|PM)?)\s*(?:-|to)\s*(\d{1,2}[:\.]\d{2}\s*(?:AM|PM)?)", re.IGNORECASE)
_PDF_TIME_RANGE_HOUR_RE = re.compile(r"(\d{1,2}\s*(?:AM|PM))\s*(?:-|to)\s*(\d{1,2}\s*(?:AM|PM))", re.IGNORECASE)
_PDF_DECORATIVE_TOKENS = ("principal", "hod", "head of department", "department", "dean", "siddhartha", "institute", "technology", "sciences", "timetable")
_PDF_BREAK_TOKENS = ("short break", "short braek", "lunch break", "break", "lunch")
_PDF_SECTION_STRICT_RE = re.compile(r"\b((?:CSM|CSE|ECE|EEE|IT|MECH|CIVIL|AIML|DS))(?:\s*[-_/ ]\s*([A-Z0-9]{1,4}))?\b", re.IGNORECASE)
_PDF_SECTION_CANDIDATE_RE = re.compile(r"\b((?:CSM|CSE|ECE|EEE|IT|MECH|CIVIL|AIML|DS))(?:\s*[-_/ ]\s*([A-Z0-9]{1,4}))?\b", re.IGNORECASE)
_PDF_SECTION_CONTEXT_TOKENS = ("class", "section", "branch", "dept", "department", "program", "programme", "course")


def _normalize_day_name(token: str) -> str:
    key = _clean_text(token).lower()[:3]
    return _PDF_DAY_ALIASES.get(key, token.title())


def _extract_pdf_day(text: str) -> tuple[str, str]:
    match = _PDF_DAY_RE.search(text or "")
    if not match:
        return "", ""
    return _normalize_day_name(match.group(1)), match.group(0)


def _extract_pdf_time_range(text: str) -> tuple[str, str, str]:
    match = _PDF_TIME_RANGE_RE.search(text or "") or _PDF_TIME_RANGE_HOUR_RE.search(text or "")
    if not match:
        return "", "", ""
    start = _format_time_str(match.group(1)) or ""
    end = _format_time_str(match.group(2)) or ""
    return start, end, match.group(0)


def _pdf_text_has_time(text: str) -> bool:
    return bool(_PDF_TIME_RANGE_RE.search(text or "") or _PDF_TIME_RANGE_HOUR_RE.search(text or ""))


def _pdf_header_has_time(text: str) -> bool:
    return bool(re.search(r"\d{1,2}[:\.]\d{2}", text or "") or re.search(r"\b\d{1,2}\s*(AM|PM)\b", text or "", flags=re.IGNORECASE))


def _pdf_is_break(text: str) -> bool:
    return _row_has_token(text, *_PDF_BREAK_TOKENS)


def _pdf_is_decorative_line(text: str) -> bool:
    return _row_has_token(text, *_PDF_DECORATIVE_TOKENS)


def _pdf_section_from_line(text: str) -> str:
    raw, normalized = _pdf_extract_section_candidate(text)
    return normalized


def _pdf_normalize_section_value(dept: str, section: str) -> str:
    dept = re.sub(r"[\s_]+", "-", _clean_text(dept).upper())
    section = re.sub(r"[\s_]+", "-", _clean_text(section).upper())
    dept = re.sub(r"-+", "-", dept).strip("-")
    section = re.sub(r"-+", "-", section).strip("-")
    if not dept or not section:
        return ""
    return f"{dept}-{section}"


def _pdf_record_section_attempt(report: Dict[str, object], source: str, raw_text: str, normalized: str, accepted: bool, reason: str = "") -> None:
    attempts = report.get("section_detection_attempts") or []
    payload = {
        "source": source,
        "raw": _clean_text(raw_text),
        "normalized": _clean_text(normalized),
        "accepted": bool(accepted),
    }
    if reason:
        payload["reason"] = reason
    if len(attempts) < PDF_DIAG_SAMPLE_CAP:
        attempts.append(payload)
    report["section_detection_attempts"] = attempts[:PDF_DIAG_SAMPLE_CAP]
    logger.debug(
        "PDF section detection attempt source=%s raw=%r normalized=%s accepted=%s reason=%s",
        source,
        payload["raw"],
        payload["normalized"],
        accepted,
        reason,
    )


def _pdf_excerpt(text: str, limit: int = 240) -> str:
    cleaned = re.sub(r"\s+", " ", _clean_text(text)).strip()
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[: max(0, limit - 3)].rstrip() + "..."


def _pdf_extract_section_candidate(text: str, require_context: bool = False) -> tuple[str, str]:
    cleaned = _clean_text(text)
    if not cleaned:
        return "", ""
    cleaned = re.sub(r"[‐‑‒–—―]", "-", cleaned)
    if _pdf_is_decorative_line(cleaned) or _pdf_is_break(cleaned):
        return "", ""
    if _pdf_text_has_time(cleaned) or _PDF_DAY_RE.search(cleaned):
        return "", ""
    if require_context and not _pdf_line_has_section_context(cleaned):
        return "", ""

    match = _PDF_SECTION_STRICT_RE.search(cleaned)
    if not match:
        return "", ""
    raw = _clean_text(match.group(0))
    normalized = _pdf_normalize_section_value(match.group(1), match.group(2))
    if not normalized:
        return raw, ""
    return raw, normalized


def _pdf_should_consider_section_line(text: str) -> bool:
    if _pdf_is_decorative_line(text) or _pdf_is_break(text):
        return False
    if _pdf_text_has_time(text):
        return False
    if _PDF_DAY_RE.search(text or ""):
        return False
    norm = _normalize_key(text)
    if any(token in norm for token in ("day", "time", "break", "lunch")):
        return False
    return True


def _pdf_line_has_section_context(text: str) -> bool:
    norm = _normalize_key(text)
    return any(token in norm for token in _PDF_SECTION_CONTEXT_TOKENS)


def _pdf_add_rejected_candidate(candidate: str, report: Dict[str, object]) -> None:
    rejected = report.get("rejected_section_candidates") or []
    if candidate and candidate not in rejected:
        rejected.append(candidate)
    report["rejected_section_candidates"] = rejected[:PDF_DIAG_SAMPLE_CAP]


def _pdf_collect_rejected_candidates(text: str, report: Dict[str, object]) -> None:
    for match in _PDF_SECTION_CANDIDATE_RE.finditer(text or ""):
        candidate = _clean_text(match.group(0))
        if not candidate:
            continue
        if _PDF_SECTION_STRICT_RE.search(candidate):
            continue
        _pdf_add_rejected_candidate(candidate, report)


def _pdf_collect_section_candidates(lines: Iterable[str], report: Dict[str, object], source: str, require_context: bool = False) -> List[str]:
    sections: List[str] = []
    source_label = source
    if source.startswith("title_"):
        source_label = "timetable_title"
    elif source.startswith("nearby_label_"):
        source_label = "nearby_label"
    elif source.startswith("faculty_table_"):
        source_label = "faculty_table"
    elif source.startswith("table_rows_"):
        source_label = "table_rows"
    for line in lines:
        cleaned = _clean_text(line)
        if not cleaned or not _pdf_should_consider_section_line(cleaned):
            continue
        if require_context and not _pdf_line_has_section_context(cleaned):
            continue
        raw, section = _pdf_extract_section_candidate(cleaned, require_context=require_context)
        if section:
            if not report.get("detected_section_raw"):
                report["detected_section_raw"] = raw
                report["detected_section_normalized"] = section
            _pdf_record_section_attempt(report, source, raw, section, True)
            if section not in sections:
                sections.append(section)
            if not report.get("detected_section_source"):
                report["detected_section_source"] = source_label
                report["detected_section_source_line"] = cleaned
        else:
            candidate_match = _PDF_SECTION_CANDIDATE_RE.search(cleaned)
            if candidate_match:
                _pdf_add_rejected_candidate(candidate_match.group(0), report)
                _pdf_record_section_attempt(report, source, candidate_match.group(0), "", False, "regex_mismatch")
            elif require_context and _pdf_line_has_section_context(cleaned):
                _pdf_record_section_attempt(report, source, cleaned, "", False, "context_without_section")
    return sections


def _pdf_bbox_key(bbox) -> tuple:
    if not bbox or len(bbox) < 4:
        return ()
    return (int(round(bbox[0])), int(round(bbox[1])), int(round(bbox[2])), int(round(bbox[3])))


def _pdf_find_tables(page) -> List[Dict]:
    tables: List[Dict] = []
    seen = set()
    for settings in PDF_TABLE_SETTINGS:
        try:
            found = page.find_tables(table_settings=settings) or []
        except Exception:
            found = []
        for table in found:
            bbox = getattr(table, "bbox", None)
            key = _pdf_bbox_key(bbox)
            if not key or key in seen:
                continue
            seen.add(key)
            tables.append({"table": table, "bbox": bbox, "key": key, "settings": settings})
    tables.sort(key=lambda t: (t["bbox"][1], t["bbox"][0]) if t.get("bbox") else (0, 0))
    return tables


def _pdf_table_extract_matrix(table) -> List[List[str]]:
    try:
        raw = table.extract() or []
    except Exception:
        raw = []
    if not raw:
        return []
    max_cols = max((len(r) for r in raw if r), default=0)
    rows = []
    for row in raw:
        row = row or []
        cleaned = [_clean_text(c) for c in row]
        if max_cols and len(cleaned) < max_cols:
            cleaned.extend([""] * (max_cols - len(cleaned)))
        rows.append(cleaned)
    return rows


def _pdf_locate_timetable_header(rows: List[List[str]]) -> tuple[int, List[str]]:
    for idx in range(min(3, len(rows))):
        row = rows[idx] or []
        if not row:
            continue
        has_day_header = any("day" in _normalize_key(c) for c in row if c)
        time_hits = sum(1 for c in row if _pdf_header_has_time(c))
        if time_hits >= 2 and (has_day_header or any(_PDF_DAY_RE.search(c or "") for c in row)):
            return idx, row
    return -1, []


def _pdf_locate_faculty_header(rows: List[List[str]]) -> tuple[int, List[str]]:
    for idx in range(min(3, len(rows))):
        row = rows[idx] or []
        header_key = _normalize_key(" ".join(row))
        if "faculty" in header_key and ("sub code" in header_key or "subject code" in header_key or "subject name" in header_key):
            return idx, row
    return -1, []


def _pdf_is_faculty_table_rows(rows: List[List[str]]) -> bool:
    header_idx, header_cells = _pdf_locate_faculty_header(rows)
    if header_idx < 0:
        return False
    header_key = _normalize_key(" ".join(header_cells))
    return ("sub code" in header_key or "subject code" in header_key) and "faculty" in header_key


def _pdf_parse_faculty_rows(rows: List[List[str]]) -> List[Dict]:
    header_idx, header_cells = _pdf_locate_faculty_header(rows)
    if header_idx < 0:
        return []
    header_key = [_normalize_key(c) for c in header_cells]
    code_idx = None
    name_idx = None
    faculty_idx = None
    for idx, key in enumerate(header_key):
        if code_idx is None and ("sub code" in key or "subject code" in key):
            code_idx = idx
        elif name_idx is None and "subject name" in key:
            name_idx = idx
        elif faculty_idx is None and "faculty" in key:
            faculty_idx = idx
    entries: List[Dict] = []
    # Fallback heuristics: if columns couldn't be located, pick sensible defaults
    col_count = len(header_cells)
    if faculty_idx is None:
        faculty_idx = col_count - 1
    if name_idx is None:
        name_idx = max(0, col_count - 2)
    # carry-forward last seen values to handle merged rows
    last_code = last_name = last_faculty = ""
    for row in rows[header_idx + 1:]:
        if not row or not any(_clean_text(c) for c in row):
            continue
        sub_code = _clean_text(row[code_idx]) if code_idx is not None and code_idx < len(row) else ""
        subject_name = _clean_text(row[name_idx]) if name_idx is not None and name_idx < len(row) else ""
        faculty_name = _clean_text(row[faculty_idx]) if faculty_idx is not None and faculty_idx < len(row) else ""
        # If any column is empty, try to inherit from previous non-empty value (merged rows)
        if not sub_code:
            sub_code = last_code
        else:
            last_code = sub_code
        if not subject_name:
            subject_name = last_name
        else:
            last_name = subject_name
        if not faculty_name:
            faculty_name = last_faculty
        else:
            last_faculty = faculty_name
        if not (sub_code or subject_name or faculty_name):
            continue
        if _normalize_key(sub_code) in {"sub code", "subject code"}:
            continue
        entries.append({"sub_code": sub_code, "subject_name": _normalize_timetable_subject_name(subject_name), "faculty_name": _normalize_timetable_faculty_name(faculty_name)})
    return entries


def _pdf_cell_bbox(cell) -> Optional[tuple]:
    if cell is None:
        return None
    if isinstance(cell, dict):
        if all(k in cell for k in ("x0", "top", "x1", "bottom")):
            return (cell["x0"], cell["top"], cell["x1"], cell["bottom"])
        if "bbox" in cell:
            return cell.get("bbox")
    if hasattr(cell, "bbox"):
        return getattr(cell, "bbox")
    return None


def _pdf_cell_text(cell) -> str:
    if cell is None:
        return ""
    if isinstance(cell, dict):
        return _clean_text(cell.get("text", ""))
    if hasattr(cell, "text"):
        return _clean_text(getattr(cell, "text") or "")
    return ""


def _pdf_table_column_bounds(table, col_count: int) -> List[tuple]:
    bounds = []
    try:
        cols = getattr(table, "columns", None)
        if cols:
            for col in cols:
                bbox = _pdf_cell_bbox(col)
                if bbox and len(bbox) >= 3:
                    bounds.append((bbox[0], bbox[2]))
    except Exception:
        bounds = []
    if bounds and len(bounds) >= col_count:
        return bounds
    return bounds


def _pdf_col_span_for_bbox(bbox: tuple, col_bounds: List[tuple]) -> tuple[Optional[int], Optional[int]]:
    x0 = bbox[0]
    x1 = bbox[2]
    indices = [i for i, (c0, c1) in enumerate(col_bounds) if c1 > x0 and c0 < x1]
    if not indices:
        return None, None
    return min(indices), max(indices)


def _pdf_table_row_spans(table, row_index: int, col_bounds: List[tuple]) -> Optional[List[Dict]]:
    if not table or not col_bounds:
        return None
    try:
        rows = getattr(table, "rows", None)
        if not rows or row_index >= len(rows):
            return None
        row = rows[row_index]
        cells = getattr(row, "cells", None)
        if not cells:
            return None
        spans = []
        for cell in cells:
            text = _pdf_cell_text(cell)
            if not text:
                continue
            bbox = _pdf_cell_bbox(cell)
            if not bbox:
                continue
            start_col, end_col = _pdf_col_span_for_bbox(bbox, col_bounds)
            if start_col is None or end_col is None:
                continue
            spans.append({"start_col": start_col, "end_col": end_col, "text": text})
        if spans:
            spans.sort(key=lambda s: (s["start_col"], s["end_col"]))
            return spans
    except Exception:
        return None
    return None


def _pdf_row_spans_from_values(row_values: List[str], header_slots: List[Dict], day_col: int) -> List[Dict]:
    spans: List[Dict] = []
    current_text = ""
    current_start = None
    for idx, cell in enumerate(row_values):
        if idx == day_col:
            continue
        if idx < len(header_slots) and header_slots[idx].get("is_break"):
            if current_start is not None and current_text:
                spans.append({"start_col": current_start, "end_col": idx - 1, "text": current_text})
            current_text = ""
            current_start = None
            continue
        text = _clean_text(cell)
        if text:
            if current_start is None:
                current_text = text
                current_start = idx
            else:
                if text == current_text and _row_has_token(text, "lab", "practical"):
                    continue
                spans.append({"start_col": current_start, "end_col": idx - 1, "text": current_text})
                current_text = text
                current_start = idx
        else:
            if current_start is not None and current_text and _row_has_token(current_text, "lab", "practical"):
                continue
            if current_start is not None and current_text:
                spans.append({"start_col": current_start, "end_col": idx - 1, "text": current_text})
            current_text = ""
            current_start = None
    if current_start is not None and current_text:
        spans.append({"start_col": current_start, "end_col": len(row_values) - 1, "text": current_text})
    return spans


def _pdf_split_span_on_breaks(span: Dict, header_slots: List[Dict], day_col: int) -> List[tuple]:
    segments = []
    current = None
    for col in range(span["start_col"], span["end_col"] + 1):
        if col == day_col:
            continue
        if col < len(header_slots) and header_slots[col].get("is_break"):
            if current:
                segments.append((current[0], current[1]))
            current = None
            continue
        if current is None:
            current = [col, col]
        else:
            current[1] = col
    if current:
        segments.append((current[0], current[1]))
    return segments


def _pdf_find_day_col(header_cells: List[str], rows: List[List[str]], header_idx: int) -> Optional[int]:
    for idx, cell in enumerate(header_cells):
        if "day" in _normalize_key(cell):
            return idx
    day_counts = {}
    for row in rows[header_idx + 1:]:
        for idx, cell in enumerate(row):
            day, _ = _extract_pdf_day(cell)
            if day:
                day_counts[idx] = day_counts.get(idx, 0) + 1
    if not day_counts:
        return None
    return max(day_counts.items(), key=lambda kv: kv[1])[0]


def _pdf_build_header_slots(header_cells: List[str]) -> List[Dict]:
    slots = []
    for idx, text in enumerate(header_cells):
        start, end, _ = _extract_pdf_time_range(text)
        slots.append({
            "index": idx,
            "label": _clean_text(text),
            "start_time": start,
            "end_time": end,
            "is_break": _pdf_is_break(text),
        })
    return slots


def _pdf_debug_table_sample(rows: List[List[str]], limit: int = 3) -> List[List[str]]:
    sample = []
    for row in (rows or [])[:limit]:
        sample.append([_clean_text(cell) for cell in (row or [])])
    return sample


def _pdf_parse_timetable_table(
    table_info: Dict,
    section_name: str,
    branch_name: str,
    semester: Optional[int],
    faculty_map: Dict[str, Dict],
    report: Dict[str, object],
) -> Iterator[Dict]:
    rows = table_info.get("rows") or []
    header_idx = table_info.get("header_idx", -1)
    header_cells = table_info.get("header_cells") or []
    if header_idx < 0 or not header_cells:
        report["validation_errors"].append("Timetable header row not detected (Day/time headers missing).")
        raise TimetablePDFValidationError("Timetable header row not detected. Ensure the grid has a Day column and time slot headers.")

    raw_samples = report.setdefault("raw_table_samples", [])
    if len(raw_samples) < PDF_DIAG_SAMPLE_CAP:
        raw_samples.append(_pdf_debug_table_sample(rows))

    day_col = _pdf_find_day_col(header_cells, rows, header_idx)
    if day_col is None:
        report["validation_errors"].append("Day column not detected in the timetable grid.")
        raise TimetablePDFValidationError("Day column not detected in the timetable grid. Ensure the first column is labeled Day and lists MON/TUE/WED...")

    header_slots = _pdf_build_header_slots(header_cells)
    time_slots = [slot for slot in header_slots if slot["index"] != day_col and not slot.get("is_break") and (slot.get("start_time") or slot.get("end_time") or slot.get("label"))]
    if not any(slot.get("start_time") or slot.get("end_time") for slot in time_slots):
        report["validation_errors"].append("Time slot headers not detected in the timetable grid.")
        raise TimetablePDFValidationError("Time slot headers not detected. Ensure the header row contains time ranges like 09:00-10:00.")

    detected_days = set()
    detected_subjects = report.get("extracted_subjects_sample") or []
    detected_time_slots = report.get("detected_time_slots") or []
    report.setdefault("valid_rows", 0)
    report.setdefault("skipped_rows", 0)
    if not detected_time_slots:
        for slot in time_slots:
            start = slot.get("start_time") or ""
            end = slot.get("end_time") or ""
            label = f"{start}-{end}" if start and end else (slot.get("label") or "")
            if label and label not in detected_time_slots:
                detected_time_slots.append(label)
    report["detected_time_slots"] = detected_time_slots[:PDF_DIAG_SAMPLE_CAP]

    # track duplicates to avoid emitting the same logical slot multiple times
    seen_slots_local = set()
    previous_slot: Optional[Dict[str, str]] = None

    col_bounds = _pdf_table_column_bounds(table_info.get("table"), len(header_cells))
    for row_index, row in enumerate(rows[header_idx + 1:], start=header_idx + 1):
        if not row:
            continue
        if not any(_clean_text(cell) for cell in row):
            logger.info("SKIPPED_EMPTY_ROW reason=blank_pdf_row row_index=%s", row_index)
            continue
        if len(row) < len(header_cells):
            row = row + [""] * (len(header_cells) - len(row))
        day_raw = _clean_text(row[day_col]) if day_col < len(row) else ""
        day, _ = _extract_pdf_day(day_raw)
        if not day:
            logger.info("SKIPPED_EMPTY_ROW reason=missing_day row_index=%s text=%s", row_index, " | ".join(row))
            continue
        detected_days.add(day)

        spans = _pdf_table_row_spans(table_info.get("table"), row_index, col_bounds)
        if spans is None:
            spans = _pdf_row_spans_from_values(row, header_slots, day_col)

        for span in spans:
            text = _clean_text(span.get("text"))
            if not text or _pdf_is_break(text):
                continue
            for seg_start, seg_end in _pdf_split_span_on_breaks(span, header_slots, day_col):
                start_time, end_time = _expand_time_bounds(header_cells, seg_start, seg_end)
                if not start_time:
                    start_time = header_slots[seg_start].get("start_time") if seg_start < len(header_slots) else ""
                if not end_time:
                    end_time = header_slots[seg_end].get("end_time") if seg_end < len(header_slots) else ""
                if not start_time and not end_time:
                    continue
                for subject_piece in _split_subjects(text):
                    resolved = _resolve_subject(subject_piece, faculty_map)
                    subject_name = _clean_text(resolved.get("subject_name") or subject_piece)
                    faculty_name = _clean_text(resolved.get("faculty_name") or "")
                    semester_value = "" if semester is None else semester
                    # Ignore administrative rows that are not real timetable entries
                    admin_tokens = ("PRINCIPAL", "PRINCIPAL-VICE", "VICE", "HOD", "DEAN", "SIDDHARTHA", "INSTITUTE", "TECHNOLOGY", "SCIENCES")
                    if any(tok.lower() in subject_piece.lower() or tok.lower() in text.lower() for tok in admin_tokens):
                        if report is not None and len(report.get("skipped_admin_rows", [])) < PDF_DIAG_SAMPLE_CAP:
                            report.setdefault("skipped_admin_rows", []).append({"index": row_index, "text": text})
                        report["skipped_rows"] = int(report.get("skipped_rows", 0) or 0) + 1
                        logger.info("SKIPPED_EMPTY_ROW reason=admin_tokens row_index=%s text=%s", row_index, text)
                        continue
                    slot = {
                        "branch": branch_name,
                        "section": section_name,
                        "semester": semester_value,
                        "day": day,
                        "start_time": start_time,
                        "end_time": end_time,
                        "subject_name": subject_name,
                        "faculty_name": faculty_name,
                        "is_lab": int(bool(_row_has_token(subject_piece, "lab", "practical") or _row_has_token(text, "lab", "practical"))),
                        "room": "",
                        "lab_theory": "LAB" if _row_has_token(subject_piece, "lab", "practical") or _row_has_token(text, "lab", "practical") else "",
                    }
                    slot = _merge_timetable_row_values(slot, previous_slot)
                    slot = _normalize_timetable_row(slot, row_text=text, previous=previous_slot)
                    if _row_has_token(text, "short braek", "short break", "lunch break", "lunch", "break"):
                        logger.info("SKIPPED_EMPTY_ROW reason=break_row row_index=%s text=%s", row_index, text)
                        report["skipped_rows"] = int(report.get("skipped_rows", 0) or 0) + 1
                        continue
                    if not is_valid_timetable_row(slot):
                        if report is not None and len(report.get("skipped_empty_rows", [])) < PDF_DIAG_SAMPLE_CAP:
                            report.setdefault("skipped_empty_rows", []).append({"index": row_index, "text": text})
                        report["skipped_rows"] = int(report.get("skipped_rows", 0) or 0) + 1
                        logger.info("SKIPPED_EMPTY_ROW reason=missing_semantic_fields row_index=%s text=%s slot=%s", row_index, text, slot)
                        continue

                    # If the core timetable fields are still missing, skip the row.
                    if not day or not start_time or not (branch_name or section_name) or (not subject_name and not faculty_name):
                        if report is not None and len(report.get("skipped_empty_rows", [])) < PDF_DIAG_SAMPLE_CAP:
                            report.setdefault("skipped_empty_rows", []).append({"index": row_index, "text": text})
                        report["skipped_rows"] = int(report.get("skipped_rows", 0) or 0) + 1
                        logger.info("SKIPPED_EMPTY_ROW reason=required_field_missing row_index=%s text=%s slot=%s", row_index, text, slot)
                        continue

                    if subject_name and subject_name not in detected_subjects:
                        detected_subjects.append(subject_name)
                    # Attempt fuzzy faculty name detection from teachers table if missing
                    if not faculty_name:
                        try:
                            db = get_db()
                            teacher_name_index = _build_teacher_name_index(db)
                            import difflib
                            key = _normalize_teacher_name(text)
                            candidates = difflib.get_close_matches(key, list(teacher_name_index.keys()), n=1, cutoff=0.65)
                            if candidates:
                                faculty_name = teacher_name_index.get(candidates[0]) or faculty_name
                        except Exception:
                            pass

                    # Log parsed pieces for diagnostics
                    if report is not None and len(report.get("parsed_rows", [])) < PDF_DIAG_SAMPLE_CAP:
                        report.setdefault("parsed_rows", []).append({
                            "index": row_index,
                            "parsed_subject": subject_name,
                            "parsed_faculty": faculty_name,
                            "parsed_section": slot.get("section"),
                            "day": slot.get("day"),
                            "start_time": slot.get("start_time"),
                            "end_time": slot.get("end_time"),
                            "text": text,
                        })
                    dup = _timetable_semantic_key(slot)
                    if dup in seen_slots_local:
                        if report is not None and len(report.get("skipped_duplicates", [])) < PDF_DIAG_SAMPLE_CAP:
                            report.setdefault("skipped_duplicates", []).append(dup)
                        report["skipped_rows"] = int(report.get("skipped_rows", 0) or 0) + 1
                        continue
                    seen_slots_local.add(dup)
                    previous_slot = slot
                    report["valid_rows"] = int(report.get("valid_rows", 0) or 0) + 1
                    logger.info(
                        "INSERTED_VALID_ROW row_index=%s section=%s day=%s start_time=%s semester=%s subject=%s faculty=%s",
                        row_index,
                        slot.get("section"),
                        slot.get("day"),
                        slot.get("start_time"),
                        slot.get("semester"),
                        slot.get("subject_name"),
                        slot.get("faculty_name"),
                    )
                    yield slot

    report["detected_days"] = sorted(detected_days)
    report["extracted_subjects_sample"] = detected_subjects[:PDF_DIAG_SAMPLE_CAP]


def _pdf_score_timetable_table(rows: List[List[str]], header_idx: int) -> int:
    if header_idx < 0 or header_idx >= len(rows):
        return 0
    header_cells = rows[header_idx] or []
    time_hits = sum(1 for c in header_cells if _pdf_header_has_time(c))
    return time_hits * 10 + len(rows)


def parse_pdf_to_slots(path: str, stats: Optional[Dict[str, object]] = None) -> Iterator[Dict]:
    if pdfplumber is None:
        raise RuntimeError("pdfplumber is not installed. Install with: pip install pdfplumber")
    report = stats if stats is not None else {}
    report.setdefault("tables_detected", 0)
    report.setdefault("timetable_tables", 0)
    report.setdefault("faculty_tables", 0)
    report.setdefault("detected_section", "")
    report.setdefault("detected_section_raw", "")
    report.setdefault("detected_section_normalized", "")
    report.setdefault("detected_section_source", "")
    report.setdefault("detected_section_source_line", "")
    report.setdefault("rejected_section_candidates", [])
    report.setdefault("section_detection_attempts", [])
    report.setdefault("header_text_samples", [])
    report.setdefault("title_text_samples", [])
    report.setdefault("faculty_text_samples", [])
    report.setdefault("detected_days", [])
    report.setdefault("detected_time_slots", [])
    report.setdefault("extracted_subjects_sample", [])
    report.setdefault("faculty_mappings_sample", [])
    report.setdefault("faculty_mappings_count", 0)
    report.setdefault("validation_errors", [])
    report.setdefault("raw_table_samples", [])
    report.setdefault("detected_sections", [])

    base_name = os.path.splitext(os.path.basename(path))[0]
    section_candidates = []
    semester_hint = _semester_from_text(base_name)

    timetable_tables = []
    faculty_tables = []
    page_sections_by_index: Dict[int, List[str]] = {}
    timetable_tables_by_page: Dict[int, List[Dict]] = {}
    faculty_tables_by_page: Dict[int, List[Dict]] = {}
    page_order: List[int] = []

    with pdfplumber.open(path) as pdf:
        for page_index, page in enumerate(pdf.pages):
            page_order.append(page_index)
            page_sections: List[str] = []

            def _remember_sections(sections: List[str]) -> None:
                for section in sections:
                    if section and section not in page_sections:
                        page_sections.append(section)

            tables = _pdf_find_tables(page)
            report["tables_detected"] += len(tables)

            header_limit = None
            for info in tables:
                bbox = info.get("bbox")
                if bbox and len(bbox) >= 2:
                    header_limit = bbox[1] if header_limit is None else min(header_limit, bbox[1])
            if header_limit is None:
                header_limit = page.height * 0.25

            header_text = ""
            try:
                header_box = (0, 0, page.width, max(0, header_limit))
                header_text = page.within_bbox(header_box).extract_text() or ""
            except Exception:
                header_text = ""

            header_lines = [l for l in header_text.splitlines() if l.strip()]
            if header_text:
                header_sections = _pdf_collect_section_candidates([header_text], report, source=f"header_block_page_{page_index + 1}")
                section_candidates.extend(header_sections)
                _remember_sections(header_sections)
            header_region_sections = _pdf_collect_section_candidates(header_lines, report, source="header_region")
            section_candidates.extend(header_region_sections)
            _remember_sections(header_region_sections)
            if header_text:
                header_samples = report.get("header_text_samples") or []
                excerpt = _pdf_excerpt(header_text)
                if excerpt and excerpt not in header_samples:
                    header_samples.append(excerpt)
                report["header_text_samples"] = header_samples[:PDF_DIAG_SAMPLE_CAP]

            page_text = page.extract_text() or ""
            page_lines = [l for l in page_text.splitlines() if l.strip()]
            title_lines = page_lines[: min(20, len(page_lines))]
            if title_lines:
                title_block_sections = _pdf_collect_section_candidates(["\n".join(title_lines)], report, source=f"title_block_page_{page_index + 1}")
                section_candidates.extend(title_block_sections)
                _remember_sections(title_block_sections)
            title_page_sections = _pdf_collect_section_candidates(title_lines, report, source=f"title_page_{page_index + 1}")
            section_candidates.extend(title_page_sections)
            _remember_sections(title_page_sections)
            if title_lines:
                title_samples = report.get("title_text_samples") or []
                for title_line in title_lines[:3]:
                    excerpt = _pdf_excerpt(title_line)
                    if excerpt and excerpt not in title_samples:
                        title_samples.append(excerpt)
                report["title_text_samples"] = title_samples[:PDF_DIAG_SAMPLE_CAP]

            nearby_sections = _pdf_collect_section_candidates(page_lines, report, source=f"nearby_label_page_{page_index + 1}", require_context=True)
            section_candidates.extend(nearby_sections)
            _remember_sections(nearby_sections)

            table_text_samples = report.get("faculty_text_samples") or []
            if not section_candidates:
                fallback_sections = _pdf_collect_section_candidates(page_lines, report, source=f"page_body_fallback_page_{page_index + 1}")
                section_candidates.extend(fallback_sections)
                _remember_sections(fallback_sections)
            _pdf_collect_rejected_candidates(page_text, report)
            if semester_hint is None:
                semester_hint = _semester_from_text(page_text)
            for info in tables:
                info["page_index"] = page_index
                rows = _pdf_table_extract_matrix(info["table"])
                if not rows:
                    continue
                info["rows"] = rows
                if len(report["raw_table_samples"]) < PDF_DIAG_SAMPLE_CAP:
                    report["raw_table_samples"].append(_pdf_debug_table_sample(rows))
                    if len(report["raw_table_samples"]) == 1:
                        logger.info("PDF raw table preview rows=%s", report["raw_table_samples"][0])
                row_samples = []
                for row in rows[:3]:
                    row_text = _pdf_excerpt(" ".join(cell for cell in row if cell))
                    if row_text:
                        row_samples.append(row_text)
                        if row_text not in table_text_samples:
                            table_text_samples.append(row_text)
                if row_samples:
                    report["faculty_text_samples"] = table_text_samples[:PDF_DIAG_SAMPLE_CAP]
                    table_row_sections = _pdf_collect_section_candidates(row_samples, report, source=f"table_rows_page_{page_index + 1}")
                    section_candidates.extend(table_row_sections)
                    _remember_sections(table_row_sections)
                header_idx, header_cells = _pdf_locate_timetable_header(rows)
                info["header_idx"] = header_idx
                info["header_cells"] = header_cells
                if _pdf_is_faculty_table_rows(rows):
                    if row_samples:
                        faculty_row_sections = _pdf_collect_section_candidates(row_samples, report, source=f"faculty_table_page_{page_index + 1}")
                        section_candidates.extend(faculty_row_sections)
                        _remember_sections(faculty_row_sections)
                    faculty_tables.append(info)
                    faculty_tables_by_page.setdefault(page_index, []).append(info)
                    continue
                if header_idx >= 0:
                    if row_samples:
                        timetable_row_sections = _pdf_collect_section_candidates(row_samples, report, source=f"timetable_table_page_{page_index + 1}")
                        section_candidates.extend(timetable_row_sections)
                        _remember_sections(timetable_row_sections)
                    timetable_tables.append(info)
                    timetable_tables_by_page.setdefault(page_index, []).append(info)
                    continue

            if page_sections:
                page_sections_by_index[page_index] = page_sections

    unique_sections = []
    for section in section_candidates:
        if section and section not in unique_sections:
            unique_sections.append(section)

    if not unique_sections:
        report["validation_errors"].append("Section name not detected in header (expected e.g. CSE-A).")
        if report.get("header_text_samples"):
            report["validation_errors"].append(f"Header text excerpt: {report['header_text_samples'][0]}")
        raise TimetablePDFValidationError("Section name not detected. Ensure the PDF header includes the section (e.g., CSE-A).")

    section_name = unique_sections[0] if unique_sections else (base_section or base_name)
    report["detected_sections"] = unique_sections[:PDF_DIAG_SAMPLE_CAP]
    report["detected_section"] = section_name
    logger.debug(
        "PDF section resolved raw=%s normalized=%s source=%s attempts=%d",
        report.get("detected_section_raw") or "",
        report.get("detected_section_normalized") or section_name,
        report.get("detected_section_source") or "",
        len(report.get("section_detection_attempts") or []),
    )

    if not timetable_tables:
        report["validation_errors"].append("Timetable grid not detected (expected Day column and time slot headers).")
        raise TimetablePDFValidationError("Timetable grid not detected. Ensure the PDF has a timetable table with Day and time-slot headers.")
    if not faculty_tables:
        report["validation_errors"].append("Faculty mapping table not detected (expected Subject Code/Name/Faculty columns).")
        raise TimetablePDFValidationError("Faculty mapping table not detected. Ensure the PDF includes the Subject Code / Subject Name / Faculty Name table below the grid.")

    report["timetable_tables"] = len(timetable_tables)
    report["faculty_tables"] = len(faculty_tables)

    faculty_entries: List[Dict] = []
    for table_info in faculty_tables:
        entries = _pdf_parse_faculty_rows(table_info.get("rows") or [])
        faculty_entries.extend(entries)
    report["faculty_mappings_count"] = len(faculty_entries)
    if faculty_tables and not faculty_entries:
        report["validation_errors"].append("Faculty table detected but no rows parsed (check Subject Code/Name/Faculty columns).")
        raise TimetablePDFValidationError("Faculty mapping table detected but rows could not be parsed. Ensure the columns are Subject Code, Subject Name, and Faculty Name.")
    if faculty_entries:
        report["faculty_mappings_sample"] = faculty_entries[:PDF_DIAG_SAMPLE_CAP]

    faculty_map = _faculty_lookup(faculty_entries)
    last_section_name = section_name
    emitted_slots = 0
    for page_index in page_order:
        page_timetable_tables = timetable_tables_by_page.get(page_index) or []
        if not page_timetable_tables:
            continue
        page_sections = page_sections_by_index.get(page_index) or []
        if page_sections:
            last_section_name = page_sections[0]
        page_section_name = last_section_name or section_name or base_name
        page_branch_name = _section_branch_name(page_section_name, base_name)
        page_faculty_entries: List[Dict] = []
        for table_info in faculty_tables_by_page.get(page_index) or []:
            page_faculty_entries.extend(_pdf_parse_faculty_rows(table_info.get("rows") or []))
        page_faculty_map = _faculty_lookup(page_faculty_entries) if page_faculty_entries else faculty_map
        if page_faculty_entries:
            _merge_faculty_entries(page_faculty_map, faculty_entries)
        for table_info in page_timetable_tables:
            for slot in _pdf_parse_timetable_table(table_info, page_section_name, page_branch_name, semester_hint, page_faculty_map, report):
                emitted_slots += 1
                yield slot

    report["parsed_slot_count"] = emitted_slots


def _row_exists(db, table: str, where_clause: str, params: tuple) -> bool:
    try:
        row = _db_execute(db, f"SELECT 1 FROM {table} WHERE {where_clause} LIMIT 1", params).fetchone()
        return row is not None
    except Exception:
        return False


def _normalize_key(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", _clean_text(value).lower()).strip()


def _first_non_empty(row_map: Dict[str, str], *aliases: str) -> str:
    for alias in aliases:
        value = _clean_text(row_map.get(_normalize_key(alias), ""))
        if value:
            return value
    return ""


def _row_has_token(text: str, *tokens: str) -> bool:
    haystack = _clean_text(text).lower()
    return any(token.lower() in haystack for token in tokens)


def _split_time_value(time_value: str, end_value: str = ""):
    start_time = _clean_text(time_value)
    end_time = _clean_text(end_value)
    if start_time and not end_time:
        if "-" in start_time:
            return _split_time_range(start_time)
        if "to" in start_time.lower():
            parts = re.split(r"\bto\b", start_time, flags=re.IGNORECASE)
            if len(parts) >= 2:
                return _format_time_str(parts[0].strip()) or "", _format_time_str(parts[1].strip()) or ""
    if not start_time and end_time:
        return "", _format_time_str(end_time) or ""
    if start_time and end_time:
        return _format_time_str(start_time) or "", _format_time_str(end_time) or ""
    if start_time:
        return _split_time_range(start_time)
    return "", ""


def _normalize_slot_row(row: Dict[str, str], row_text: str = "") -> Dict[str, str]:
    normalized = {key: _clean_text(value) for key, value in row.items()}
    normalized["branch"] = _normalize_timetable_branch_name(
        _first_non_empty(normalized, "branch", "dept", "department", "program", "course", "branch name"),
        row_text=row_text,
    )
    normalized["section"] = _normalize_timetable_section_name(
        _first_non_empty(normalized, "section", "class", "division", "batch", "group"),
        branch_value=normalized["branch"],
        row_text=row_text,
    )
    if not normalized["branch"] and normalized["section"]:
        normalized["branch"] = _normalize_timetable_branch_name(normalized["section"].split("-", 1)[0], row_text=row_text)
    semester_value = _first_non_empty(normalized, "semester", "sem", "term", "year")
    normalized["semester"] = _safe_int(semester_value)
    normalized["day"] = _normalize_display_text(_first_non_empty(normalized, "day", "weekday", "date"))
    time_value = _first_non_empty(normalized, "time", "slot", "period", "session")
    start_value = _first_non_empty(normalized, "start", "start time", "from", "begin")
    end_value = _first_non_empty(normalized, "end", "end time", "to", "until", "finish")
    if not start_value and not end_value and time_value:
        start_value, end_value = _split_time_range(time_value)
    else:
        start_value, end_value = _split_time_value(start_value or time_value, end_value)
    normalized["start_time"] = start_value
    normalized["end_time"] = end_value
    normalized["subject_name"] = _normalize_timetable_subject_name(
        _first_non_empty(normalized, "subject", "course", "paper", "topic", "title")
    )
    normalized["faculty_name"] = _normalize_timetable_faculty_name(_first_non_empty(normalized, "faculty", "teacher", "instructor", "lecturer", "staff"))
    normalized["room"] = _normalize_display_text(_first_non_empty(normalized, "room", "classroom", "hall", "venue", "lab room"))
    normalized["lab_theory"] = _normalize_timetable_lab_theory(_first_non_empty(normalized, "lab_theory", "lab theory", "type", "category"), row_text=row_text)
    normalized["is_lab"] = int(bool(_row_has_token(normalized["subject_name"], "lab", "practical") or _row_has_token(row_text, "lab", "practical") or _row_has_token(normalized["lab_theory"], "lab", "practical")))
    return normalized


def _valid_slot_row(row: Dict[str, str]) -> bool:
    return _is_valid_academic_timetable_row(row)


_TIMETABLE_BLOCKED_TOKENS = (
    "PRINCIPAL",
    "PRINCIPAL-VICE",
    "VICE",
    "HOD",
    "DEAN",
    "TIMETABLE",
    "SIDDHARTHA",
    "INSTITUTE",
    "TECHNOLOGY",
    "SCIENCES",
    "SHORT BREAK",
    "SHORT BRAEK",
    "LUNCH BREAK",
    "LUNCH",
    "BREAK",
)

_ACADEMIC_DEPARTMENT_CODES = {
    "AIML",
    "AIDS",
    "CIVIL",
    "CSE",
    "CSD",
    "CSM",
    "DS",
    "ECE",
    "EEE",
    "IT",
    "MECH",
}


def _contains_blocked_timetable_text(value: str) -> bool:
    text = _clean_text(value).upper()
    return bool(text) and any(token in text for token in _TIMETABLE_BLOCKED_TOKENS)


def _normalize_academic_department_code(value: str) -> str:
    text = _clean_text(value).upper()
    if not text or _contains_blocked_timetable_text(text):
        return ""
    match = re.fullmatch(r"([A-Z]{2,5})", text)
    if match and match.group(1) in _ACADEMIC_DEPARTMENT_CODES:
        return match.group(1)
    return ""


def _normalize_academic_section_code(value: str) -> str:
    text = _clean_text(value).upper()
    if not text or _contains_blocked_timetable_text(text):
        return ""
    text = re.sub(r"[\s_/]+", "-", text)
    text = re.sub(r"-+", "-", text).strip("-")
    match = re.match(r"^([A-Z]{2,5})-([A-Z0-9]+)$", text)
    if not match:
        return ""
    department_code = match.group(1)
    if department_code not in _ACADEMIC_DEPARTMENT_CODES:
        return ""
    return f"{department_code}-{match.group(2)}"


def _is_valid_academic_timetable_row(row: Dict[str, str]) -> bool:
    branch_value, section_value = _normalize_timetable_branch_section(row.get("branch"), row.get("section"), row_text=_clean_text(row.get("day")) + " " + _clean_text(row.get("subject_name")))
    subject_value = _normalize_timetable_subject_name(row.get("subject_name"))
    faculty_value = _clean_text(row.get("faculty_name"))
    day_value = _clean_text(row.get("day"))
    start_value = _clean_text(row.get("start_time"))
    end_value = _clean_text(row.get("end_time"))

    if not (day_value and start_value and end_value):
        return False
    if not subject_value or not faculty_value:
        return False
    if _contains_blocked_timetable_text(branch_value) or _contains_blocked_timetable_text(section_value):
        return False

    department_code = _normalize_academic_department_code(branch_value)
    if not department_code and section_value:
        department_code = _normalize_academic_department_code(section_value.split("-", 1)[0])
    section_code = _normalize_academic_section_code(section_value)
    if not section_code and branch_value and section_value and "-" in section_value:
        section_code = _normalize_academic_section_code(section_value)
    if not section_code:
        return False
    if department_code not in _ACADEMIC_DEPARTMENT_CODES:
        return False
    if not section_code.startswith(f"{department_code}-"):
        return False
    return True


def is_valid_timetable_row(row):
    return _is_valid_academic_timetable_row(row)


def _split_subjects(subj_raw: str):
    """Split merged subject strings into individual subject names.

    Examples: 'PYTHON/AEP LAB' -> ['PYTHON LAB', 'AEP LAB']
    """
    if not subj_raw:
        return [""]
    s = subj_raw.strip()
    # normalize separators
    parts = [p.strip() for p in re.split(r"[/,&]", s) if p.strip()]
    results = []
    for p in parts:
        # if original had LAB suffix, preserve it
        if re.search(r"\blab\b", s, flags=re.IGNORECASE) and not re.search(r"\blab\b", p, flags=re.IGNORECASE):
            results.append((p + " LAB").strip())
        else:
            results.append(p)
    return results


# --- Utilities --------------------------------------------------------------

def _safe_int(v):
    try:
        return int(str(v).strip())
    except Exception:
        return None


def _normalize_time(val: str) -> str:
    if not val:
        return ""
    val = val.strip()
    # Accept formats like 09:00-10:00 or 09.00 - 10.00 or 9 AM - 10 AM
    if "-" in val:
        a, b = val.split("-", 1)
        return _format_time_str(a.strip()) or "", _format_time_str(b.strip()) or ""
    # Single time -> treat as start
    return _format_time_str(val)


def _format_time_str(s: str) -> Optional[str]:
    s = s.strip()
    s = re.sub(r"\s+", " ", s)
    s = re.sub(r"(?i)(\d)(am|pm)\b", r"\1 \2", s)
    # Try HH:MM
    for fmt in ("%H:%M", "%I:%M %p", "%I %p", "%H.%M"):
        try:
            dt = datetime.strptime(s, fmt)
            return dt.strftime("%H:%M")
        except Exception:
            continue
    # Numeric hour
    try:
        h = int(s)
        return f"{h:02d}:00"
    except Exception:
        return None


def _split_time_range(s: str):
    s = _clean_text(s)
    s = re.sub(r"(?i)(\d)(am|pm)\b", r"\1 \2", s)
    if "-" in s:
        a, b = s.split("-", 1)
        return _format_time_str(a.strip()) or "", _format_time_str(b.strip()) or ""
    return (_format_time_str(s) or "", "")


def _current_local_datetime(now: Optional[datetime] = None) -> datetime:
    """Return a timezone-aware local datetime for timetable comparisons."""
    if now is None:
        return datetime.now().astimezone()
    if getattr(now, "tzinfo", None) is None:
        return now.astimezone()
    return now.astimezone()


def _row_to_dict(row):
    try:
        return dict(row)
    except Exception:
        return row


def _lookup_branch_id(db, branch_name: str):
    if not branch_name:
        return None
    try:
        row = _db_execute(db, "SELECT id FROM branches WHERE LOWER(name)=LOWER(?) LIMIT 1", (branch_name,)).fetchone()
        if row:
            return row[0] if not hasattr(row, "keys") else row["id"]
    except Exception:
        return None
    return None


def _resolve_or_create_branch_id(db, branch_name: str, branch_cache: Optional[Dict[str, Optional[int]]] = None):
    branch_name = _clean_text(branch_name)
    if not branch_name:
        return None

    cache_key = branch_name.lower()
    if branch_cache is not None and cache_key in branch_cache:
        return branch_cache[cache_key]

    candidates = [branch_name]
    branch_part, _ = split_branch_section(branch_name)
    if branch_part and branch_part != branch_name:
        candidates.append(branch_part)
    normalized_branch = _normalize_timetable_branch_name(branch_name)
    if normalized_branch and normalized_branch not in candidates:
        candidates.append(normalized_branch)
    if "-" in branch_name:
        candidates.append(_clean_text(branch_name.split("-", 1)[0]))
    if " " in branch_name:
        candidates.append(_clean_text(branch_name.split(" ", 1)[0]))
    if branch_name.endswith("SECTION"):
        candidates.append(branch_name[:-7].strip())

    for candidate in candidates:
        if not candidate:
            continue
        branch_id = _lookup_branch_id(db, candidate)
        if branch_id is not None:
            if branch_cache is not None:
                branch_cache[cache_key] = branch_id
                branch_cache[candidate.lower()] = branch_id
            return branch_id

    try:
        _db_execute(db, _insert_ignore_sql(db, "branches", ["name", "location"]), (branch_name, "Auto-imported timetable branch"))
        try:
            db.commit()
        except Exception as commit_err:
            logger.warning("Failed to commit branch creation: branch=%s error=%s", branch_name, commit_err)
            pass
    except Exception as create_err:
        logger.exception("Failed to auto-create branch: branch=%s error=%s", branch_name, create_err)
        return None

    branch_id = _lookup_branch_id(db, branch_name)
    if branch_id is None:
        logger.warning("Branch creation or lookup failed: branch=%s cache_key=%s", branch_name, cache_key)
    if branch_cache is not None:
        branch_cache[cache_key] = branch_id
    return branch_id


def _normalize_subject_name(value: str) -> str:
    text = _clean_text(value).lower()
    text = text.replace("&", " and ")
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _normalize_teacher_name(value: str) -> str:
    text = _clean_text(value).lower()
    # Ignore punctuation/whitespace variants like "Dr. A. Kumar" vs "dr a kumar".
    return re.sub(r"[^a-z0-9]", "", text)


def _subject_lookup_variants(subject_name: str) -> List[str]:
    raw = _clean_text(subject_name)
    if not raw:
        return []
    variants = []

    def _add(v: str):
        v = _normalize_subject_name(v)
        if v and v not in variants:
            variants.append(v)

    _add(raw)
    _add(raw.replace("LAB", "").replace("lab", ""))
    for part in re.split(r"[/,&]", raw):
        _add(part)
        _add(part.replace("LAB", "").replace("lab", ""))
    acronym = _subject_acronym(raw)
    if acronym:
        _add(acronym)
    return variants


def _build_subject_lookup_index(db) -> Dict[str, int]:
    index: Dict[str, int] = {}
    try:
        rows = _db_execute(db, "SELECT id, name FROM subjects").fetchall()
    except Exception:
        return index
    for row in rows:
        subject_id = row[0] if not hasattr(row, "keys") else row["id"]
        subject_name = row[1] if not hasattr(row, "keys") else row["name"]
        for variant in _subject_lookup_variants(subject_name):
            if variant not in index:
                index[variant] = int(subject_id)

    try:
        alias_rows = _db_execute(db, "SELECT alias, canonical_name FROM subject_aliases").fetchall()
    except Exception:
        alias_rows = []
    if alias_rows:
        subject_rows = { _normalize_subject_name(row[1] if not hasattr(row, "keys") else row["name"]): int(row[0] if not hasattr(row, "keys") else row["id"]) for row in rows }
        for row in alias_rows:
            alias = _normalize_subject_name(row[0] if not hasattr(row, "keys") else row["alias"])
            canonical = _normalize_subject_name(row[1] if not hasattr(row, "keys") else row["canonical_name"])
            subject_id = subject_rows.get(canonical)
            if subject_id is None:
                continue
            for variant in _subject_lookup_variants(alias) + _subject_lookup_variants(canonical):
                if variant and variant not in index:
                    index[variant] = subject_id
    return index


def _resolve_subject_id(subject_name: str, subject_cache: Dict[str, Optional[int]], subject_index: Dict[str, int]) -> Optional[int]:
    cache_key = _normalize_subject_name(subject_name)
    if cache_key in subject_cache:
        return subject_cache[cache_key]
    for variant in _subject_lookup_variants(subject_name):
        subject_id = subject_index.get(variant)
        if subject_id is not None:
            subject_cache[cache_key] = subject_id
            return subject_id
    subject_cache[cache_key] = None
    return None


def _build_teacher_lookup_index(db) -> Dict[str, int]:
    index: Dict[str, int] = {}
    try:
        rows = _db_execute(db, "SELECT id, name FROM teachers").fetchall()
    except Exception:
        return index
    for row in rows:
        teacher_id = row[0] if not hasattr(row, "keys") else row["id"]
        teacher_name = row[1] if not hasattr(row, "keys") else row["name"]
        norm_name = _normalize_teacher_name(teacher_name)
        if norm_name and norm_name not in index:
            index[norm_name] = int(teacher_id)
    return index


def _build_teacher_name_index(db) -> Dict[str, str]:
    """Return mapping of normalized teacher name -> canonical display name."""
    idx: Dict[str, str] = {}
    try:
        rows = _db_execute(db, "SELECT id, name FROM teachers").fetchall()
    except Exception:
        return idx
    for row in rows:
        teacher_name = row[1] if not hasattr(row, "keys") else row["name"]
        norm_name = _normalize_teacher_name(teacher_name)
        if norm_name and norm_name not in idx:
            idx[norm_name] = teacher_name
    return idx


def _resolve_teacher_id(teacher_name: str, teacher_cache: Dict[str, Optional[int]], teacher_index: Dict[str, int]) -> Optional[int]:
    cache_key = _normalize_teacher_name(teacher_name)
    if not cache_key:
        return None
    if cache_key in teacher_cache:
        return teacher_cache[cache_key]
    # Exact match first
    teacher_id = teacher_index.get(cache_key)
    if teacher_id is not None:
        teacher_cache[cache_key] = teacher_id
        return teacher_id
    # Fuzzy match against known teacher normalized names
    try:
        import difflib

        candidates = difflib.get_close_matches(cache_key, list(teacher_index.keys()), n=1, cutoff=0.7)
        if candidates:
            teacher_id = teacher_index.get(candidates[0])
            teacher_cache[cache_key] = teacher_id
            return teacher_id
    except Exception:
        pass
    teacher_cache[cache_key] = None
    return None


def _refresh_timetable_entry_ids(db) -> Dict[str, int]:
    subject_index = _build_subject_lookup_index(db)
    teacher_index = _build_teacher_lookup_index(db)
    subject_cache: Dict[str, Optional[int]] = {}
    teacher_cache: Dict[str, Optional[int]] = {}
    update_count = 0
    total_rows = 0
    try:
        rows = _db_execute(db, "SELECT id, subject_name, faculty_name FROM timetable_entries WHERE subject_id IS NULL OR teacher_id IS NULL").fetchall()
        for row in rows:
            total_rows += 1
            entry_id = row[0] if not hasattr(row, 'keys') else row['id']
            subject_name = row[1] if not hasattr(row, 'keys') else row['subject_name']
            faculty_name = row[2] if not hasattr(row, 'keys') else row['faculty_name']
            subject_id = _resolve_subject_id(subject_name or "", subject_cache, subject_index) if subject_name else None
            teacher_id = _resolve_teacher_id(faculty_name or "", teacher_cache, teacher_index) if faculty_name else None
            if subject_id is None and teacher_id is None:
                continue
            updates = []
            params = []
            if subject_id is not None:
                updates.append("subject_id = %s")
                params.append(subject_id)
            if teacher_id is not None:
                updates.append("teacher_id = %s")
                params.append(teacher_id)
            params.append(entry_id)
            try:
                _db_execute(db, f"UPDATE timetable_entries SET {', '.join(updates)} WHERE id = %s", tuple(params))
                update_count += 1
            except Exception:
                logger.exception("Failed to refresh timetable entry ids for entry_id=%s", entry_id)
        try:
            db.commit()
        except Exception:
            pass
    except Exception:
        logger.exception("Failed to refresh timetable entries subject/teacher IDs")
    return {"total_rows": total_rows, "updated": update_count}


def get_upcoming_classes(db, branch: str = "", section: str = "", limit: int = 3, now: Optional[datetime] = None):
    """Return upcoming classes for the current day using normalized data first."""
    current = _current_local_datetime(now)
    weekday = current.strftime("%A")
    cur_time = current.strftime("%H:%M")
    branch_id = _lookup_branch_id(db, branch)

    entries = []
    try:
        if branch_id is not None:
            rows = _db_execute(db,
                "SELECT te.*, s.name AS subject_name, t.name AS teacher_name, b.name AS branch_name FROM timetable_entries te LEFT JOIN subjects s ON te.subject_id = s.id LEFT JOIN teachers t ON te.teacher_id = t.id LEFT JOIN branches b ON te.branch_id = b.id WHERE te.branch_id = ? AND COALESCE(te.section, '') = ? AND te.day = ? AND te.start_time >= ? ORDER BY te.start_time LIMIT ?",
                (branch_id, section or "", weekday, cur_time, limit),
            ).fetchall()
            entries = [_row_to_dict(r) for r in rows]
            logger.info("upcoming normalized query result count=%d branch=%s section=%s day=%s", len(entries), branch, section, weekday)
    except Exception:
        entries = []

    if not entries:
        try:
            rows = _db_execute(db,
                "SELECT * FROM timetable_slots WHERE branch = ? AND section = ? AND day = ? AND start_time >= ? ORDER BY start_time LIMIT ?",
                (branch or "", section or "", weekday, cur_time, limit),
            ).fetchall()
            entries = [_row_to_dict(r) for r in rows]
            logger.info("upcoming legacy query result count=%d branch=%s section=%s day=%s", len(entries), branch, section, weekday)
        except Exception:
            entries = []

    return entries


def _infer_dashboard_scope(db) -> tuple[str, str]:
    """Choose a default branch/section for dashboard queries when session scope is absent."""
    try:
        row = _db_execute(
            db,
            """
            SELECT COALESCE(b.name, '') AS branch_name, COALESCE(te.section, '') AS section_name, COUNT(1) AS c
            FROM timetable_entries te
            LEFT JOIN branches b ON te.branch_id = b.id
            GROUP BY COALESCE(b.name, ''), COALESCE(te.section, '')
            ORDER BY c DESC
            LIMIT 1
            """,
        ).fetchone()
        if row:
            branch_name = _clean_text(row[0] if not hasattr(row, "keys") else row["branch_name"])
            section_name = _clean_text(row[1] if not hasattr(row, "keys") else row["section_name"])
            if branch_name or section_name:
                return branch_name, section_name
    except Exception:
        pass
    try:
        row = _db_execute(
            db,
            """
            SELECT COALESCE(branch, '') AS branch_name, COALESCE(section, '') AS section_name, COUNT(1) AS c
            FROM timetable_slots
            GROUP BY COALESCE(branch, ''), COALESCE(section, '')
            ORDER BY c DESC
            LIMIT 1
            """,
        ).fetchone()
        if row:
            branch_name = _clean_text(row[0] if not hasattr(row, "keys") else row["branch_name"])
            section_name = _clean_text(row[1] if not hasattr(row, "keys") else row["section_name"])
            return branch_name, section_name
    except Exception:
        pass
    return "", ""


def get_current_active_class(db, branch: str = "", section: str = "", now: Optional[datetime] = None):
    """Return the currently running class, preferring normalized timetable_entries."""
    return get_current_slot(db, branch, section, now=now)


def get_global_active_class(db, now: Optional[datetime] = None):
    """Return the first active class across the institute for the current time."""
    current = _current_local_datetime(now)
    weekday = current.strftime("%A")
    cur_time = current.strftime("%H:%M")
    try:
        rows = _db_execute(db,
            "SELECT te.*, s.name AS subject_name, t.name AS teacher_name, b.name AS branch_name FROM timetable_entries te LEFT JOIN subjects s ON te.subject_id = s.id LEFT JOIN teachers t ON te.teacher_id = t.id LEFT JOIN branches b ON te.branch_id = b.id WHERE te.day = ? AND te.start_time <= ? AND te.end_time >= ? ORDER BY te.start_time LIMIT 1",
            (weekday, cur_time, cur_time),
        ).fetchall()
        if rows:
            return _row_to_dict(rows[0])
    except Exception:
        pass
    try:
        rows = _db_execute(db,
            "SELECT * FROM timetable_slots WHERE day = ? AND start_time <= ? AND end_time >= ? ORDER BY start_time LIMIT 1",
            (weekday, cur_time, cur_time),
        ).fetchall()
        if rows:
            return _row_to_dict(rows[0])
    except Exception:
        pass
    return None


def get_faculty_schedule(db, teacher_id, now: Optional[datetime] = None):
    current = _current_local_datetime(now)
    weekday = current.strftime("%A")
    rows = []
    try:
        rows = _db_execute(db,
            "SELECT te.*, s.name AS subject_name, b.name AS branch_name, t.name AS teacher_name FROM timetable_entries te LEFT JOIN subjects s ON te.subject_id = s.id LEFT JOIN branches b ON te.branch_id = b.id LEFT JOIN teachers t ON te.teacher_id = t.id WHERE te.teacher_id = ? AND te.day = ? ORDER BY te.start_time",
            (teacher_id, weekday),
        ).fetchall()
    except Exception:
        rows = []

    if rows:
        return [_row_to_dict(r) for r in rows]

    try:
        rows = _db_execute(db,
            "SELECT * FROM timetable_slots WHERE faculty_name IS NOT NULL AND day = ? ORDER BY start_time",
            (weekday,),
        ).fetchall()
    except Exception:
        rows = []
    return [_row_to_dict(r) for r in rows]


# --- DB import --------------------------------------------------------------

def _write_preview_line(preview_path: str, preview_state: Dict[str, int], payload: Dict):
    if not preview_path or PREVIEW_ROW_CAP <= 0:
        return
    if preview_state["written"] >= PREVIEW_ROW_CAP:
        return
    try:
        with open(preview_path, "a", encoding="utf-8") as pf:
            pf.write(json.dumps(payload, default=str) + "\n")
        preview_state["written"] += 1
    except Exception:
        logger.exception("Failed to write import preview line")


def _iter_docx_slots_with_fallback(path: str) -> Iterator[Dict]:
    yield from iter_docx_section_slots(path)


def import_slots_streaming(db, slots_iter: Iterable[Dict]):
    """Import slots from an iterator to avoid retaining full parsed timetable in memory."""
    raw_counters = {
        "processed": 0,
        "inserted": 0,
        "skipped_total": 0,
        "skipped_invalid": 0,
        "skipped_duplicate": 0,
        "failures": 0,
    }
    normalized_counters = {
        "processed": 0,
        "inserted": 0,
        "skipped_total": 0,
        "skipped_branch": 0,
        "skipped_duplicate": 0,
        "missing_subjects": 0,
        "missing_teachers": 0,
        "failures": 0,
        "skip_reasons": {},
    }
    normalized_diagnostics = {
        "unresolved_subject_rows": [],
        "unresolved_teacher_rows": [],
    }
    unresolved_sample_cap = 40

    preview_path = None
    preview_state = {"written": 0}
    if PREVIEW_ROW_CAP > 0:
        preview_path = os.path.join(os.path.dirname(__file__), "uploads", "last_import_debug.jsonl")
        os.makedirs(os.path.dirname(preview_path), exist_ok=True)

    # Commit every 5-10 rows for low-memory environments
    batch_size = max(5, min(10, int(BATCH_INSERT_SIZE)))
    inserted_since_commit = 0
    batch_commits = 0
    start_ts = time.time()
    peak_mem_estimate_bytes = 0

    branch_cache = {}
    subject_cache = {}
    teacher_cache = {}
    seen_norm_keys = set()
    subject_index = _build_subject_lookup_index(db)
    teacher_index = _build_teacher_lookup_index(db)
    branch_subject_counts: Dict[str, set] = {}
    for row_index, slot_in in enumerate(slots_iter, start=1):
        mem_estimate = (
            (len(branch_cache) + len(subject_cache) + len(teacher_cache)) * 80
            + preview_state["written"] * 64
        )
        if mem_estimate > peak_mem_estimate_bytes:
            peak_mem_estimate_bytes = mem_estimate

        row = {
            "branch": _clean_text(slot_in.get("branch")),
            "section": _clean_text(slot_in.get("section")),
            "semester": _safe_int(slot_in.get("semester")),
            "day": _clean_text(slot_in.get("day")),
            "start_time": _clean_text(slot_in.get("start_time")),
            "end_time": _clean_text(slot_in.get("end_time")),
            "subject_name": _clean_text(slot_in.get("subject_name")),
            "faculty_name": _clean_text(slot_in.get("faculty_name")),
            "is_lab": int(bool(slot_in.get("is_lab"))),
            "room": _clean_text(slot_in.get("room")),
        }
        print("NORMALIZED_ROW", row)
        logger.info("Streaming normalized row %s: %s", row_index, row)
        raw_counters["processed"] += 1

        # If the row is completely empty, skip immediately
        if not any(_clean_text(value) for value in row.values()):
            raw_counters["skipped_total"] += 1
            raw_counters["skipped_invalid"] += 1
            _write_preview_line(preview_path, preview_state, {"index": row_index, "reason": "completely_empty_row", "row": row})
            continue

        # Core validation: require a real academic section plus subject and faculty.
        if not _is_valid_academic_timetable_row(row):
            raw_counters["skipped_total"] += 1
            raw_counters["skipped_invalid"] += 1
            _write_preview_line(preview_path, preview_state, {"index": row_index, "reason": "invalid_row", "row": row})
            continue

        branch_subject_key = f"{row['branch']}|{row['section']}"
        branch_subject_counts.setdefault(branch_subject_key, set()).add(row["subject_name"])

        try:
            cur = _db_execute(
                db,
                _insert_ignore_sql(db, "timetable_slots", ["branch", "section", "semester", "day", "start_time", "end_time", "subject_name", "faculty_name", "is_lab", "room"]),
                (row["branch"], row["section"], row["semester"], row["day"], row["start_time"], row["end_time"], row["subject_name"], row["faculty_name"], row["is_lab"], row["room"]),
            )
            if hasattr(cur, "rowcount") and int(cur.rowcount or 0) == 0:
                raw_counters["skipped_total"] += 1
                raw_counters["skipped_duplicate"] += 1
            else:
                raw_counters["inserted"] += 1
                inserted_since_commit += 1
        except Exception as e:
            if "unique" in str(e).lower() or "duplicate" in str(e).lower():
                raw_counters["skipped_total"] += 1
                raw_counters["skipped_duplicate"] += 1
                continue
            raw_counters["failures"] += 1
            _write_preview_line(preview_path, preview_state, {"index": row_index, "reason": "raw_insert_exception", "row": row})
            raise

        normalized_counters["processed"] += 1
        branch_key = row["branch"].strip().lower()
        if branch_key not in branch_cache:
            branch_cache[branch_key] = _resolve_or_create_branch_id(db, row["branch"], branch_cache)
        branch_id = branch_cache.get(branch_key)
        if branch_id is None:
            normalized_counters["skipped_total"] += 1
            normalized_counters["skipped_branch"] += 1
            _write_preview_line(preview_path, preview_state, {"index": row_index, "reason": "missing_branch", "row": row})
        else:
            subject_id = _resolve_subject_id(row["subject_name"], subject_cache, subject_index)
            if subject_id is None:
                normalized_counters["missing_subjects"] += 1
                if len(normalized_diagnostics["unresolved_subject_rows"]) < unresolved_sample_cap:
                    normalized_diagnostics["unresolved_subject_rows"].append(
                        {
                            "index": row_index,
                            "section": row["section"],
                            "day": row["day"],
                            "start_time": row["start_time"],
                            "subject_name": row["subject_name"],
                            "subject_variants": _subject_lookup_variants(row["subject_name"]),
                        }
                    )
                _write_preview_line(preview_path, preview_state, {"index": row_index, "reason": "missing_subject_id", "row": row})

            teacher_id = _resolve_teacher_id(row["faculty_name"], teacher_cache, teacher_index)
            if row["faculty_name"] and teacher_id is None:
                normalized_counters["missing_teachers"] += 1
                if len(normalized_diagnostics["unresolved_teacher_rows"]) < unresolved_sample_cap:
                    normalized_diagnostics["unresolved_teacher_rows"].append(
                        {
                            "index": row_index,
                            "section": row["section"],
                            "day": row["day"],
                            "start_time": row["start_time"],
                            "faculty_name": row["faculty_name"],
                            "teacher_normalized": _normalize_teacher_name(row["faculty_name"]),
                        }
                    )
                _write_preview_line(preview_path, preview_state, {"index": row_index, "reason": "missing_teacher_id", "row": row})

            norm_row = {
                "branch_id": branch_id,
                "section": row["section"],
                "semester": row["semester"],
                "day": row["day"],
                "start_time": row["start_time"],
                "end_time": row["end_time"],
                "subject_name": row["subject_name"],
                "faculty_name": row["faculty_name"],
                "subject_id": subject_id,
                "teacher_id": teacher_id,
                "is_lab": row["is_lab"],
                "room": row["room"],
            }
            # Keep unresolved-ID rows distinct by folding normalized subject/faculty text
            # into the in-memory dedupe key while preserving canonical IDs when available.
            norm_key = _dup_key(
                norm_row["branch_id"],
                norm_row["day"],
                norm_row["start_time"],
                norm_row["end_time"],
                norm_row["section"],
                norm_row["subject_id"] if norm_row["subject_id"] is not None else f"sub:{_normalize_subject_name(row['subject_name'])}",
                norm_row["teacher_id"] if norm_row["teacher_id"] is not None else f"teach:{_normalize_teacher_name(row['faculty_name'])}",
            )
            if norm_key in seen_norm_keys:
                normalized_counters["skipped_total"] += 1
                normalized_counters["skipped_duplicate"] += 1
                normalized_counters["skip_reasons"]["in_memory_norm_key"] = normalized_counters["skip_reasons"].get("in_memory_norm_key", 0) + 1
                _write_preview_line(preview_path, preview_state, {"index": row_index, "reason": "normalized_in_memory_duplicate", "row": norm_row})
                continue
            seen_norm_keys.add(norm_key)
            try:
                cur = _db_execute(
                    db,
                    _insert_ignore_sql(db, "timetable_entries", ["branch_id", "section", "semester", "day", "start_time", "end_time", "subject_id", "teacher_id", "subject_name", "faculty_name", "is_lab", "room"]),
                    (norm_row["branch_id"], norm_row["section"], norm_row["semester"], norm_row["day"], norm_row["start_time"], norm_row["end_time"], norm_row["subject_id"], norm_row["teacher_id"], norm_row["subject_name"], norm_row["faculty_name"], norm_row["is_lab"], norm_row["room"]),
                )
                if hasattr(cur, "rowcount") and int(cur.rowcount or 0) == 0:
                    normalized_counters["skipped_total"] += 1
                    normalized_counters["skipped_duplicate"] += 1
                    normalized_counters["skip_reasons"]["db_conflict_or_ignore"] = normalized_counters["skip_reasons"].get("db_conflict_or_ignore", 0) + 1
                    _write_preview_line(preview_path, preview_state, {"index": row_index, "reason": "normalized_db_duplicate", "row": norm_row})
                else:
                    normalized_counters["inserted"] += 1
                    inserted_since_commit += 1
                    branch_subject_counts.setdefault(branch_subject_key, set()).add(row["subject_name"])
                    logger.info(
                        "inserted_row row_index=%s branch_id=%s section=%s semester=%s day=%s start_time=%s end_time=%s subject_id=%s teacher_id=%s room=%s",
                        row_index,
                        norm_row["branch_id"],
                        norm_row["section"],
                        norm_row["semester"],
                        norm_row["day"],
                        norm_row["start_time"],
                        norm_row["end_time"],
                        norm_row["subject_id"],
                        norm_row["teacher_id"],
                        norm_row["room"],
                    )
            except Exception as e:
                if "unique" in str(e).lower() or "duplicate" in str(e).lower():
                    normalized_counters["skipped_total"] += 1
                    normalized_counters["skipped_duplicate"] += 1
                    normalized_counters["skip_reasons"]["db_unique_exception"] = normalized_counters["skip_reasons"].get("db_unique_exception", 0) + 1
                else:
                    normalized_counters["failures"] += 1
                    _write_preview_line(preview_path, preview_state, {"index": row_index, "reason": "normalized_insert_exception", "row": norm_row})
                    raise

        if inserted_since_commit >= batch_size:
            try:
                db.commit()
                batch_commits += 1
                inserted_since_commit = 0
                if len(branch_cache) > 64 or len(subject_cache) > 64 or len(teacher_cache) > 64:
                    logger.debug("Clearing timetable import caches (branch=%d, subject=%d, teacher=%d)", len(branch_cache), len(subject_cache), len(teacher_cache))
                    branch_cache.clear()
                    subject_cache.clear()
                    teacher_cache.clear()
                logger.info(
                    "import batch commit section=%s processed=%d commits=%d mem_estimate_kb=%.1f",
                    row.get("section"),
                    row_index,
                    batch_commits,
                    mem_estimate / 1024.0,
                )
                _write_preview_line(
                    preview_path,
                    preview_state,
                    {"type": "batch_commit", "processed": row_index, "raw_inserted": raw_counters["inserted"], "normalized_inserted": normalized_counters["inserted"], "timestamp": time.time()},
                )
            except Exception as batch_commit_err:
                logger.exception("Batch commit failed at row_index=%d: %s", row_index, batch_commit_err)
                try:
                    db.rollback()
                    logger.info("Rolled back failed batch at row_index=%d", row_index)
                except:
                    pass
                raise

        if row_index % max(10, batch_size) == 0:
            elapsed = time.time() - start_ts
            rate = row_index / elapsed if elapsed > 0 else 0
            logger.info(
                "import_slots_streaming progress section=%s processed=%d raw_inserted=%d normalized_inserted=%d skipped=%d commits=%d elapsed=%.1fs rate=%.1f rows/s peak_mem_estimate_kb=%.1f",
                row["section"],
                row_index,
                raw_counters["inserted"],
                normalized_counters["inserted"],
                raw_counters["skipped_total"] + normalized_counters["skipped_total"],
                batch_commits,
                elapsed,
                rate,
                peak_mem_estimate_bytes / 1024.0,
            )

    if inserted_since_commit > 0:
        try:
            db.commit()
            batch_commits += 1
            logger.info("Final batch commit successful: row_count=%d", inserted_since_commit)
        except Exception as final_commit_err:
            logger.exception("Final batch commit failed: %s", final_commit_err)
            try:
                db.rollback()
                logger.info("Rolled back final batch after commit failure")
            except:
                pass
            raise

    # Final cleanup to free memory
    branch_cache.clear()
    subject_cache.clear()
    teacher_cache.clear()
    seen_norm_keys.clear()
    gc.collect()

    final_mem_estimate = (
        (len(branch_cache) + len(subject_cache) + len(teacher_cache)) * 80
        + preview_state["written"] * 64
    )
    if final_mem_estimate > peak_mem_estimate_bytes:
        peak_mem_estimate_bytes = final_mem_estimate

    elapsed_seconds = time.time() - start_ts
    try:
        logger.info("import_slots_streaming branch subjects: %s", {key: sorted(values) for key, values in branch_subject_counts.items()})
    except Exception:
        logger.exception("Failed to log streaming branch subject summary")
    return {
        "raw_insert": {"counters": raw_counters},
        "normalized_insert": {"counters": normalized_counters, "diagnostics": normalized_diagnostics, "branch_subject_counts": {key: sorted(values) for key, values in branch_subject_counts.items()}},
        "preview_path": preview_path if preview_state["written"] else None,
        "preview_written": preview_state["written"],
        "batch_commits": batch_commits,
        "elapsed_seconds": elapsed_seconds,
        "peak_memory_estimate_bytes": peak_mem_estimate_bytes,
    }

def import_slots(db, slots: List[Dict]):
    """Import raw timetable slots into `timetable_slots` with verbose diagnostics.

    Returns a dict: {"counters": {...}, "skipped_rows": [...]}
    """
    counters = {
        "total": len(slots),
        "parsed": 0,
        "inserted": 0,
        "skipped_total": 0,
        "skipped_invalid": 0,
        "skipped_duplicate": 0,
        "failures": 0,
        "invalid_section": 0,
        "normalization_failures": 0,
    }
    skipped_rows = []
    skipped_rows_omitted = 0
    preview_path = os.path.join(os.path.dirname(__file__), "uploads", "last_import_debug.jsonl")
    os.makedirs(os.path.dirname(preview_path), exist_ok=True)
    preview_written = 0
    batch_commit_counts = 0
    progress_log_interval = max(50, BATCH_INSERT_SIZE)
    start_time = time.time()
    use_tracemalloc = False
    if ENABLE_IMPORT_TRACEMALLOC:
        try:
            tracemalloc.start()
            use_tracemalloc = True
        except Exception:
            use_tracemalloc = False

    try:
        logger.info("import_slots: parsed_rows_count=%d", len(slots))
        branch_subject_counts: Dict[str, set] = {}
        inserted_since_commit = 0
        for row_index, s in enumerate(slots, start=1):
            counters["parsed"] += 1
            row = {
                "branch": _clean_text(s.get("branch")),
                "section": _clean_text(s.get("section")),
                "semester": _safe_int(s.get("semester")),
                "day": _clean_text(s.get("day")),
                "start_time": _clean_text(s.get("start_time")),
                "end_time": _clean_text(s.get("end_time")),
                "subject_name": _clean_text(s.get("subject_name")),
                "faculty_name": _clean_text(s.get("faculty_name")),
                "is_lab": int(bool(s.get("is_lab"))),
                "room": _clean_text(s.get("room")),
            }

            logger.info("Processing row %s: raw=%s normalized=%s", row_index, s, row)

            if not _is_valid_academic_timetable_row(row):
                counters["skipped_total"] += 1
                counters["skipped_invalid"] += 1
                reason = "invalid_row"
                if not _clean_text(row.get("section")):
                    counters["invalid_section"] += 1
                    reason = "missing_section"
                logger.info("import_slots skipped row %s: reason=%s normalized=%s", row_index, reason, row)
                # stream skipped row to preview file (capped)
                if preview_written < PREVIEW_ROW_CAP:
                    try:
                        with open(preview_path, "a", encoding="utf-8") as pf:
                            pf.write(json.dumps({"index": row_index, "raw": s, "normalized": row, "reason": reason}, default=str) + "\n")
                        preview_written += 1
                    except Exception:
                        logger.exception("Failed to write skipped preview line")
                # keep a small in-memory sample for immediate return
                skipped_rows_omitted = _append_skipped_sample(skipped_rows, skipped_rows_omitted, {"index": row_index, "raw": s, "normalized": row, "reason": reason})
                continue

            branch_subject_key = f"{row['branch']}|{row['section']}"
            branch_subject_counts.setdefault(branch_subject_key, set()).add(row["subject_name"])

            duplicate_where = "COALESCE(branch, '') = COALESCE(%s, '') AND COALESCE(section, '') = COALESCE(%s, '') AND COALESCE(day, '') = COALESCE(%s, '') AND COALESCE(start_time, '') = COALESCE(%s, '') AND COALESCE(end_time, '') = COALESCE(%s, '') AND COALESCE(subject_name, '') = COALESCE(%s, '') AND COALESCE(faculty_name, '') = COALESCE(%s, '')"
            try:
                if _row_exists(
                    db,
                    "timetable_slots",
                    duplicate_where,
                    (row['branch'], row['section'], row['day'], row['start_time'], row['end_time'], row['subject_name'], row['faculty_name']),
                ):
                    counters["skipped_total"] += 1
                    counters["skipped_duplicate"] += 1
                    reason = "duplicate"
                    logger.info("import_slots skipped duplicate row %s: %s", row_index, row)
                    skipped_rows_omitted = _append_skipped_sample(skipped_rows, skipped_rows_omitted, {"index": row_index, "raw": s, "normalized": row, "reason": reason})
                    continue
            except Exception:
                counters["normalization_failures"] += 1
                logger.exception("Duplicate check failed at row %s | row=%s", row_index, row)
                logger.error(traceback.format_exc())
                skipped_rows_omitted = _append_skipped_sample(skipped_rows, skipped_rows_omitted, {"index": row_index, "raw": s, "normalized": row, "reason": "dup_check_exception"})
                raise

            try:
                _db_execute(
                    db,
                    _insert_ignore_sql(db, "timetable_slots", ["branch", "section", "semester", "day", "start_time", "end_time", "subject_name", "faculty_name", "is_lab", "room"]),
                    (row['branch'], row['section'], row['semester'], row['day'], row['start_time'], row['end_time'], row['subject_name'], row['faculty_name'], row['is_lab'], row['room']),
                )
                counters["inserted"] += 1
                inserted_since_commit += 1
                # commit in batches to reduce transaction memory
                if inserted_since_commit >= BATCH_INSERT_SIZE:
                    try:
                        db.commit()
                        batch_commit_counts += 1
                        # write a lightweight batch commit diagnostic line
                        try:
                            with open(preview_path, "a", encoding="utf-8") as pf:
                                pf.write(json.dumps({"type": "batch_commit", "batch_count": inserted_since_commit, "inserted_total": counters.get("inserted"), "timestamp": time.time()}, default=str) + "\n")
                        except Exception:
                            logger.exception("Failed to write batch commit diagnostic")
                    except Exception:
                        logger.exception("DB commit failed during batch commit in import_slots")
                    inserted_since_commit = 0
                # periodic progress log
                if row_index % progress_log_interval == 0:
                    logger.info("import_slots progress: processed=%d inserted=%d skipped=%d", row_index, counters.get("inserted"), counters.get("skipped_total"))
            except Exception:
                counters["failures"] += 1
                logger.exception("Import failed at row %s | row=%s", row_index, row)
                logger.error(traceback.format_exc())
                if preview_written < PREVIEW_ROW_CAP:
                    try:
                        with open(preview_path, "a", encoding="utf-8") as pf:
                            pf.write(json.dumps({"index": row_index, "raw": s, "normalized": row, "reason": "insert_exception"}, default=str) + "\n")
                        preview_written += 1
                    except Exception:
                        logger.exception("Failed to write insert_exception preview line")
                skipped_rows_omitted = _append_skipped_sample(skipped_rows, skipped_rows_omitted, {"index": row_index, "raw": s, "normalized": row, "reason": "insert_exception"})
                raise

        try:
            db.commit()
            if inserted_since_commit > 0:
                batch_commit_counts += 1
                try:
                    with open(preview_path, "a", encoding="utf-8") as pf:
                        pf.write(json.dumps({"type": "batch_commit", "batch_count": inserted_since_commit, "inserted_total": counters.get("inserted"), "timestamp": time.time()}, default=str) + "\n")
                except Exception:
                    logger.exception("Failed to write final batch commit diagnostic")
        except Exception:
            logger.exception("DB commit failed after import_slots")
            logger.error(traceback.format_exc())
            raise

    except Exception:
        logger.exception("import_slots failed")
        logger.error(traceback.format_exc())
        raise

    logger.info(
        "import_slots summary: %s",
        {k: counters.get(k) for k in ("total", "parsed", "inserted", "skipped_total", "skipped_invalid", "skipped_duplicate", "failures", "invalid_section", "normalization_failures")},
    )
    try:
        logger.info("import_slots branch subjects: %s", {key: sorted(values) for key, values in branch_subject_counts.items()})
    except Exception:
        logger.exception("Failed to log raw branch subject summary")
    # capture peak memory if tracemalloc was used
    peak = None
    try:
        if use_tracemalloc:
            current, peak = tracemalloc.get_traced_memory()
            tracemalloc.stop()
    except Exception:
        peak = None

    return {"counters": counters, "skipped_rows": skipped_rows, "skipped_rows_omitted": skipped_rows_omitted, "preview_path": preview_path if preview_written else None, "batch_commits": batch_commit_counts, "elapsed_seconds": time.time() - start_time, "memory_peak_bytes": peak}


def import_slots_normalized(db, slots: List[Dict]):
    """Insert slots into normalized `timetable_entries` table when possible.
    Best-effort: resolve branch -> branches.id, subject -> subjects.id, faculty -> teachers.id.
    Falls back to leaving subject_id/teacher_id NULL if resolution fails.
    Returns a dict: {"counters": {...}, "skipped_rows": [...]}
    """
    counters = {
        "total": len(slots),
        "parsed": 0,
        "inserted": 0,
        "skipped_total": 0,
        "skipped_branch": 0,
        "skipped_unresolved_subject": 0,
        "skipped_unresolved_teacher": 0,
        "skipped_invalid": 0,
        "skipped_duplicate": 0,
        "failures": 0,
        "missing_subjects": 0,
        "missing_teachers": 0,
        "invalid_section": 0,
        "normalization_failures": 0,
    }
    skipped_rows = []
    skipped_rows_omitted = 0
    preview_path = os.path.join(os.path.dirname(__file__), "uploads", "last_import_debug.jsonl")
    os.makedirs(os.path.dirname(preview_path), exist_ok=True)
    preview_written = 0
    batch_commit_counts = 0
    progress_log_interval = max(50, BATCH_INSERT_SIZE)
    start_time = time.time()
    use_tracemalloc = False
    if ENABLE_IMPORT_TRACEMALLOC:
        try:
            tracemalloc.start()
            use_tracemalloc = True
        except Exception:
            use_tracemalloc = False
    branch_cache = {}
    subject_cache = {}
    teacher_cache = {}
    subject_index = _build_subject_lookup_index(db)
    teacher_index = _build_teacher_lookup_index(db)
    seen_norm_keys = set()
    branch_subject_counts: Dict[str, set] = {}
    try:
        inserted_since_commit = 0
        for row_index, s in enumerate(slots, start=1):
            counters["parsed"] += 1
            bname = _clean_text(s.get("branch"))
            sec = _clean_text(s.get("section"))
            sem = _safe_int(s.get("semester"))
            day = _clean_text(s.get("day"))
            start = _clean_text(s.get("start_time"))
            end = _clean_text(s.get("end_time"))
            subj_name = _clean_text(s.get("subject_name"))
            fac_name = _clean_text(s.get("faculty_name"))
            is_lab = int(bool(s.get("is_lab")))
            room = _clean_text(s.get("room"))

            normalized_row = {
                "branch": bname,
                "section": sec,
                "semester": sem,
                "day": day,
                "start_time": start,
                "end_time": end,
                "subject_name": subj_name,
                "faculty_name": fac_name,
                "is_lab": is_lab,
                "room": room,
            }
            print("NORMALIZED_ROW", normalized_row)
            logger.info("Processing normalized row %s: raw=%s normalized=%s", row_index, s, normalized_row)
            if not _is_valid_academic_timetable_row({"branch": bname, "section": sec, "day": day, "start_time": start, "end_time": end, "subject_name": subj_name, "faculty_name": fac_name}):
                counters["skipped_total"] += 1
                counters["skipped_invalid"] += 1
                counters["invalid_section"] += 1
                reason = "invalid_row"
                logger.info("import_slots_normalized skipped row %s: reason=%s normalized=%s", row_index, reason, normalized_row)
                if preview_written < PREVIEW_ROW_CAP:
                    try:
                        with open(preview_path, "a", encoding="utf-8") as pf:
                            pf.write(json.dumps({"index": row_index, "raw": s, "normalized": normalized_row, "reason": reason}, default=str) + "\n")
                        preview_written += 1
                    except Exception:
                        logger.exception("Failed to write normalized skipped preview line")
                skipped_rows_omitted = _append_skipped_sample(skipped_rows, skipped_rows_omitted, {"index": row_index, "raw": s, "normalized": normalized_row, "reason": reason})
                continue

            branch_subject_key = f"{bname}|{sec}"
            branch_subject_counts.setdefault(branch_subject_key, set()).add(subj_name)

            branch_key = bname.strip().lower()
            branch_id = branch_cache.get(branch_key)
            if branch_id is None:
                branch_id = _resolve_or_create_branch_id(db, bname, branch_cache)
                branch_cache[branch_key] = branch_id
            if branch_id is None:
                counters["skipped_total"] += 1
                counters["skipped_branch"] += 1
                reason = "missing_branch"
                logger.info("import_slots_normalized skipped unresolved branch for row %s: %s", row_index, normalized_row)
                if preview_written < PREVIEW_ROW_CAP:
                    try:
                        with open(preview_path, "a", encoding="utf-8") as pf:
                            pf.write(json.dumps({"index": row_index, "raw": s, "normalized": normalized_row, "reason": reason}, default=str) + "\n")
                        preview_written += 1
                    except Exception:
                        logger.exception("Failed to write normalized branch preview line")
                skipped_rows_omitted = _append_skipped_sample(skipped_rows, skipped_rows_omitted, {"index": row_index, "raw": s, "normalized": normalized_row, "reason": reason})
                continue

            subject_id = _resolve_subject_id(subj_name, subject_cache, subject_index)
            if subj_name and subject_id is None:
                counters["skipped_unresolved_subject"] += 1
                counters["missing_subjects"] += 1

            teacher_id = _resolve_teacher_id(fac_name, teacher_cache, teacher_index)
            if fac_name and teacher_id is None:
                counters["skipped_unresolved_teacher"] += 1
                counters["missing_teachers"] += 1

            entry_row = {
                "branch_id": branch_id,
                "section": sec,
                "semester": sem,
                "day": day,
                "start_time": start,
                "end_time": end,
                "subject_id": subject_id,
                "teacher_id": teacher_id,
                "is_lab": is_lab,
                "room": room,
            }
            norm_key = _dup_key(
                entry_row["branch_id"],
                entry_row["section"],
                entry_row["semester"],
                entry_row["day"],
                entry_row["start_time"],
                entry_row["end_time"],
                entry_row["subject_id"] if entry_row["subject_id"] is not None else _normalize_subject_name(subj_name),
                entry_row["teacher_id"] if entry_row["teacher_id"] is not None else _normalize_teacher_name(fac_name),
                entry_row["room"],
            )
            if norm_key in seen_norm_keys:
                counters["skipped_total"] += 1
                counters["skipped_duplicate"] += 1
                reason = "in_memory_duplicate"
                logger.info("import_slots_normalized skipped duplicate row %s: %s", row_index, entry_row)
                if preview_written < PREVIEW_ROW_CAP:
                    try:
                        with open(preview_path, "a", encoding="utf-8") as pf:
                            pf.write(json.dumps({"index": row_index, "raw": s, "normalized": normalized_row, "reason": reason}, default=str) + "\n")
                        preview_written += 1
                    except Exception:
                        logger.exception("Failed to write normalized duplicate preview line")
                skipped_rows_omitted = _append_skipped_sample(skipped_rows, skipped_rows_omitted, {"index": row_index, "raw": s, "normalized": normalized_row, "reason": reason})
                continue

            duplicate_where = "COALESCE(CAST(branch_id AS TEXT), '') = COALESCE(CAST(%s AS TEXT), '') AND COALESCE(section, '') = COALESCE(%s, '') AND COALESCE(CAST(semester AS TEXT), '') = COALESCE(CAST(%s AS TEXT), '') AND COALESCE(day, '') = COALESCE(%s, '') AND COALESCE(start_time, '') = COALESCE(%s, '') AND COALESCE(end_time, '') = COALESCE(%s, '') AND COALESCE(CAST(subject_id AS TEXT), '') = COALESCE(CAST(%s AS TEXT), '') AND COALESCE(CAST(teacher_id AS TEXT), '') = COALESCE(CAST(%s AS TEXT), '') AND COALESCE(room, '') = COALESCE(%s, '')"
            try:
                if _row_exists(
                    db,
                    "timetable_entries",
                    duplicate_where,
                    (entry_row['branch_id'], entry_row['section'], entry_row['semester'], entry_row['day'], entry_row['start_time'], entry_row['end_time'], entry_row['subject_id'], entry_row['teacher_id'], entry_row['room']),
                ):
                    counters["skipped_total"] += 1
                    counters["skipped_duplicate"] += 1
                    reason = "db_duplicate"
                    logger.info("import_slots_normalized skipped duplicate row %s: %s", row_index, entry_row)
                    if preview_written < PREVIEW_ROW_CAP:
                        try:
                            with open(preview_path, "a", encoding="utf-8") as pf:
                                pf.write(json.dumps({"index": row_index, "raw": s, "normalized": normalized_row, "reason": reason}, default=str) + "\n")
                            preview_written += 1
                        except Exception:
                            logger.exception("Failed to write normalized duplicate preview line")
                    skipped_rows_omitted = _append_skipped_sample(skipped_rows, skipped_rows_omitted, {"index": row_index, "raw": s, "normalized": normalized_row, "reason": reason})
                    continue
            except Exception:
                counters["normalization_failures"] += 1
                logger.exception("Duplicate check failed at row %s | row=%s", row_index, normalized_row)
                logger.error(traceback.format_exc())
                skipped_rows_omitted = _append_skipped_sample(skipped_rows, skipped_rows_omitted, {"index": row_index, "raw": s, "normalized": normalized_row, "reason": "dup_check_exception"})
                raise

            seen_norm_keys.add(norm_key)
            try:
                logger.info(
                    "Normalized values | branch=%s subject=%s faculty=%s section=%s day=%s time=%s-%s room=%s",
                    bname,
                    subj_name,
                    fac_name,
                    sec,
                    day,
                    start,
                    end,
                    room,
                )
                _db_execute(
                    db,
                    _insert_ignore_sql(db, "timetable_entries", ["branch_id", "section", "semester", "day", "start_time", "end_time", "subject_id", "teacher_id", "is_lab", "room"]),
                    (entry_row['branch_id'], entry_row['section'], entry_row['semester'], entry_row['day'], entry_row['start_time'], entry_row['end_time'], entry_row['subject_id'], entry_row['teacher_id'], entry_row['is_lab'], entry_row['room']),
                )
                counters["inserted"] += 1
                branch_subject_counts.setdefault(branch_subject_key, set()).add(subj_name)
            except Exception:
                counters["failures"] += 1
                counters["normalization_failures"] += 1
                logger.exception("Import failed at normalized row %s | row=%s", row_index, normalized_row)
                logger.error(traceback.format_exc())
                if preview_written < PREVIEW_ROW_CAP:
                    try:
                        with open(preview_path, "a", encoding="utf-8") as pf:
                            pf.write(json.dumps({"index": row_index, "raw": s, "normalized": normalized_row, "reason": "insert_exception"}, default=str) + "\n")
                        preview_written += 1
                    except Exception:
                        logger.exception("Failed to write normalized insert_exception preview line")
                skipped_rows_omitted = _append_skipped_sample(skipped_rows, skipped_rows_omitted, {"index": row_index, "raw": s, "normalized": normalized_row, "reason": "insert_exception"})
                raise

            inserted_since_commit += 1
            if inserted_since_commit >= BATCH_INSERT_SIZE:
                try:
                    db.commit()
                    batch_commit_counts += 1
                    try:
                        with open(preview_path, "a", encoding="utf-8") as pf:
                            pf.write(json.dumps({"type": "batch_commit_normalized", "batch_count": inserted_since_commit, "inserted_total": counters.get("inserted"), "timestamp": time.time()}, default=str) + "\n")
                    except Exception:
                        logger.exception("Failed to write normalized batch commit diagnostic")
                except Exception:
                    logger.exception("DB commit failed during batch commit in import_slots_normalized")
                inserted_since_commit = 0
            if row_index % progress_log_interval == 0:
                logger.info("import_slots_normalized progress: processed=%d inserted=%d skipped=%d", row_index, counters.get("inserted"), counters.get("skipped_total"))

            if subject_id is None or teacher_id is None:
                logger.info(
                    "import_slots_normalized inserting with unresolved subject/teacher branch=%s subject=%s teacher=%s row_index=%s",
                    branch_id,
                    subject_id,
                    teacher_id,
                    row_index,
                )

            if subject_id is None and subj_name:
                logger.info("import_slots_normalized missing subject mapping row_index=%s subject=%s", row_index, subj_name)
            if teacher_id is None and fac_name:
                logger.info("import_slots_normalized missing teacher mapping row_index=%s faculty=%s", row_index, fac_name)

            try:
                _db_execute(
                    db,
                    _insert_ignore_sql(db, "timetable_entries", ["branch_id", "section", "semester", "day", "start_time", "end_time", "subject_id", "teacher_id", "is_lab", "room"]),
                    (entry_row['branch_id'], entry_row['section'], entry_row['semester'], entry_row['day'], entry_row['start_time'], entry_row['end_time'], entry_row['subject_id'], entry_row['teacher_id'], entry_row['is_lab'], entry_row['room']),
                )
                counters["inserted"] += 1
            except Exception:
                counters["failures"] += 1
                counters["normalization_failures"] += 1
                logger.exception("Import failed at normalized row %s | row=%s", row_index, normalized_row)
                logger.error(traceback.format_exc())
                skipped_rows_omitted = _append_skipped_sample(skipped_rows, skipped_rows_omitted, {"index": row_index, "raw": s, "normalized": normalized_row, "reason": "insert_exception"})
                raise

        try:
            db.commit()
            if inserted_since_commit > 0:
                batch_commit_counts += 1
                try:
                    with open(preview_path, "a", encoding="utf-8") as pf:
                        pf.write(json.dumps({"type": "batch_commit_normalized", "batch_count": inserted_since_commit, "inserted_total": counters.get("inserted"), "timestamp": time.time()}, default=str) + "\n")
                except Exception:
                    logger.exception("Failed to write final normalized batch commit diagnostic")
        except Exception:
            logger.exception("DB commit failed after import_slots_normalized")
            logger.error(traceback.format_exc())
            raise

    except Exception:
        logger.exception("import_slots_normalized failed")
        logger.error(traceback.format_exc())
        raise

    logger.info(
        "import_slots_normalized summary: %s",
        {k: counters.get(k) for k in ("total", "parsed", "inserted", "skipped_total", "skipped_branch", "skipped_unresolved_subject", "skipped_unresolved_teacher", "skipped_invalid", "skipped_duplicate", "failures", "missing_subjects", "missing_teachers", "invalid_section", "normalization_failures")},
    )
    try:
        logger.info("import_slots_normalized branch subjects: %s", {key: sorted(values) for key, values in branch_subject_counts.items()})
    except Exception:
        logger.exception("Failed to log normalized branch subject summary")
    peak = None
    try:
        if use_tracemalloc:
            current, peak = tracemalloc.get_traced_memory()
            tracemalloc.stop()
    except Exception:
        peak = None

    return {"counters": counters, "skipped_rows": skipped_rows, "skipped_rows_omitted": skipped_rows_omitted, "preview_path": preview_path if preview_written else None, "batch_commits": batch_commit_counts, "elapsed_seconds": time.time() - start_time, "memory_peak_bytes": peak}


# --- Lookup helpers --------------------------------------------------------

def get_current_slot(db, branch: str, section: str, now: Optional[datetime] = None):
    if now is None:
        now = datetime.now()
    weekday = _clean_text(now.strftime("%A"))
    cur_time = _clean_text(now.strftime("%H:%M"))
    branch = _clean_text(branch)
    section = _clean_text(section)
    logger.info("get_current_slot start branch=%s section=%s day=%s time=%s", branch, section, weekday, cur_time)
    # Prefer normalized timetable_entries when available (uses branch_id)
    try:
        branch_row = _db_execute(db, "SELECT id FROM branches WHERE LOWER(TRIM(name))=LOWER(TRIM(%s)) LIMIT 1", (branch,)).fetchone()
        branch_id = branch_row[0] if branch_row else None
    except Exception:
        branch_id = None
    logger.info("get_current_slot resolved branch_id=%s", branch_id)

    if branch_id is not None:
        rows = _db_execute(db,
            "SELECT te.*, s.name AS subject_name, t.name AS teacher_name FROM timetable_entries te LEFT JOIN subjects s ON te.subject_id = s.id LEFT JOIN teachers t ON te.teacher_id = t.id WHERE te.branch_id = %s AND LOWER(TRIM(COALESCE(te.section, ''))) = LOWER(TRIM(%s)) AND LOWER(TRIM(COALESCE(te.day, ''))) = LOWER(TRIM(%s)) AND te.start_time <= %s AND te.end_time >= %s ORDER BY te.start_time LIMIT 1",
            (branch_id, section or "", weekday, cur_time, cur_time),
        ).fetchall()
        logger.info("active timetable query result count=%d source=normalized branch=%s section=%s day=%s", len(rows), branch, section, weekday)
        if rows:
            logger.info("get_current_slot matched normalized timetable_entries row_id=%s subject=%s teacher=%s", rows[0]["id"] if hasattr(rows[0], "keys") and "id" in rows[0].keys() else None, rows[0]["subject_name"] if hasattr(rows[0], "keys") and "subject_name" in rows[0].keys() else None, rows[0]["teacher_name"] if hasattr(rows[0], "keys") and "teacher_name" in rows[0].keys() else None)
            return rows[0]

    # Fallback to legacy timetable_slots text-based lookup
    rows = _db_execute(db,
        "SELECT * FROM timetable_slots WHERE LOWER(TRIM(branch)) = LOWER(TRIM(%s)) AND LOWER(TRIM(COALESCE(section, ''))) = LOWER(TRIM(%s)) AND LOWER(TRIM(day)) = LOWER(TRIM(%s)) AND start_time <= %s AND end_time >= %s ORDER BY start_time LIMIT 1",
        (branch, section, weekday, cur_time, cur_time),
    ).fetchall()
    logger.info("active timetable query result count=%d source=legacy branch=%s section=%s day=%s", len(rows), branch, section, weekday)
    if rows:
        logger.info("get_current_slot matched legacy timetable_slots row subject=%s faculty=%s", rows[0]["subject_name"] if hasattr(rows[0], "keys") and "subject_name" in rows[0].keys() else None, rows[0]["faculty_name"] if hasattr(rows[0], "keys") and "faculty_name" in rows[0].keys() else None)
        return rows[0]
    logger.info("get_current_slot no active slot found for branch=%s section=%s day=%s time=%s", branch, section, weekday, cur_time)
    return None


def _ensure_column(db, table_name: str, column_name: str, column_definition: str):
    try:
        is_pg = _is_postgres_db(db)
        if is_pg:
            rows = _db_execute(
                db,
                "SELECT column_name FROM information_schema.columns WHERE table_schema = 'public' AND table_name = %s",
                (table_name,),
            ).fetchall()
            cols = {r[0] if not hasattr(r, "keys") else r["column_name"] for r in rows}
        else:
            rows = _db_execute(db, f"PRAGMA table_info({table_name})").fetchall()
            cols = {r[1] if not hasattr(r, "keys") else r["name"] for r in rows}
        if cols and column_name not in cols:
            _db_execute(db, f"ALTER TABLE {table_name} ADD COLUMN {column_name} {column_definition}")
            try:
                db.commit()
            except Exception:
                pass
    except Exception as e:
        try:
            db.rollback()
        except Exception:
            pass
        logger.warning(f"Failed to ensure column {table_name}.{column_name}: {e}")


def derive_subject_code(s_name: str) -> str:
    s_clean = _clean_text(s_name)
    if not s_clean:
        return ""
    tokens = [t for t in re.split(r"[^A-Za-z0-9]+", s_clean) if t and t.lower() not in ("and", "of", "the", "lab", "practical")]
    if len(tokens) >= 2:
        return "".join([t[0].upper() for t in tokens])
    elif tokens:
        return tokens[0].upper()[:6]
    return s_clean.upper()[:6]


def derive_teacher_username(t_name: str) -> str:
    clean_t = _clean_text(t_name)
    clean_base = re.sub(r"^(mr|ms|mrs|dr|prof)\.?\s+", "", clean_t, flags=re.I)
    sanitized = re.sub(r"[^a-zA-Z0-9]+", "_", clean_base).strip("_").lower()
    if not sanitized:
        sanitized = "teacher"
    return sanitized


def auto_setup_academic_from_slots(db, slots: List[Dict]) -> Dict[str, Any]:
    """Single-Source Academic Setup from Timetable Slots:
    1. Ensures Branch records exist in `branches`.
    2. Ensures Section records exist in `sections`.
    3. Automatically creates missing Subject records in `subjects` with codes and branch IDs.
    4. Automatically creates missing Teacher user accounts in `teachers` with default passwords ("1234").
    5. Automatically maps `Teacher -> Subject -> Branch -> Section -> Semester` in `teacher_subjects`, `teacher_branches`, and `teacher_subject_assignments`.
    6. Does NOT create demo branches/sections/teachers/subjects.
    7. Prints a summary of created entities.
    """
    is_pg = _is_postgres_db(db)

    _ensure_column(db, "subjects", "code", "TEXT")
    _ensure_column(db, "subjects", "branch_id", "INTEGER")
    _ensure_column(db, "teachers", "username", "TEXT")
    _ensure_column(db, "teachers", "password", "TEXT")
    _ensure_column(db, "teachers", "password_hash", "TEXT")
    _ensure_column(db, "teachers", "status", "TEXT")
    _ensure_column(db, "teachers", "branch_id", "INTEGER")
    _ensure_column(db, "teachers", "subject_id", "INTEGER")

    try:
        if is_pg:
            _db_execute(db, "CREATE TABLE IF NOT EXISTS sections (id SERIAL PRIMARY KEY, branch_id INTEGER, name TEXT)")
        else:
            _db_execute(db, "CREATE TABLE IF NOT EXISTS sections (id INTEGER PRIMARY KEY AUTOINCREMENT, branch_id INTEGER, name TEXT)")
        try:
            db.commit()
        except Exception:
            pass
    except Exception:
        pass

    summary = {
        "branches_created": 0,
        "sections_created": 0,
        "subjects_created": 0,
        "subjects_mapped": 0,
        "teachers_created": 0,
        "teachers_mapped": 0,
        "assignments_created": 0,
    }

    created_branches_list: List[str] = []
    created_sections_list: List[str] = []
    created_subjects_list: List[str] = []
    created_teachers_list: List[str] = []
    created_assignments_list: List[str] = []

    branch_cache: Dict[str, int] = {}
    section_cache: Dict[tuple, int] = {}
    subject_cache: Dict[tuple, int] = {}
    teacher_cache: Dict[str, int] = {}

    def get_or_create_branch(b_name: str) -> Optional[int]:
        clean_b = _clean_text(b_name)
        if not clean_b:
            return None
        norm_b = clean_b.upper()
        if norm_b in branch_cache:
            return branch_cache[norm_b]

        row = _db_execute(db, "SELECT id FROM branches WHERE UPPER(TRIM(name)) = UPPER(TRIM(%s)) LIMIT 1", (clean_b,)).fetchone()
        if row:
            bid = row[0] if not hasattr(row, "keys") else row["id"]
            branch_cache[norm_b] = int(bid)
            return int(bid)

        try:
            _db_execute(db, "INSERT INTO branches (name) VALUES (%s)", (clean_b,))
            try:
                db.commit()
            except Exception:
                pass
            r2 = _db_execute(db, "SELECT id FROM branches WHERE UPPER(TRIM(name)) = UPPER(TRIM(%s)) LIMIT 1", (clean_b,)).fetchone()
            if r2:
                bid = r2[0] if not hasattr(r2, "keys") else r2["id"]
                branch_cache[norm_b] = int(bid)
                summary["branches_created"] += 1
                created_branches_list.append(clean_b)
                return int(bid)
        except Exception as e:
            logger.exception("Failed to auto-create branch %s: %s", clean_b, e)
            try: db.rollback()
            except Exception: pass
        return None

    def get_or_create_section(sec_name: str, b_id: Optional[int]) -> Optional[int]:
        clean_sec = _clean_text(sec_name)
        if not clean_sec:
            return None
        norm_sec = clean_sec.upper()
        sec_key = (norm_sec, b_id or 0)
        if sec_key in section_cache:
            return section_cache[sec_key]

        row = _db_execute(db, "SELECT id FROM sections WHERE UPPER(TRIM(name)) = UPPER(TRIM(%s)) AND (branch_id = %s OR branch_id IS NULL) LIMIT 1", (clean_sec, b_id)).fetchone()
        if row:
            sid = row[0] if not hasattr(row, "keys") else row["id"]
            section_cache[sec_key] = int(sid)
            return int(sid)

        try:
            _db_execute(db, "INSERT INTO sections (name, branch_id) VALUES (%s, %s)", (clean_sec, b_id))
            try:
                db.commit()
            except Exception:
                pass
            r2 = _db_execute(db, "SELECT id FROM sections WHERE UPPER(TRIM(name)) = UPPER(TRIM(%s)) AND (branch_id = %s OR branch_id IS NULL) ORDER BY id DESC LIMIT 1", (clean_sec, b_id)).fetchone()
            if r2:
                sid = r2[0] if not hasattr(r2, "keys") else r2["id"]
                section_cache[sec_key] = int(sid)
                summary["sections_created"] += 1
                created_sections_list.append(clean_sec)
                return int(sid)
        except Exception as e:
            logger.exception("Failed auto-creating section %s: %s", clean_sec, e)
            try: db.rollback()
            except Exception: pass
        return None

    def get_or_create_subject(s_name: str, b_id: Optional[int], s_code: Optional[str] = None) -> Optional[int]:
        clean_s = _clean_text(s_name)
        if not clean_s:
            return None
        key = (clean_s.lower(), b_id or 0)
        if key in subject_cache:
            return subject_cache[key]

        code_val = s_code or derive_subject_code(clean_s)
        try:
            existing = _db_execute(db, "SELECT id, name, code FROM subjects").fetchall()
        except Exception:
            existing = []

        for r in existing:
            sid = r[0] if not hasattr(r, "keys") else r["id"]
            sname = r[1] if not hasattr(r, "keys") else r["name"]
            scode = r[2] if not hasattr(r, "keys") and len(r) > 2 else (r["code"] if "code" in r.keys() else None)
            if _normalize_subject_name(clean_s) == _normalize_subject_name(sname) or (scode and scode.upper() == code_val.upper()):
                subject_cache[key] = int(sid)
                summary["subjects_mapped"] += 1
                if not scode and code_val:
                    try:
                        _db_execute(db, "UPDATE subjects SET code = %s WHERE id = %s", (code_val, int(sid)))
                        db.commit()
                    except Exception:
                        pass
                return int(sid)

        try:
            display_name = clean_s.title()
            _db_execute(db, "INSERT INTO subjects (name, code, branch_id) VALUES (%s, %s, %s)", (display_name, code_val, b_id))
            try:
                db.commit()
            except Exception:
                pass
            r2 = _db_execute(db, "SELECT id FROM subjects WHERE LOWER(TRIM(name)) = LOWER(TRIM(%s)) ORDER BY id DESC LIMIT 1", (display_name,)).fetchone()
            if r2:
                sid = r2[0] if not hasattr(r2, "keys") else r2["id"]
                subject_cache[key] = int(sid)
                summary["subjects_created"] += 1
                created_subjects_list.append(f"{display_name} ({code_val})")
                return int(sid)
        except Exception as e:
            logger.exception("Failed auto-creating subject %s: %s", clean_s, e)
            try: db.rollback()
            except Exception: pass
        return None

    def get_or_create_teacher(t_name: str, default_branch_id: Optional[int] = None, default_subject_id: Optional[int] = None) -> Optional[int]:
        clean_t = _clean_text(t_name)
        if not clean_t:
            return None
        norm_t = clean_t.lower()
        if norm_t in teacher_cache:
            return teacher_cache[norm_t]

        try:
            existing = _db_execute(db, "SELECT id, name, username FROM teachers").fetchall()
        except Exception:
            existing = []

        for r in existing:
            tid = r[0] if not hasattr(r, "keys") else r["id"]
            tname = r[1] if not hasattr(r, "keys") else r["name"]
            if _normalize_teacher_name(clean_t) == _normalize_teacher_name(tname) or clean_t.lower() == str(tname).lower():
                teacher_cache[norm_t] = int(tid)
                summary["teachers_mapped"] += 1
                return int(tid)

        try:
            uname_base = derive_teacher_username(clean_t)
            uname_final = uname_base
            idx = 1
            while True:
                r_chk = _db_execute(db, "SELECT id FROM teachers WHERE username = %s LIMIT 1", (uname_final,)).fetchone()
                if not r_chk:
                    break
                uname_final = f"{uname_base}_{idx}"
                idx += 1

            pwd_hash = generate_password_hash("1234")
            display_name = clean_t
            _db_execute(
                db,
                "INSERT INTO teachers (name, username, password, password_hash, status, branch_id, subject_id) VALUES (%s, %s, %s, %s, %s, %s, %s)",
                (display_name, uname_final, "1234", pwd_hash, "active", default_branch_id, default_subject_id),
            )
            try:
                db.commit()
            except Exception:
                pass
            r2 = _db_execute(db, "SELECT id FROM teachers WHERE username = %s LIMIT 1", (uname_final,)).fetchone()
            if r2:
                tid = r2[0] if not hasattr(r2, "keys") else r2["id"]
                teacher_cache[norm_t] = int(tid)
                summary["teachers_created"] += 1
                created_teachers_list.append(display_name)
                return int(tid)
        except Exception as e:
            logger.exception("Failed auto-creating teacher %s: %s", clean_t, e)
            try: db.rollback()
            except Exception: pass
        return None

    def add_teacher_assignments(t_id: int, s_id: Optional[int], b_id: Optional[int], sec_str: str, sem_str: str, f_name: str, s_name: str):
        if not t_id:
            return
        if s_id:
            try:
                if is_pg:
                    _db_execute(db, "INSERT INTO teacher_subjects (teacher_id, subject_id) VALUES (%s, %s) ON CONFLICT DO NOTHING", (t_id, s_id))
                else:
                    _db_execute(db, "INSERT OR IGNORE INTO teacher_subjects (teacher_id, subject_id) VALUES (%s, %s)", (t_id, s_id))
            except Exception:
                pass
        if b_id:
            try:
                if is_pg:
                    _db_execute(db, "INSERT INTO teacher_branches (teacher_id, branch_id) VALUES (%s, %s) ON CONFLICT DO NOTHING", (t_id, b_id))
                else:
                    _db_execute(db, "INSERT OR IGNORE INTO teacher_branches (teacher_id, branch_id) VALUES (%s, %s)", (t_id, b_id))
            except Exception:
                pass

        if s_id and b_id:
            sec_clean = _clean_text(sec_str)
            sem_clean = str(sem_str or "").strip()
            if sem_clean and not sem_clean.startswith("Semester"):
                sem_clean = f"Semester {sem_clean}"
            try:
                dup = _db_execute(
                    db,
                    "SELECT id FROM teacher_subject_assignments WHERE teacher_id = %s AND subject_id = %s AND branch_id = %s AND COALESCE(section, '') = COALESCE(%s, '') AND COALESCE(semester, '') = COALESCE(%s, '') LIMIT 1",
                    (t_id, s_id, b_id, sec_clean, sem_clean),
                ).fetchone()
                if not dup:
                    _db_execute(
                        db,
                        "INSERT INTO teacher_subject_assignments (teacher_id, subject_id, branch_id, section, semester, academic_year) VALUES (%s, %s, %s, %s, %s, %s)",
                        (t_id, s_id, b_id, sec_clean, sem_clean, "2025-2026"),
                    )
                    summary["assignments_created"] += 1
                    created_assignments_list.append(f"{f_name} -> {s_name} ({sec_clean})")
            except Exception as e:
                logger.exception("Failed creating teacher_subject_assignment: %s", e)

    for slot in slots:
        b_name = slot.get("branch")
        s_name = slot.get("subject_name")
        f_name = slot.get("faculty_name")
        sec_name = slot.get("section")
        sem_val = slot.get("semester")

        if not b_name and not s_name:
            continue

        b_id = get_or_create_branch(b_name)
        get_or_create_section(sec_name, b_id)

        s_code = slot.get("sub_code") or slot.get("subject_code")
        s_id = get_or_create_subject(s_name, b_id, s_code)
        t_id = get_or_create_teacher(f_name, default_branch_id=b_id, default_subject_id=s_id)

        if t_id:
            add_teacher_assignments(t_id, s_id, b_id, sec_name, sem_val, f_name, s_name)

        slot["branch_id"] = b_id
        slot["subject_id"] = s_id
        slot["teacher_id"] = t_id

    try:
        db.commit()
    except Exception as e:
        logger.exception("Final DB commit after auto academic setup failed: %s", e)

    # Print requested summary format
    print("Branches Created:")
    if created_branches_list:
        for b in sorted(set(created_branches_list)):
            print(f"* {b}")
    else:
        print("(None created - reused existing)")

    print("\nSections Created:")
    if created_sections_list:
        for s in sorted(set(created_sections_list)):
            print(f"* {s}")
    else:
        print("(None created - reused existing)")

    print("\nSubjects Created:")
    if created_subjects_list:
        for sb in sorted(set(created_subjects_list)):
            print(f"* {sb}")
    else:
        print("(None created - reused existing)")

    print("\nTeachers Created:")
    if created_teachers_list:
        for t in sorted(set(created_teachers_list)):
            print(f"* {t}")
    else:
        print("(None created - reused existing)")

    print("\nAssignments Created:")
    if created_assignments_list:
        for a in sorted(set(created_assignments_list)):
            print(f"* {a}")
    else:
        print("(None created - reused existing)")

    summary["branches_list"] = sorted(set(created_branches_list))
    summary["sections_list"] = sorted(set(created_sections_list))
    summary["subjects_list"] = sorted(set(created_subjects_list))
    summary["teachers_list"] = sorted(set(created_teachers_list))
    summary["assignments_list"] = sorted(set(created_assignments_list))

    return summary


# --- Routes registration ---------------------------------------------------

def register_routes(app, db_getter=None):
    globals()["get_db"] = db_getter

    # Idempotency guard: avoid Flask endpoint/path collisions if this
    # registration function is called more than once in a process.
    if getattr(app, "_timetable_routes_registered", False):
        return
    existing = {
        "timetable_home",
        "timetable_manage",
        "timetable_faculty_schedules",
        "timetable_admin_bulk_resolve",
    }
    if any(ep in app.view_functions for ep in existing):
        logger.warning("Skipping timetable route registration because endpoints already exist")
        app._timetable_routes_registered = True
        return
    app._timetable_routes_registered = True

    @app.route("/timetable")
    def timetable_home():
        db = None
        rows_count = 0
        normalized_count = 0
        table_ready = False
        upcoming_classes = []
        active_slot = None
        try:
            if db_getter is None:
                raise RuntimeError("Database getter is not configured")
            db = db_getter()
            ensure_timetable_tables(db)
            row = _db_execute(db, "SELECT COUNT(1) AS c FROM timetable_slots").fetchone()
            rows_count = int(row["c"] if row and row["c"] is not None else 0)
            nrow = _db_execute(db, "SELECT COUNT(1) AS c FROM timetable_entries").fetchone()
            normalized_count = int(nrow["c"] if nrow and nrow["c"] is not None else 0)
            table_ready = True
            scoped_branch = session.get("teacher_branch_name") or session.get("teacher_branch") or ""
            scoped_section = session.get("teacher_section") or ""
            if not scoped_branch and not scoped_section:
                scoped_branch, scoped_section = _infer_dashboard_scope(db)
            active_slot = get_current_active_class(db, scoped_branch, scoped_section)
            upcoming_classes = get_upcoming_classes(db, scoped_branch, scoped_section, limit=4)
            logger.info(
                "timetable_home diagnostics scope_branch=%s scope_section=%s slots=%d entries=%d active=%s upcoming=%d",
                scoped_branch,
                scoped_section,
                rows_count,
                normalized_count,
                bool(active_slot),
                len(upcoming_classes),
            )
        except Exception as e:
            logger.exception("Failed to render timetable status page")
            return render_template_string(
                """
                {% extends "layout.html" %}
                {% block content %}
                <div class="card">
                    <h1>Timetable Status</h1>
                    <p>Timetable routes are loaded, but the database check failed.</p>
                    <p>{{ error }}</p>
                    <p><a class="button" href="{{ url_for('timetable_manage') }}">Open timetable management</a></p>
                </div>
                {% endblock %}
                """,
                error=str(e),
            )
        finally:
            if db:
                try:
                    db.close()
                except Exception:
                    pass

        return render_template(
            "timetable_dashboard.html",
            table_ready=table_ready,
            rows_count=rows_count,
            normalized_count=normalized_count,
            active_slot=active_slot,
            upcoming_classes=upcoming_classes,
        )

    @app.route("/timetable/dashboard")
    def timetable_dashboard_alias():
        return timetable_home()

    @app.route("/timetable/upload")
    def timetable_upload():
        return redirect(url_for("timetable_manage"))
    @app.route("/timetable/manage", methods=("GET", "POST"))
    def timetable_manage():
        if session.get("role") != "admin":
            return redirect(url_for("dashboard"))
        db = get_db()
        ensure_timetable_tables(db)
        staged_path = os.path.join(os.path.dirname(__file__), "uploads", "staged_timetable.json")

        if request.method == "POST":
            action = request.form.get("action")

            if action == "delete_timetable":
                try:
                    db = get_db()
                    _db_execute(db, "DELETE FROM timetable_slots")
                    _db_execute(db, "DELETE FROM timetable_entries")
                    try:
                        db.commit()
                    except Exception:
                        pass
                    if os.path.exists(staged_path):
                        try: os.remove(staged_path)
                        except Exception: pass
                    flash("Timetable deleted successfully.", "success")
                except Exception:
                    logger.exception("Failed to delete timetable data")
                    flash("Failed to delete timetable.", "error")
                return redirect(url_for("timetable_manage"))

            elif action == "cancel_import":
                if os.path.exists(staged_path):
                    try:
                        os.remove(staged_path)
                    except Exception:
                        pass
                flash("Timetable import cancelled.", "info")
                return redirect(url_for("timetable_manage"))

            elif action == "confirm_import":
                if not os.path.exists(staged_path):
                    flash("No pending timetable upload found to confirm. Please upload a file first.", "error")
                    return redirect(url_for("timetable_manage"))

                try:
                    with open(staged_path, "r", encoding="utf-8") as f:
                        staged_data = json.load(f)
                    slots = staged_data.get("slots", [])
                    filename = staged_data.get("filename", "upload")

                    if not slots:
                        flash("Staged timetable has no valid slots.", "error")
                        return redirect(url_for("timetable_manage"))

                    # Academic Setup Validation Gate (Block Import if critical errors exist)
                    if validate_staged_slots:
                        val_report = validate_staged_slots(slots, db)
                        if not val_report.get("can_import", True):
                            critical_errs = val_report.get("block_import", {}).get("critical_errors", [])
                            flash(f"Timetable import blocked due to {len(critical_errs)} critical validation error(s). Please review and correct the errors.", "error")
                            for err in critical_errs[:5]:
                                flash(f"• {err}", "error")
                            return redirect(url_for("timetable_manage"))

                    # 1. Automatic Academic Setup (branches, subjects, teachers, assignments)
                    summary = auto_setup_academic_from_slots(db, slots)

                    # 2. Import timetable entries into DB
                    import_info = import_slots_streaming(db, slots)

                    # 3. Publish Announcement
                    try:
                        _db_execute(db, "INSERT INTO sys_announcements (title, content, target_audience, created_by) VALUES (?, ?, ?, ?)",
                                    ("📅 Timetable Updated", f"A new timetable ({filename}) has been published. Please check your dashboard for the updated schedule.", "all", "Admin"))
                        db.commit()
                    except Exception as e:
                        print(f"[Timetable Announcement ERROR] {repr(e)}")

                    # Clean staged file
                    try:
                        os.remove(staged_path)
                    except Exception:
                        pass

                    db.commit()

                    success_msg = (
                        f"Timetable import & academic setup completed! Created {summary['subjects_created']} subjects, "
                        f"mapped {summary['subjects_mapped']} subjects, created {summary['teachers_created']} teacher accounts, "
                        f"and created {summary['assignments_created']} class assignments."
                    )
                    session["timetable_manage_banner"] = success_msg
                    session["timetable_last_imported_file"] = filename
                    session["timetable_refresh_normalized"] = "1"
                    flash(success_msg, "success")
                except Exception as e:
                    logger.exception("Confirm import failed: %s", e)
                    try:
                        db.rollback()
                    except Exception:
                        pass
                    flash(f"Failed to confirm timetable import: {e}", "error")
                return redirect(url_for("timetable_manage"))

            file = request.files.get("timetable_file")
            if not file:
                flash("Please upload a file.", "error")
                return redirect(url_for("timetable_manage"))

            filename = file.filename or "upload"
            safe_path = os.path.join(os.path.dirname(__file__), "uploads")
            os.makedirs(safe_path, exist_ok=True)
            dest = os.path.join(safe_path, filename)

            try:
                file.save(dest)
                if not os.path.exists(dest) or os.path.getsize(dest) == 0:
                    flash("Error: Uploaded file is empty or corrupted.", "error")
                    return redirect(url_for("timetable_manage"))
            except Exception as e:
                logger.exception("File save failed: %s", e)
                flash("Error: Unable to save uploaded file.", "error")
                return redirect(url_for("timetable_manage"))

            ext = os.path.splitext(filename)[1].lower()
            try:
                slots_iter = None
                if ext in (".docx",) and docx is not None:
                    summary = scan_docx_structure(dest, max_tables=None)
                    direct_tables = int(summary.get("direct_timetable_tables", 0) or 0)
                    legacy_tables = int(summary.get("timetable_tables", 0) or 0)
                    if direct_tables == 0 and legacy_tables == 0:
                        flash("No timetable tables detected in DOCX.", "error")
                        return redirect(url_for("timetable_manage"))
                    slots_iter = iter_docx_section_slots(dest, single_section_only=False, max_tables=None)
                elif ext in (".pdf",) and pdfplumber is not None:
                    pdf_stats = {}
                    slots_iter = parse_pdf_to_slots(dest, stats=pdf_stats)
                else:
                    flash("Unsupported file format or missing parser dependencies.", "error")
                    return redirect(url_for("timetable_manage"))

                if slots_iter is None:
                    flash("Error: Unable to parse file contents.", "error")
                    return redirect(url_for("timetable_manage"))

                slots = list(slots_iter)
                if not slots:
                    flash("No timetable slots were parsed from the file.", "error")
                    return redirect(url_for("timetable_manage"))

                # Save staged slots for preview before import
                with open(staged_path, "w", encoding="utf-8") as f:
                    json.dump({"filename": filename, "slots": slots}, f, indent=2, default=str)

                flash("Timetable file parsed successfully! Review the extracted academic setup below and click 'Confirm Import' to apply.", "info")
                return redirect(url_for("timetable_manage"))
            except Exception as e:
                logger.exception("Failed parsing timetable file: %s", e)
                flash(f"Failed to parse file: {e}", "error")
                return redirect(url_for("timetable_manage"))

        # GET: show management UI & pending preview if exists
        success_banner = session.pop("timetable_manage_banner", None)
        last_imported_file = session.get("timetable_last_imported_file")
        if session.pop("timetable_refresh_normalized", None):
            logger.info("Refreshing normalized timetable rows for manage view")
            refresh_stats = _refresh_timetable_entry_ids(db)

        # Check pending preview
        preview_mode = False
        preview_rows = []
        unique_subjects = []
        unique_teachers = []
        unique_sections = []
        validation_report = None
        if os.path.exists(staged_path):
            try:
                with open(staged_path, "r", encoding="utf-8") as f:
                    staged_data = json.load(f)
                staged_slots = staged_data.get("slots", [])
                if staged_slots:
                    preview_mode = True
                    subj_map = {}
                    teach_map = {}
                    sec_map = {}
                    for s in staged_slots:
                        sub = _clean_text(s.get("subject_name"))
                        fac = _clean_text(s.get("faculty_name"))
                        sec = _clean_text(s.get("section"))
                        sem = str(s.get("semester") or "").strip()
                        br = _clean_text(s.get("branch"))
                        day = _clean_text(s.get("day"))
                        t_start = _clean_text(s.get("start_time"))
                        t_end = _clean_text(s.get("end_time"))

                        if sub and sub not in subj_map:
                            subj_map[sub] = {"name": sub, "branch": br, "code": derive_subject_code(sub)}
                        if fac and fac not in teach_map:
                            teach_map[fac] = {"name": fac, "username": derive_teacher_username(fac)}
                        sec_key = f"{br}|{sec}|{sem}"
                        if sec_key not in sec_map:
                            sec_map[sec_key] = {"branch": br, "section": sec, "semester": sem}

                        preview_rows.append({
                            "subject_name": sub,
                            "faculty_name": fac,
                            "section": sec,
                            "semester": sem,
                            "branch": br,
                            "day": day,
                            "start_time": t_start,
                            "end_time": t_end,
                        })
                    unique_subjects = list(subj_map.values())
                    unique_teachers = list(teach_map.values())
                    unique_sections = list(sec_map.values())
                    if validate_staged_slots:
                        validation_report = validate_staged_slots(staged_slots, db)
            except Exception as e:
                logger.exception("Failed to load staged preview: %s", e)
        rows = []
        rows_source = "raw"
        try:
            c1 = _db_execute(db, "SELECT COUNT(*) AS c FROM timetable_entries").fetchone()
            c_entries = int(c1[0] if c1 is not None else 0)
        except Exception:
            c_entries = 0
        try:
            c2 = _db_execute(db, "SELECT COUNT(*) AS c FROM timetable_slots").fetchone()
            c_slots = int(c2[0] if c2 is not None else 0)
        except Exception:
            c_slots = 0

        # Server-side pagination & search
        page = 1
        try:
            page = int(request.args.get('page') or request.form.get('page') or 1)
        except Exception:
            page = 1
        PAGE_SIZE = 25
        q = (request.args.get('q') or request.form.get('q') or '').strip()

        # Build WHERE clause for simple text search across normalized and legacy timetable fields.
        # Use LOWER(TRIM(COALESCE(...))) for normalization and COALESCE to avoid NULLs.
        where_clauses = []
        params = []
        if q:
            raw_q = q
            norm_q = raw_q.strip()
            like = f"%{norm_q}%"
            # normalized SQL comparisons (case-insensitive, trimmed)
            where_clauses.append(
                "(LOWER(TRIM(COALESCE(s.name,''))) LIKE LOWER(TRIM(%s))"
                " OR LOWER(TRIM(COALESCE(t.name,''))) LIKE LOWER(TRIM(%s))"
                " OR LOWER(TRIM(COALESCE(te.subject_name,''))) LIKE LOWER(TRIM(%s))"
                " OR LOWER(TRIM(COALESCE(te.faculty_name,''))) LIKE LOWER(TRIM(%s))"
                " OR LOWER(TRIM(COALESCE(ts.subject_name,''))) LIKE LOWER(TRIM(%s))"
                " OR LOWER(TRIM(COALESCE(ts.faculty_name,''))) LIKE LOWER(TRIM(%s))"
                " OR LOWER(TRIM(COALESCE(te.section,''))) LIKE LOWER(TRIM(%s))"
                " OR LOWER(TRIM(COALESCE(b.name,''))) LIKE LOWER(TRIM(%s))"
                " OR LOWER(TRIM(COALESCE(ts.branch,''))) LIKE LOWER(TRIM(%s))"
                " OR LOWER(TRIM(COALESCE(te.room,''))) LIKE LOWER(TRIM(%s))"
                " OR LOWER(TRIM(COALESCE(te.day,''))) LIKE LOWER(TRIM(%s))"
                " OR LOWER(TRIM(COALESCE(CAST(COALESCE(te.semester, ts.semester) AS TEXT), ''))) LIKE LOWER(TRIM(%s)) )"
            )
            # Pass the same like param for each placeholder; TRIM happens in SQL
            params.extend([like] * 10)
            logger.info("Timetable manage search q=%s normalized=%s", raw_q, norm_q)

        where_sql = "WHERE " + " AND ".join(where_clauses) if where_clauses else ""

        try:
            # total count of normalized entries
            count_sql = f"SELECT COUNT(1) AS c FROM timetable_entries te LEFT JOIN subjects s ON te.subject_id = s.id LEFT JOIN teachers t ON te.teacher_id = t.id LEFT JOIN branches b ON te.branch_id = b.id LEFT JOIN timetable_slots ts ON LOWER(TRIM(ts.branch)) = LOWER(TRIM(b.name)) AND COALESCE(ts.section, '') = COALESCE(te.section, '') AND COALESCE(CAST(ts.semester AS TEXT), '') = COALESCE(CAST(te.semester AS TEXT), '') AND COALESCE(ts.day, '') = COALESCE(te.day, '') AND COALESCE(ts.start_time, '') = COALESCE(te.start_time, '') AND COALESCE(ts.end_time, '') = COALESCE(te.end_time, '') AND COALESCE(ts.room, '') = COALESCE(te.room, '') {where_sql}"
            logger.debug("Timetable manage count_sql=%s params=%s", count_sql, params)
            total_row = _db_execute(db, count_sql, tuple(params)).fetchone()
            total_count = int(total_row[0] if total_row is not None else 0)
        except Exception:
            logger.exception("Failed to count timetable_entries for pagination")
            total_count = 0

        entries = []
        rows_source = 'raw'
        try:
            if total_count > 0:
                offset = max(0, (page - 1) * PAGE_SIZE)
                visible_clause = "AND NOT (\n                        COALESCE(te.day, '') = ''\n                        AND COALESCE(te.start_time, '') = ''\n                        AND COALESCE(te.end_time, '') = ''\n                        AND COALESCE(te.section, '') = ''\n                        AND COALESCE(CAST(COALESCE(te.semester, ts.semester) AS TEXT), '') = ''\n                        AND COALESCE(te.room, '') = ''\n                        AND COALESCE(s.name, te.subject_name, ts.subject_name, '') = ''\n                        AND COALESCE(t.name, te.faculty_name, ts.faculty_name, '') = ''\n                        AND COALESCE(b.name, ts.branch, '') = ''\n                      )"
                if not where_sql:
                    visible_clause = visible_clause.replace("AND NOT", "WHERE NOT", 1)
                sql = f"""
                    SELECT
                        te.day,
                        te.start_time,
                        te.end_time,
                        te.section,
                        COALESCE(te.semester, ts.semester) AS semester,
                        te.room,
                        te.is_lab,
                        COALESCE(s.name, te.subject_name, ts.subject_name, '') AS subject_name,
                        COALESCE(t.name, te.faculty_name, ts.faculty_name, '') AS faculty_name,
                        COALESCE(b.name, ts.branch, '') AS branch_name
                    FROM timetable_entries te
                    LEFT JOIN subjects s ON te.subject_id = s.id
                    LEFT JOIN teachers t ON te.teacher_id = t.id
                    LEFT JOIN branches b ON te.branch_id = b.id
                    LEFT JOIN timetable_slots ts
                        ON LOWER(TRIM(ts.branch)) = LOWER(TRIM(b.name))
                       AND COALESCE(ts.section, '') = COALESCE(te.section, '')
                       AND COALESCE(CAST(ts.semester AS TEXT), '') = COALESCE(CAST(te.semester AS TEXT), '')
                       AND COALESCE(ts.day, '') = COALESCE(te.day, '')
                       AND COALESCE(ts.start_time, '') = COALESCE(te.start_time, '')
                       AND COALESCE(ts.end_time, '') = COALESCE(te.end_time, '')
                       AND COALESCE(ts.room, '') = COALESCE(te.room, '')
                    {where_sql}
                                        {visible_clause}
                    ORDER BY te.day, te.start_time
                    LIMIT %s OFFSET %s
                """
                qparams = list(params) + [PAGE_SIZE, offset]
                entries = _db_execute(db, sql, tuple(qparams)).fetchall()
                logger.debug("Timetable manage entries_sql=%s params=%s", sql, qparams)
                entries = [dict(r) for r in entries]
                logger.info("Timetable manage raw entries loaded=%s", len(entries))
                rows_source = 'normalized'
            else:
                # Fallback to legacy slots with paging
                count_sql = "SELECT COUNT(1) AS c FROM timetable_slots"
                total_row = _db_execute(db, count_sql).fetchone()
                total_count = int(total_row[0] if total_row is not None else 0)
                offset = max(0, (page - 1) * PAGE_SIZE)
                sql = "SELECT * FROM timetable_slots ORDER BY day, start_time LIMIT %s OFFSET %s"
                entries = _db_execute(db, sql, (PAGE_SIZE, offset)).fetchall()
                entries = [dict(r) for r in entries]
                rows_source = 'raw'
                logger.info("Timetable manage raw fallback slots loaded=%s", len(entries))
        except Exception:
            logger.exception("Failed to load paginated timetable rows")
            entries = []

        # Debug preview of last import
        skipped_preview = None
        try:
            preview_path = os.path.join(os.path.dirname(__file__), "uploads", "last_import_debug.json")
            if os.path.exists(preview_path):
                with open(preview_path, "r", encoding="utf-8") as f:
                    skipped_preview = f.read()
        except Exception:
            logger.exception("Failed to load skipped preview")
        # Debug counts
        try:
            raw_count_row = _db_execute(db, "SELECT COUNT(*) AS c FROM timetable_slots").fetchone()
            raw_count = int(raw_count_row[0] if raw_count_row is not None else 0)
        except Exception:
            raw_count = 0
        try:
            norm_count_row = _db_execute(db, "SELECT COUNT(*) AS c FROM timetable_entries").fetchone()
            normalized_count = int(norm_count_row[0] if norm_count_row is not None else 0)
        except Exception:
            normalized_count = 0

        def _row_to_dict(value):
            try:
                if isinstance(value, dict):
                    return dict(value)
                if hasattr(value, "keys"):
                    return {k: value[k] for k in value.keys()}
                return dict(value)
            except Exception:
                return {}

        def _has_visible_data(value):
            data = _row_to_dict(value)
            for key in ("id", "created_at"):
                data.pop(key, None)
            semantic_keys = ("day", "start_time", "end_time", "branch", "section", "semester", "subject_name", "faculty_name", "room")
            return any(_clean_text(data.get(key)) for key in semantic_keys)

        visible_entries = []
        for row in entries:
            row_dict = _row_to_dict(row)
            if _has_visible_data(row_dict):
                visible_entries.append(row_dict)
        if len(visible_entries) != len(entries):
            logger.info(
                "Filtered empty timetable rows for manage view kept=%s dropped=%s",
                len(visible_entries),
                len(entries) - len(visible_entries),
            )
        entries = visible_entries
        rows = list(entries)

        logger.info("Timetable manage fetched %s visible normalized rows", len(entries))
        logger.info("Timetable manage first 5 rows: %s", entries[:5])

        logger.debug("Timetable manage page loaded: normalized_count=%s, raw_count=%s, rows_source=%s", normalized_count, raw_count, rows_source)

        normalized_rows = rows

        # Expose pagination context
        page_size = PAGE_SIZE
        return render_template(
            "timetable_manage.html",
            rows=rows,
            entries=entries,
            skipped_preview=skipped_preview,
            rows_source=rows_source,
            raw_count=raw_count,
            normalized_count=normalized_count,
            normalized_rows=normalized_rows,
            success_banner=success_banner,
            last_imported_file=last_imported_file,
            total_count=total_count,
            page=page,
            page_size=page_size,
            q=q,
            preview_mode=preview_mode,
            preview_rows=preview_rows,
            unique_subjects=unique_subjects,
            unique_teachers=unique_teachers,
            unique_sections=unique_sections,
            validation_report=validation_report,
        )

    @app.route("/timetable/active")
    def timetable_active():
        # Returns the current active class for the logged-in teacher or for query params
        branch = request.args.get("branch") or session.get("teacher_branch_name") or ""
        section = request.args.get("section") or session.get("teacher_section") or ""
        db = get_db()
        ensure_timetable_tables(db)
        row = get_current_slot(db, branch, section)
        if not row:
            return jsonify({"active": False})
        return jsonify({
            "active": True,
            "slot": _row_to_dict(row)
        })

    @app.route("/timetable/admin/bulk_resolve", methods=("POST",))
    def timetable_admin_bulk_resolve():
        if session.get("role") != "admin":
            return redirect(url_for("dashboard"))
        db = get_db()
        create_missing = request.form.get("create_missing_subjects") in ("1", "true", "True")
        summary = {
            "subjects_checked": 0,
            "subjects_created": 0,
            "subjects_mapped": 0,
            "teachers_checked": 0,
            "teachers_mapped": 0,
        }
        try:
            # Gather distinct subject names
            rows = _db_execute(db, "SELECT DISTINCT TRIM(subject_name) AS subject_name FROM timetable_slots WHERE subject_name IS NOT NULL").fetchall()
            slot_subjects = [r[0] if not hasattr(r, 'keys') else r.get('subject_name') for r in rows if r and (r[0] if not hasattr(r, 'keys') else r.get('subject_name'))]
            rows = _db_execute(db, "SELECT DISTINCT TRIM(subject_name) AS subject_name FROM timetable_entries WHERE subject_name IS NOT NULL").fetchall()
            entry_subjects = [r[0] if not hasattr(r, 'keys') else r.get('subject_name') for r in rows if r and (r[0] if not hasattr(r, 'keys') else r.get('subject_name'))]
            distinct = sorted(set([_clean_text(s) for s in (slot_subjects + entry_subjects) if s]))

            existing = _db_execute(db, "SELECT id, name FROM subjects").fetchall()
            existing_map = { (r[1].lower() if not hasattr(r, 'keys') else r.get('name').lower()): (r[0] if not hasattr(r, 'keys') else r.get('id')) for r in existing }
            existing_names = list(existing_map.keys())

            # default branch for created subjects
            row = _db_execute(db, "SELECT id FROM branches ORDER BY id LIMIT 1").fetchone()
            default_branch_id = (row[0] if not hasattr(row, 'keys') else row.get('id')) if row else None

            for name in distinct:
                if not name:
                    continue
                summary["subjects_checked"] += 1
                lname = name.lower()
                if lname in existing_map:
                    summary["subjects_mapped"] += 1
                    continue
                # fuzzy match to existing subjects
                match = difflib.get_close_matches(name, existing_names, n=1, cutoff=0.8)
                if match:
                    matched_lower = match[0]
                    existing_map[lname] = existing_map[matched_lower]
                    summary["subjects_mapped"] += 1
                    continue
                if create_missing:
                    display = name.title()
                    try:
                        _db_execute(db, "INSERT INTO subjects (name, branch_id) VALUES (%s, %s)", (display, default_branch_id))
                        try:
                            db.commit()
                        except Exception:
                            pass
                        row = _db_execute(db, "SELECT id FROM subjects WHERE name = %s LIMIT 1", (display,)).fetchone()
                        sid = (row[0] if not hasattr(row, 'keys') else row.get('id')) if row else None
                        if sid:
                            existing_map[lname] = sid
                            existing_names.append(lname)
                            summary["subjects_created"] += 1
                    except Exception:
                        logger.exception("failed to create subject %s", display)

            # Apply mappings: update timetable_entries subject_id where subject_name matches
            for raw_lower, sid in list(existing_map.items()):
                if not sid:
                    continue
                try:
                    _db_execute(db, "UPDATE timetable_entries SET subject_id = %s WHERE subject_id IS NULL AND LOWER(TRIM(subject_name)) = LOWER(TRIM(%s))", (sid, raw_lower))
                except Exception:
                    logger.exception("failed to update timetable_entries for subject %s", raw_lower)

            # Teachers: normalize names and attempt to map
            rows = _db_execute(db, "SELECT DISTINCT TRIM(faculty_name) AS faculty_name FROM timetable_slots WHERE faculty_name IS NOT NULL").fetchall()
            slot_teachers = [r[0] if not hasattr(r, 'keys') else r.get('faculty_name') for r in rows if r and (r[0] if not hasattr(r, 'keys') else r.get('faculty_name'))]
            rows = _db_execute(db, "SELECT DISTINCT TRIM(faculty_name) AS faculty_name FROM timetable_entries WHERE faculty_name IS NOT NULL").fetchall()
            entry_teachers = [r[0] if not hasattr(r, 'keys') else r.get('faculty_name') for r in rows if r and (r[0] if not hasattr(r, 'keys') else r.get('faculty_name'))]
            tdistinct = sorted(set([_clean_text(t) for t in (slot_teachers + entry_teachers) if t]))
            existing_t = _db_execute(db, "SELECT id, name FROM teachers").fetchall()
            existing_t_map = { (r[1].lower() if not hasattr(r, 'keys') else r.get('name').lower()): (r[0] if not hasattr(r, 'keys') else r.get('id')) for r in existing_t }
            existing_t_names = list(existing_t_map.keys())

            def normalize_teacher_name(n):
                if not n:
                    return ""
                t = re.sub(r'^(mr|ms|mrs|dr)\.?\s+', '', n, flags=re.I)
                t = re.sub(r'[^\w\s]', ' ', t)
                return ' '.join(t.split()).strip()

            for tname in tdistinct:
                if not tname:
                    continue
                summary["teachers_checked"] += 1
                norm = normalize_teacher_name(tname)
                ln = norm.lower()
                if ln in existing_t_map:
                    summary["teachers_mapped"] += 1
                    tid = existing_t_map[ln]
                    try:
                        _db_execute(db, "UPDATE timetable_entries SET teacher_id = %s WHERE teacher_id IS NULL AND LOWER(TRIM(faculty_name)) = LOWER(TRIM(%s))", (tid, tname))
                    except Exception:
                        logger.exception("failed to update teacher mapping for %s", tname)
                    continue
                match = difflib.get_close_matches(norm, existing_t_names, n=1, cutoff=0.85)
                if match:
                    matched = match[0]
                    tid = existing_t_map[matched]
                    summary["teachers_mapped"] += 1
                    try:
                        _db_execute(db, "UPDATE timetable_entries SET teacher_id = %s WHERE teacher_id IS NULL AND LOWER(TRIM(faculty_name)) = LOWER(TRIM(%s))", (tid, tname))
                    except Exception:
                        logger.exception("failed to update teacher mapping for %s", tname)

            try:
                db.commit()
            except Exception:
                pass
            flash(f"Bulk resolve completed: {summary}", "success")
        except Exception as e:
            logger.exception("bulk_resolve failed")
            flash(f"Bulk resolve failed: {e}", "error")
        return redirect(url_for("timetable_manage"))

    @app.route("/timetable/faculty-schedules")
    def timetable_faculty_schedules():
        if session.get("role") not in ("admin", "teacher"):
            return redirect(url_for("login"))
        db = get_db()
        ensure_timetable_tables(db)
        teacher_id = session.get("teacher_id") if session.get("role") == "teacher" else None
        schedules = []
        if teacher_id:
            schedules = get_faculty_schedule(db, teacher_id)
        else:
            try:
                rows = _db_execute(
                    db,
                    "SELECT te.*, s.name AS subject_name, t.name AS teacher_name, b.name AS branch_name FROM timetable_entries te LEFT JOIN subjects s ON te.subject_id = s.id LEFT JOIN teachers t ON te.teacher_id = t.id LEFT JOIN branches b ON te.branch_id = b.id ORDER BY te.day, te.start_time",
                ).fetchall()
                schedules = [_row_to_dict(r) for r in rows]
            except Exception:
                schedules = []
        return render_template("faculty_schedules.html", schedules=schedules, teacher_id=teacher_id)

    @app.route('/timetable/students')
    def timetable_students():
        # Return students for a given branch and section (or teacher session)
        branch = request.args.get('branch') or session.get('teacher_branch_name') or ''
        section = request.args.get('section') or session.get('teacher_section') or ''
        db = get_db()
        try:
            rows = _db_execute(
                db,
                "SELECT s.id, s.name, s.roll_no FROM students s JOIN branches b ON s.branch_id = b.id WHERE b.name = ? AND s.section = ? ORDER BY s.name",
                (branch, section),
            ).fetchall()
            students = [{"id": r["id"], "name": r["name"], "roll_no": r["roll_no"]} for r in rows]
            return jsonify({"ok": True, "students": students})
        except Exception as e:
            logger.exception('failed to fetch students')
            return jsonify({"ok": False, "error": str(e)}), 500

    @app.route("/attendance/mark_current", methods=("POST",))
    def attendance_mark_current():
        # Mark attendance for current slot (teacher must be logged in)
        if session.get("role") != "teacher":
            return jsonify({"ok": False, "error": "unauthorized"}), 403
        data = request.get_json() or {}
        action = data.get("action")  # present / absent / bulk_absent
        branch = session.get("teacher_branch_name") or request.args.get("branch")
        section = session.get("teacher_section") or request.args.get("section")
        db = get_db()
        slot = get_current_slot(db, branch, section)
        if not slot:
            return jsonify({"ok": False, "error": "no active slot"}), 400

        # Determine subject_id / teacher_id / branch_id from normalized entry when present
        try:
            # slot may be sqlite3.Row or dict-like
            subject_id = slot.get("subject_id") if isinstance(slot, dict) else (slot["subject_id"] if "subject_id" in slot.keys() else None)
        except Exception:
            subject_id = None

        # Do not resolve subject_id from subjects table; rely on timetable_entries only

        # Determine branch_id and students list
        try:
            branch_id = slot.get("branch_id") if isinstance(slot, dict) else (slot["branch_id"] if "branch_id" in slot.keys() else None)
        except Exception:
            branch_id = None

        if branch_id:
            students = _db_execute(db, "SELECT id, name FROM students WHERE branch_id = ? AND (COALESCE(section, '') = COALESCE(?, '')) ORDER BY name", (branch_id, section or "")).fetchall()
        else:
            # fallback to previous name-based lookup
            students = _db_execute(db, "SELECT s.id, s.name FROM students s JOIN branches b ON s.branch_id = b.id WHERE b.name = ? AND s.section = ? ORDER BY s.name", (branch, section)).fetchall()

        # Use today's date and default period=1 for current slot marking
        from datetime import date as _d
        today_str = _d.today().isoformat()
        period = 1

        # Prevent duplicate attendance for the same subject today+period
        duplicate_check = False
        try:
            if subject_id:
                dup_row = _db_execute(db, "SELECT COUNT(1) AS c FROM attendance WHERE subject_id = ? AND date = ? AND period = ?", (subject_id, today_str, period)).fetchone()
                dup_count = dup_row[0] if dup_row is not None else 0
                if int(dup_count) > 0:
                    duplicate_check = True
        except Exception:
            duplicate_check = False

        if duplicate_check:
            return jsonify({"ok": False, "error": "attendance already recorded for this subject today"}), 400

        teacher_id = session.get("teacher_id")

        if action == "bulk_absent":
            marked = 0
            for s in students:
                sid = s[0] if not isinstance(s, dict) else s.get("id")
                try:
                    _db_execute(
                        db,
                        "INSERT INTO attendance (student_id, branch_id, branch_section, section, subject_id, teacher_id, subject_name, date, period, status) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                        (sid, branch_id or None, None, section or None, subject_id or None, teacher_id or None, (slot.get("subject_name") if isinstance(slot, dict) else None), today_str, period, "Absent"),
                    )
                    marked += 1
                except Exception:
                    pass
            try:
                db.commit()
            except Exception:
                pass
            return jsonify({"ok": True, "marked": marked})
        elif action in ("present", "absent"):
            student_ids = data.get("student_ids", [])
            marked = 0
            for sid in student_ids:
                st = "Present" if action == "present" else "Absent"
                try:
                    _db_execute(
                        db,
                        "INSERT INTO attendance (student_id, branch_id, branch_section, section, subject_id, teacher_id, subject_name, date, period, status) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                        (sid, branch_id or None, None, section or None, subject_id or None, teacher_id or None, (slot.get("subject_name") if isinstance(slot, dict) else None), today_str, period, st),
                    )
                    marked += 1
                except Exception:
                    pass
            try:
                db.commit()
            except Exception:
                pass
            return jsonify({"ok": True, "marked": marked})
        else:
            return jsonify({"ok": False, "error": "unknown action"}), 400

    logger.info("Timetable routes registered")

    @app.route("/health/timetable")
    def timetable_health():
        """Health-check for timetable initialization and DB compatibility.

        Access: admin users OR non-production/debug OR Render internal hosts.
        Returns JSON describing status, counts, and postgres compatibility.
        """
        # Access control: allow only admin or non-prod/debug or Render internal
        try:
            is_admin = bool(session.get("role") == "admin")
        except Exception:
            is_admin = False

        allow_internal = bool(app.debug or os.environ.get("RENDER") or os.environ.get("RENDER_INTERNAL_HOSTNAME"))
        if not (is_admin or allow_internal):
            return jsonify({"ok": False, "error": "unauthorized"}), 403

        result = {"ok": False, "postgres_compatible": False, "tables": {}, "messages": []}
        db = None
        try:
            if db_getter is None:
                raise RuntimeError("Database getter is not configured")
            db = db_getter()

            # Verify connectivity
            try:
                # simple query to ensure connection is alive
                cur = _db_execute(db, "SELECT 1")
                _ = cur.fetchone()
                result["messages"].append("db_connectivity_ok")
            except Exception:
                result["messages"].append("db_connectivity_failed")
                raise

            # Ensure tables are present (this may create them)
            ensure_timetable_tables(db)
            result["messages"].append("ensure_timetable_tables_executed")

            # Postgres compatibility flag
            pg_ok = _is_postgres_db(db)
            result["postgres_compatible"] = pg_ok
            if pg_ok:
                result["messages"].append("postgres_compatibility_detected")

            # Counts
            try:
                r1 = _db_execute(db, "SELECT COUNT(1) AS c FROM timetable_slots").fetchone()
                r2 = _db_execute(db, "SELECT COUNT(1) AS c FROM timetable_entries").fetchone()
                slots_count = int(r1[0] if r1 is not None and r1[0] is not None else 0)
                entries_count = int(r2[0] if r2 is not None and r2[0] is not None else 0)
                result["tables"]["timetable_slots"] = slots_count
                result["tables"]["timetable_entries"] = entries_count
                result["messages"].append("counts_retrieved")
                diagnostics = _table_diagnostics(db)
                result["diagnostics"] = diagnostics
                result["messages"].append("diagnostics_retrieved")
            except Exception:
                # If counts failed, attempt information_schema check for Postgres
                try:
                    if _is_postgres_db(db):
                        q = "SELECT to_regclass('public.timetable_slots') IS NOT NULL AS slots_exists, to_regclass('public.timetable_entries') IS NOT NULL AS entries_exists"
                        rr = _db_execute(db, q).fetchone()
                        result["tables"]["timetable_slots_exists"] = bool(rr[0])
                        result["tables"]["timetable_entries_exists"] = bool(rr[1])
                        result["messages"].append("schema_presence_checked")
                except Exception:
                    result["messages"].append("counts_unavailable")

            result["ok"] = True
            result["messages"].append("timetable_ok")
            # Log summary server-side
            logger.info("Timetable tables verified")
            if result.get("tables"):
                logger.info(f"Timetable schema initialized: slots={result['tables'].get('timetable_slots')} entries={result['tables'].get('timetable_entries')}")
            if result.get("diagnostics"):
                logger.info(
                    "Timetable diagnostics: slots_dup_groups=%s slots_dup_rows=%s entries_dup_groups=%s entries_dup_rows=%s",
                    result["diagnostics"].get("slots_duplicate_groups"),
                    result["diagnostics"].get("slots_duplicate_rows"),
                    result["diagnostics"].get("entries_duplicate_groups"),
                    result["diagnostics"].get("entries_duplicate_rows"),
                )
                logger.info("Timetable null diagnostics: slots=%s entries=%s", result["diagnostics"].get("slots_nulls"), result["diagnostics"].get("entries_nulls"))
            if pg_ok:
                logger.info("PostgreSQL timetable compatibility OK")

            return jsonify(result)
        except Exception as e:
            logger.exception("timetable health check failed")
            msg = {"ok": False, "error": str(type(e).__name__) + ": " + str(e)}
            # include non-sensitive diagnostics
            msg.update({k: v for k, v in result.items() if k in ("postgres_compatible", "tables", "messages")})
            return jsonify(msg), 500
        finally:
            if db:
                try:
                    db.close()
                except Exception:
                    pass


# Register when imported into app.py
try:
    from flask import session
    # plugin-style: if app is imported as module, require explicit registration
except Exception:
    pass
