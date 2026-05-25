"""Split combined timetable DOCX/PDF into per-section files and import them.

Saves per-section files in uploads/timetable_<SECTION>.docx and runs the
existing timetable parser/import pipeline for each.

Produces:
- uploads/last_import_debug.jsonl (appendable JSON-lines with diagnostics)
- uploads/last_import_summary.json (summary of all sections)

Usage: run from project root: .venv\Scripts\python.exe scripts\split_and_import_timtables.py
"""
import os
import re
import json
import time
import logging
from collections import defaultdict

try:
    import docx
except Exception:
    docx = None

try:
    import pdfplumber
except Exception:
    pdfplumber = None

# ensure project import path
PR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if PR not in os.sys.path:
    os.sys.path.insert(0, PR)

from app import get_db
import timetable

logger = logging.getLogger("scripts.split_import")
logging.basicConfig(level=logging.INFO)

UPLOADS = os.path.join(PR, "uploads")
os.makedirs(UPLOADS, exist_ok=True)

SECTION_RE = re.compile(r"\b([A-Z]{2,6}(?:[-\s]?[A-Z0-9]{1,3})?)\b")
COMMON_SECTIONS = {"CSE", "ECE", "MECH", "CIVIL", "CSM", "PRINCIPAL", "VICE"}

SUMMARY_PATH = os.path.join(UPLOADS, "last_import_summary.json")
DEBUG_PATH = os.path.join(UPLOADS, "last_import_debug.jsonl")


def normalize_section_name(s: str) -> str:
    s = s.strip().upper().replace(" ", "-")
    s = re.sub(r"[^A-Z0-9\-]", "", s)
    s = re.sub(r"-+", "-", s)
    return s


def extract_sections_from_docx(path: str):
    if docx is None:
        raise RuntimeError("python-docx not available")
    doc = docx.Document(path)
    sections = defaultdict(lambda: {"headings": [], "tables": []})

    # iterate body elements to preserve sequence of paragraphs and tables
    table_idx = 0
    current_section = None
    for child in doc.element.body:
        tag = child.tag
        if tag.endswith('}p'):
            # paragraph
            para = docx.text.paragraph.Paragraph(child, doc)
            text = para.text.strip()
            if not text:
                continue
            m = SECTION_RE.search(text)
            if m:
                name = normalize_section_name(m.group(1))
                # prefer names that include known section tokens or patterns like 'CSE-A'
                sections[name]["headings"].append(text)
                current_section = name
                continue
            # if no section pattern, keep current_section unchanged
        elif tag.endswith('}tbl'):
            # table element corresponds to doc.tables[table_idx]
            try:
                tbl = doc.tables[table_idx]
            except Exception:
                tbl = None
            table_idx += 1
            if current_section is None:
                # try infer from first row first cell
                inferred = None
                try:
                    cell0 = tbl.rows[0].cells[0].text.strip()
                    m2 = SECTION_RE.search(cell0)
                    if m2:
                        inferred = normalize_section_name(m2.group(1))
                except Exception:
                    inferred = None
                if inferred:
                    current_section = inferred
                    sections[current_section]["headings"].append(cell0[:100])
                else:
                    current_section = "UNASSIGNED"
            sections[current_section]["tables"].append(tbl)
    return sections


def write_section_docx(section_name: str, heading_texts, tables, dest_path: str):
    # create a new docx and copy headings and tables
    if docx is None:
        raise RuntimeError("python-docx not available")
    newd = docx.Document()
    # add heading(s)
    for h in heading_texts or []:
        newd.add_paragraph(h)
    # copy tables
    for t in tables:
        # create table with same dimensions
        rows = len(t.rows)
        cols = len(t.columns)
        nt = newd.add_table(rows=rows, cols=cols)
        for r in range(rows):
            for c in range(cols):
                try:
                    nt.rows[r].cells[c].text = t.rows[r].cells[c].text
                except Exception:
                    nt.rows[r].cells[c].text = ""
        newd.add_paragraph()  # spacer
    newd.save(dest_path)


def split_pdf_to_docx_sections(path: str):
    if pdfplumber is None:
        raise RuntimeError("pdfplumber not available")
    text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text.append(page.extract_text() or "")
    full = "\n".join(text)
    # split by headings matching SECTION_RE
    parts = re.split(r"(^.*$)", full, flags=re.M)
    # naive approach: find lines matching SECTION_RE and split by them
    lines = full.splitlines()
    sections = defaultdict(lambda: {"headings": [], "text": []})
    current = None
    for ln in lines:
        if SECTION_RE.search(ln):
            name = normalize_section_name(SECTION_RE.search(ln).group(1))
            current = name
            sections[current]["headings"].append(ln.strip())
            continue
        if current is None:
            current = "UNASSIGNED"
        sections[current]["text"].append(ln)
    # write each section as a docx with text content
    out_paths = []
    for sec, data in sections.items():
        fname = f"timetable_{sec}.docx"
        dest = os.path.join(UPLOADS, fname)
        if docx is None:
            # fallback to writing plain text file
            with open(dest + ".txt", "w", encoding="utf-8") as f:
                f.write("\n".join(data.get("headings", []) + data.get("text", [])))
            out_paths.append(dest + ".txt")
            continue
        newd = docx.Document()
        for h in data.get("headings", []):
            newd.add_paragraph(h)
        for ln in data.get("text", []):
            newd.add_paragraph(ln)
        newd.save(dest)
        out_paths.append(dest)
    return out_paths


def run_import_on_file(filepath: str, db):
    name = os.path.basename(filepath)
    ext = os.path.splitext(name)[1].lower()
    try:
        if ext == ".docx" and docx is not None:
            slots = timetable.parse_docx_table(filepath)
            if not slots:
                slots = timetable.parse_docx_grid(filepath)
        elif ext == ".pdf" and pdfplumber is not None:
            slots = timetable.parse_pdf_to_slots(filepath)
        else:
            logger.warning("Unsupported file type for import: %s", filepath)
            return {"file": filepath, "error": "unsupported"}
    except Exception as e:
        logger.exception("Parsing failed for %s", filepath)
        return {"file": filepath, "error": repr(e)}

    logger.info("Parsed %d slots from %s", len(slots), filepath)

    # call import functions
    ins = timetable.import_slots(db, slots)
    norm = timetable.import_slots_normalized(db, slots)
    # aggregate
    result = {"file": filepath, "parsed": len(slots), "inserted": ins.get("counters", {}).get("inserted", 0), "skipped": ins.get("counters", {}).get("skipped_total", 0), "normalized_inserted": norm.get("counters", {}).get("inserted", 0), "normalized_skipped": norm.get("counters", {}).get("skipped_total", 0), "ins": ins, "norm": norm}
    return result


def main():
    # find candidate combined files
    candidates = []
    for fn in os.listdir(UPLOADS):
        if fn.lower().startswith("timetable_"):
            continue
        if fn.lower().endswith(".docx") or fn.lower().endswith('.pdf'):
            candidates.append(os.path.join(UPLOADS, fn))
    if not candidates:
        print("No candidate DOCX/PDF files found in uploads/ to split. Place combined file(s) in uploads/ and re-run.")
        return

    db = get_db()
    timetable.ensure_timetable_tables(db)

    overall = {"sections": {}, "files_processed": [], "total_parsed": 0, "total_inserted": 0}

    # open debug output for appending
    with open(DEBUG_PATH, "a", encoding="utf-8") as debugf:
        for c in candidates:
            print("Processing candidate:", c)
            if c.lower().endswith('.docx'):
                try:
                    sections = extract_sections_from_docx(c)
                except Exception as e:
                    logger.exception("Failed to extract sections from %s", c)
                    continue
                # write per-section docx files and import
                for sec, data in sections.items():
                    sec_name = sec if sec else "UNASSIGNED"
                    safe_name = f"timetable_{sec_name}.docx"
                    dest = os.path.join(UPLOADS, safe_name)
                    write_section_docx(sec_name, data.get('headings', []), data.get('tables', []), dest)
                    res = run_import_on_file(dest, db)
                    overall['sections'][sec_name] = res
                    overall['files_processed'].append(dest)
                    overall['total_parsed'] += res.get('parsed', 0) or 0
                    overall['total_inserted'] += res.get('inserted', 0) or 0
                    # write debug per-section
                    debugf.write(json.dumps({"section": sec_name, "result": res}, default=str) + "\n")
            elif c.lower().endswith('.pdf'):
                try:
                    out_paths = split_pdf_to_docx_sections(c)
                except Exception as e:
                    logger.exception("Failed to split PDF %s", c)
                    continue
                for dest in out_paths:
                    res = run_import_on_file(dest, db)
                    sec_name = os.path.splitext(os.path.basename(dest))[0].replace('timetable_', '')
                    overall['sections'][sec_name] = res
                    overall['files_processed'].append(dest)
                    overall['total_parsed'] += res.get('parsed', 0) or 0
                    overall['total_inserted'] += res.get('inserted', 0) or 0
                    debugf.write(json.dumps({"section": sec_name, "result": res}, default=str) + "\n")

    # write summary
    overall['timestamp'] = time.time()
    with open(SUMMARY_PATH, 'w', encoding='utf-8') as sf:
        json.dump(overall, sf, indent=2, default=str)

    print("Split and import completed. Summary written to", SUMMARY_PATH)


if __name__ == '__main__':
    main()
